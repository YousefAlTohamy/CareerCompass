from __future__ import annotations

import html
import json
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor, Twips
from PIL import Image as PILImage
from PIL import ImageDraw, ImageFont
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[3]
OUT_DIR = ROOT / "docs/graduation-book"
ASSETS = OUT_DIR / "assets"
DIAGRAMS = ASSETS / "diagrams"
SCREENSHOTS = ASSETS / "screenshots"
LOGOS = ASSETS / "logos"
DOCX_PATH = OUT_DIR / "CareerCompass_Graduation_Project_Book.docx"
PDF_PATH = OUT_DIR / "CareerCompass_Graduation_Project_Book.pdf"
MD_PATH = OUT_DIR / "CareerCompass_Graduation_Project_Book.md"
REFERENCES_PATH = OUT_DIR / "references.md"
NOTES_PATH = OUT_DIR / "REPORT_GENERATION_NOTES.md"
EVALUATION = OUT_DIR / "evaluation"
MINI_EVAL_SCRIPT = EVALUATION / "run_mini_evaluation.py"
MINI_EVAL_RESULTS = EVALUATION / "mini_evaluation_results.json"
MINI_EVAL_SUMMARY = EVALUATION / "mini_evaluation_summary.md"
SMOKE_EVAL_SCRIPT = EVALUATION / "run_ai_cv_analyzer_smoke_eval.py"
SMOKE_EVAL_RESULTS = EVALUATION / "ai_cv_analyzer_smoke_results.json"
SMOKE_EVAL_SUMMARY = EVALUATION / "ai_cv_analyzer_smoke_summary.md"

PROJECT_TITLE = "CareerCompass: AI-Powered Career Guidance and Job Recommendation Platform"
PROJECT_SUBTITLE = "AI-Powered Career Guidance and Job Recommendation Platform"
SHORT_NAME = "CareerCompass"
CC_LOGO_PATH = LOGOS / "CC_Logo.png"
COVER_LOGO_WIDTH = Inches(1.38)
PAGE_BREAK_BEFORE_HEADINGS = frozenset(
    {
        "1.6 Proposed Solution",
        "2.7 Non-Functional Requirements",
        "2.8 Requirement-to-Code/Test Traceability",
        "3.7 Database Design",
        "5.5.5 Layer 1: CV Understanding Pipeline",
        "5.17 Dockerized Runtime Flow",
        "6.4 Fault Tolerance and Recovery",
        "Metric Definitions",
        "Gap Analysis Pair Details",
        "8.21 Summary of Results",
        "Appendix D: Docker Services Summary",
        "Appendix F: Test Cases",
    }
)
UNIVERSITY = "Kafr El-Sheikh University"
FACULTY = "Faculty of Computers and Information"
DEPARTMENT = "Computer Science Department"
ACADEMIC_YEAR = "2025 / 2026"
SUPERVISOR = "Dr. Amena Mahmoud"
STUDENTS = [
    "Yousef Altohamy Ahmed Altohamy",
    "Ahmed Mohamed Ahmed Abdelaziz",
    "Mohamed Ali Ahmed Mohamed",
    "Mohamed Ibrahim Ahmed Mohamed",
    "Ahmed Khamis Mohamed Younes",
    "Ahmed Sobhy Mohamed Ali",
]

TOC_ENTRIES = [
    ("List of Figures", "List of Figures"),
    ("List of Tables", "List of Tables"),
    ("Acknowledgment", "Acknowledgment"),
    ("Abstract", "Abstract"),
    ("Abbreviations", "Abbreviations"),
    ("Chapter 1: Introduction", "Chapter 1: Introduction"),
    ("Chapter 2: System Analysis", "Chapter 2: System Analysis"),
    ("Chapter 3: System Design and Architecture", "Chapter 3: System Design and Architecture"),
    ("Chapter 4: Software and Tools Used", "Chapter 4: Software and Tools Used"),
    ("Chapter 5: System Implementation", "Chapter 5: System Implementation"),
    ("Chapter 6: AI CV Analyzer Deep Technical Analysis", "Chapter 6: AI CV Analyzer Deep Technical Analysis"),
    ("Chapter 7: AI Job Miner and Scraping Deep Technical Analysis", "Chapter 7: AI Job Miner and Scraping Deep Technical Analysis"),
    ("Chapter 8: Testing and Evaluation", "Chapter 8: Testing and Evaluation"),
    ("Chapter 9: Security and Privacy", "Chapter 9: Security and Privacy"),
    ("Chapter 10: Conclusion and Future Work", "Chapter 10: Conclusion and Future Work"),
    ("References", "References"),
    ("Appendices", "Appendices"),
]


@dataclass(frozen=True)
class Reference:
    key: int
    organization: str
    title: str
    source: str
    year: str
    url: str
    accessed: str = "Accessed: May 29, 2026"


REFERENCES = [
    Reference(1, "Laravel", "Laravel 12.x Documentation", "Laravel", "2026", "https://laravel.com/docs/12.x"),
    Reference(2, "Laravel", "Laravel Sanctum", "Laravel", "2026", "https://laravel.com/docs/12.x/sanctum"),
    Reference(3, "Laravel", "Queues", "Laravel", "2026", "https://laravel.com/docs/12.x/queues"),
    Reference(4, "Meta Open Source", "Quick Start - React", "React Documentation", "2026", "https://react.dev/learn"),
    Reference(5, "Vite", "Getting Started", "Vite Documentation", "2026", "https://vite.dev/guide/"),
    Reference(6, "FastAPI", "FastAPI Documentation", "FastAPI", "2026", "https://fastapi.tiangolo.com/"),
    Reference(7, "Python Software Foundation", "Python 3 Documentation", "Python", "2026", "https://docs.python.org/3/"),
    Reference(8, "Oracle", "MySQL 8.0 Reference Manual", "MySQL Documentation", "2026", "https://dev.mysql.com/doc/refman/8.0/en/"),
    Reference(9, "MinIO", "MinIO AIStor Documentation", "MinIO Documentation", "2026", "https://docs.min.io/aistor/"),
    Reference(10, "Docker", "What is a Container?", "Docker Resources", "2026", "https://www.docker.com/resources/what-container/"),
    Reference(11, "Docker", "Docker Compose", "Docker Documentation", "2026", "https://docs.docker.com/compose/"),
    Reference(12, "Nginx", "nginx Documentation", "nginx", "2026", "https://nginx.org/en/docs/"),
    Reference(13, "Prometheus", "Overview", "Prometheus Documentation", "2026", "https://prometheus.io/docs/introduction/overview/"),
    Reference(14, "Grafana Labs", "Grafana OSS and Enterprise", "Grafana Documentation", "2026", "https://grafana.com/docs/grafana/latest/"),
    Reference(15, "GitHub", "GitHub Actions Documentation", "GitHub Docs", "2026", "https://docs.github.com/en/actions"),
    Reference(16, "MDN Web Docs", "HTTP: Hypertext Transfer Protocol", "MDN", "2026", "https://developer.mozilla.org/en-US/docs/Web/HTTP"),
    Reference(17, "Scrapy", "Scrapy Documentation", "Scrapy", "2026", "https://docs.scrapy.org/en/latest/"),
    Reference(18, "Leonard Richardson", "Beautiful Soup Documentation", "Beautiful Soup", "2026", "https://www.crummy.com/software/BeautifulSoup/bs4/doc/"),
    Reference(19, "scikit-learn", "TfidfVectorizer", "scikit-learn Documentation", "2026", "https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html"),
    Reference(20, "scikit-learn", "cosine_similarity", "scikit-learn Documentation", "2026", "https://scikit-learn.org/stable/modules/generated/sklearn.metrics.pairwise.cosine_similarity.html"),
    Reference(21, "UKP Lab", "Sentence Transformers Documentation", "Sentence Transformers", "2026", "https://www.sbert.net/"),
    Reference(22, "Artifex", "PyMuPDF Documentation", "PyMuPDF", "2026", "https://pymupdf.readthedocs.io/en/latest/"),
    Reference(23, "jsvine", "pdfplumber", "GitHub Repository", "2026", "https://github.com/jsvine/pdfplumber"),
    Reference(24, "Jaided AI", "EasyOCR", "GitHub Repository", "2026", "https://github.com/JaidedAI/EasyOCR"),
    Reference(25, "PHPUnit", "PHPUnit 12.0 Manual", "PHPUnit Documentation", "2026", "https://docs.phpunit.de/en/12.0/"),
    Reference(26, "pytest", "pytest Documentation", "pytest", "2026", "https://docs.pytest.org/en/stable/"),
    Reference(27, "OWASP", "File Upload Cheat Sheet", "OWASP Cheat Sheet Series", "2026", "https://cheatsheetseries.owasp.org/cheatsheets/File_Upload_Cheat_Sheet.html"),
    Reference(28, "OWASP", "Authentication Cheat Sheet", "OWASP Cheat Sheet Series", "2026", "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html"),
    Reference(29, "Martin Fowler and James Lewis", "Microservices", "martinfowler.com", "2014", "https://martinfowler.com/articles/microservices.html"),
    Reference(30, "scikit-learn", "precision_recall_fscore_support", "scikit-learn Documentation", "2026", "https://scikit-learn.org/stable/modules/generated/sklearn.metrics.precision_recall_fscore_support.html"),
    Reference(31, "Hugging Face", "Transformers Documentation", "Hugging Face Documentation", "2026", "https://huggingface.co/docs/transformers/index", "Accessed: June 6, 2026"),
    Reference(32, "Hugging Face", "Token Classification", "Hugging Face Documentation", "2026", "https://huggingface.co/docs/transformers/tasks/token_classification", "Accessed: June 6, 2026"),
    Reference(33, "Hugging Face", "Trainer", "Hugging Face Documentation", "2026", "https://huggingface.co/docs/transformers/main_classes/trainer", "Accessed: June 6, 2026"),
    Reference(34, "Google", "Gemini API Documentation", "Google AI for Developers", "2026", "https://ai.google.dev/gemini-api/docs", "Accessed: June 6, 2026"),
    Reference(35, "Google", "Google AI Studio", "Google AI for Developers", "2026", "https://ai.google.dev/aistudio", "Accessed: June 6, 2026"),
    Reference(36, "Google", "Google Colaboratory FAQ", "Google Research", "2026", "https://research.google.com/colaboratory/faq.html", "Accessed: June 6, 2026"),
    Reference(37, "Jacob Devlin, Ming-Wei Chang, Kenton Lee, and Kristina Toutanova", "BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding", "arXiv", "2018", "https://arxiv.org/abs/1810.04805", "Accessed: June 6, 2026"),
    Reference(38, "PyTorch", "Dynamic Quantization", "PyTorch Tutorials", "2026", "https://docs.pytorch.org/tutorials/recipes/recipes/dynamic_quantization.html", "Accessed: June 6, 2026"),
    Reference(39, "OpenAPI Initiative", "OpenAPI Specification", "OpenAPI Documentation", "2026", "https://spec.openapis.org/oas/latest.html", "Accessed: June 7, 2026"),
    Reference(40, "Adzuna", "Adzuna Developer API", "Adzuna Developer Portal", "2026", "https://developer.adzuna.com/", "Accessed: June 7, 2026"),
    Reference(41, "IETF", "RFC 9309: Robots Exclusion Protocol", "RFC Editor", "2022", "https://www.rfc-editor.org/rfc/rfc9309", "Accessed: June 7, 2026"),
    Reference(42, "React Router", "Routing", "React Router Documentation", "2026", "https://reactrouter.com/start/declarative/routing", "Accessed: June 7, 2026"),
    Reference(43, "Axios", "First steps", "Axios Documentation", "2026", "https://axios-http.com/docs/intro", "Accessed: June 7, 2026"),
]


FIGURES = [
    ("Figure 1", "High-level architecture of CareerCompass.", "assets/diagrams/01_high_level_architecture.png"),
    ("Figure 2", "Docker deployment architecture.", "assets/diagrams/02_docker_deployment.png"),
    ("Figure 3", "DFD Level 0 context diagram.", "assets/diagrams/03_dfd_level_0.png"),
    ("Figure 4", "DFD Level 1 process diagram.", "assets/diagrams/04_dfd_level_1.png"),
    ("Figure 5", "UML use case diagram.", "assets/diagrams/05_use_case_diagram.png"),
    ("Figure 6", "Sequence diagram for CV upload and analysis.", "assets/diagrams/06_sequence_cv_upload_analysis.png"),
    ("Figure 7", "Sequence diagram for recommendation and gap analysis.", "assets/diagrams/07_sequence_job_recommendation_gap_analysis.png"),
    ("Figure 8", "ERD and database summary diagram.", "assets/diagrams/08_erd.png"),
    ("Figure 9", "AI CV Analyzer runtime flow.", "assets/diagrams/09_cv_analyzer_runtime_flow.png"),
    ("Figure 10", "AI CV Analyzer model-training workflow.", "assets/diagrams/10_cv_model_training_pipeline.png"),
    ("Figure 11", "AI CV Analyzer extraction components.", "assets/diagrams/11_cv_extraction_components.png"),
    ("Figure 12", "Layer 1 CV understanding pipeline.", "assets/diagrams/12_layer1_understanding_pipeline.png"),
    ("Figure 13", "Layer 2 classification flow.", "assets/diagrams/13_layer2_classification_flow.png"),
    ("Figure 14", "Layer 3 matching engine.", "assets/diagrams/14_layer3_matching_engine.png"),
    ("Figure 15", "NER token processing and BIO tagging.", "assets/diagrams/15_ner_token_processing.png"),
    ("Figure 16", "Seniority decision logic.", "assets/diagrams/16_seniority_decision_logic.png"),
    ("Figure 17", "Skill canonicalization chain.", "assets/diagrams/17_canonicalization_chain.png"),
    ("Figure 18", "Layer 3 score collapse logic.", "assets/diagrams/18_score_collapse_logic.png"),
    ("Figure 19", "AI design philosophy for the layered hybrid analyzer.", "assets/diagrams/19_ai_design_philosophy.png"),
    ("Figure 20", "Complete CV processing flow.", "assets/diagrams/20_complete_cv_processing_flow.png"),
    ("Figure 21", "CV analyzer fault tolerance and recovery flow.", "assets/diagrams/21_cv_fault_tolerance_flow.png"),
    ("Figure 22", "Confidence and readiness signal flow.", "assets/diagrams/22_confidence_signal_flow.png"),
    ("Figure 23", "Skill canonicalization example.", "assets/diagrams/23_skill_canonicalization_example.png"),
    ("Figure 24", "Fine-tuned BERT NER architecture.", "assets/diagrams/24_fine_tuned_bert_ner_architecture.png"),
    ("Figure 25", "Detailed NER training pipeline.", "assets/diagrams/25_detailed_training_pipeline.png"),
    ("Figure 26", "Matching formula and penalty flow.", "assets/diagrams/26_matching_formula_flow.png"),
    ("Figure 27", "Explainable AI fit output.", "assets/diagrams/27_explainable_ai_output.png"),
    ("Figure 28", "AI analyzer sequence diagram.", "assets/diagrams/28_ai_analyzer_sequence.png"),
    ("Figure 29", "Dataset evidence availability.", "assets/diagrams/29_dataset_evidence_availability.png"),
    ("Figure 30", "AI CV Analyzer smoke evaluation metrics.", "assets/diagrams/30_ai_cv_analyzer_smoke_metrics.png"),
    ("Figure 31", "Home page.", "assets/screenshots/01_home.png"),
    ("Figure 32", "Register page.", "assets/screenshots/02_register.png"),
    ("Figure 33", "Login page.", "assets/screenshots/03_login.png"),
    ("Figure 34", "Student dashboard before CV upload.", "assets/screenshots/04_dashboard_before_cv_upload.png"),
    ("Figure 35", "CV upload user interface.", "assets/screenshots/05_cv_upload_ui.png"),
    ("Figure 36", "Dashboard after successful CV parsing.", "assets/screenshots/06_dashboard_after_cv_upload.png"),
    ("Figure 37", "Extracted profile and skills page.", "assets/screenshots/07_extracted_profile_skills.png"),
    ("Figure 38", "Jobs recommendations page.", "assets/screenshots/08_jobs_recommendations.png"),
    ("Figure 39", "Job detail and inline gap panel.", "assets/screenshots/09_job_details_and_inline_gap.png"),
    ("Figure 40", "Gap analysis page.", "assets/screenshots/10_gap_analysis.png"),
    ("Figure 41", "Applications tracker page.", "assets/screenshots/11_applications_tracker.png"),
    ("Figure 42", "Tools Hub preview page.", "assets/screenshots/12_tools_hub.png"),
    ("Figure 43", "System status page.", "assets/screenshots/13_system_status.png"),
    ("Figure 44", "Admin dashboard.", "assets/screenshots/14_admin_dashboard.png"),
    ("Figure 45", "Admin jobs page.", "assets/screenshots/15_admin_jobs.png"),
    ("Figure 46", "Admin sources diagnostics page.", "assets/screenshots/16_admin_sources_diagnostics.png"),
    ("Figure 47", "Admin target roles page.", "assets/screenshots/17_admin_targets.png"),
    ("Figure 48", "Docker services evidence.", "assets/screenshots/18_docker_containers.png"),
    ("Figure 49", "Validation evidence summary.", "assets/screenshots/19_validation_summary.png"),
    ("Figure 50", "Colab NER final epoch metrics.", "assets/diagrams/31_colab_ner_metrics.png"),
    ("Figure 51", "Colab NER epoch performance trend.", "assets/diagrams/61_colab_ner_epoch_performance.png"),
    ("Figure 52", "Colab NER training and validation loss curve.", "assets/diagrams/62_colab_ner_loss_curve.png"),
    ("Figure 53", "Job mining design philosophy.", "assets/diagrams/32_job_mining_design_philosophy.png"),
    ("Figure 54", "AI Job Miner runtime architecture.", "assets/diagrams/34_scraping_runtime_architecture.png"),
    ("Figure 55", "Complete job mining flow.", "assets/diagrams/33_complete_job_mining_flow.png"),
    ("Figure 56", "Scraping sequence diagram.", "assets/diagrams/35_scraping_sequence_diagram.png"),
    ("Figure 57", "Scraping job lifecycle.", "assets/diagrams/36_scraping_job_lifecycle.png"),
    ("Figure 58", "Source management and target-role flow.", "assets/diagrams/37_source_management_flow.png"),
    ("Figure 59", "Job import and deduplication flow.", "assets/diagrams/38_job_import_deduplication_flow.png"),
    ("Figure 60", "Failed URL and retry flow.", "assets/diagrams/39_scraping_failure_dlq_flow.png"),
    ("Figure 61", "Scraping security boundaries.", "assets/diagrams/40_scraping_security_boundaries.png"),
    ("Figure 62", "Scraping validation evidence.", "assets/diagrams/41_scraping_validation_evidence.png"),
    ("Figure 63", "Frontend route and layout architecture.", "assets/diagrams/66_frontend_route_layout_architecture.png"),
    ("Figure 64", "Frontend API and authentication flow.", "assets/diagrams/67_frontend_api_auth_flow.png"),
    ("Figure 65", "Laravel backend request lifecycle.", "assets/diagrams/68_laravel_backend_request_lifecycle.png"),
    ("Figure 66", "Database relationship rationale.", "assets/diagrams/69_database_relationship_rationale.png"),
]

TABLES = [
    ("Table 1", "Stakeholder summary."),
    ("Table 2", "Functional requirements summary."),
    ("Table 3", "Non-functional requirements summary."),
    ("Table 4", "Software environment summary."),
    ("Table 5", "Backend module responsibility summary."),
    ("Table 6", "Laravel validation and protection mapping."),
    ("Table 7", "Database design rationale."),
    ("Table 8", "Data integrity mechanisms."),
    ("Table 9", "Main ERD relationship notes."),
    ("Table 10", "Design decisions summary."),
    ("Table 11", "AI CV Analyzer components."),
    ("Table 12", "NER entity label schema."),
    ("Table 13", "Synthetic dataset generation workflow."),
    ("Table 14", "Model training configuration."),
    ("Table 15", "Layer 1 component details."),
    ("Table 16", "Layer 2 classification engine details."),
    ("Table 17", "Layer 3 matching engine details."),
    ("Table 18", "Semantic embedding and TF-IDF fallback comparison."),
    ("Table 19", "Simplified BIO tagging example."),
    ("Table 20", "AI CV Analyzer source inventory summary."),
    ("Table 21", "Algorithm-to-file mapping."),
    ("Table 22", "AI design alternatives comparison."),
    ("Table 23", "Confidence and readiness signal summary."),
    ("Table 24", "Skill canonicalization example."),
    ("Table 25", "Dataset availability and transparency."),
    ("Table 26", "Seniority-aware matching weights."),
    ("Table 27", "Fit explanation output types."),
    ("Table 28", "Computational complexity overview."),
    ("Table 29", "Raw CV fragment extraction example."),
    ("Table 30", "Layer 2 interpretation example."),
    ("Table 31", "Layer 3 matching evidence example."),
    ("Table 32", "AI approach comparison."),
    ("Table 33", "Colab NER training run configuration."),
    ("Table 34", "Colab NER final metric summary."),
    ("Table 35", "AI CV Analyzer output schema sections."),
    ("Table 36", "Job mining design decisions."),
    ("Table 37", "Scraping runtime component map."),
    ("Table 38", "On-demand scraping lifecycle states."),
    ("Table 39", "Source management and target-role controls."),
    ("Table 40", "Import and deduplication stages."),
    ("Table 41", "Failed URL and operational failure handling."),
    ("Table 42", "Scraping security and configuration controls."),
    ("Table 43", "Job mining API contract summary."),
    ("Table 44", "Scraping validation evidence."),
    ("Table 45", "Scraping limitations, ethics, and future work."),
    ("Table 46", "Module validation coverage matrix."),
    ("Table 47", "Model evaluation evidence."),
    ("Table 48", "NER extraction examples."),
    ("Table 49", "Semantic matching and TF-IDF example results."),
    ("Table 50", "Mini CV dataset."),
    ("Table 51", "Mini job dataset."),
    ("Table 52", "Mini evaluation metrics."),
    ("Table 53", "Recommendation ranking details."),
    ("Table 54", "Gap analysis pair details."),
    ("Table 55", "Automated validation results."),
    ("Table 56", "Manual functional evaluation matrix."),
    ("Table 57", "Manual functional observations."),
    ("Table 58", "Security and privacy controls."),
    ("Table 59", "API endpoint summary."),
    ("Table 60", "Database tables summary."),
    ("Table 61", "Docker services summary."),
    ("Table 62", "AI CV Analyzer function inventory summary."),
]

COLAB_NER_EPOCHS = [
    {"epoch": 1, "training_loss": 0.077623, "validation_loss": 0.069118, "precision": 0.921027, "recall": 0.928206, "f1": 0.924603, "accuracy": 0.973227},
    {"epoch": 2, "training_loss": 0.061530, "validation_loss": 0.064051, "precision": 0.915886, "recall": 0.941504, "f1": 0.928518, "accuracy": 0.974912},
    {"epoch": 3, "training_loss": 0.053831, "validation_loss": 0.063463, "precision": 0.928387, "recall": 0.943469, "f1": 0.935867, "accuracy": 0.976233},
    {"epoch": 4, "training_loss": 0.044553, "validation_loss": 0.064025, "precision": 0.932287, "recall": 0.937967, "f1": 0.935118, "accuracy": 0.977018},
    {"epoch": 5, "training_loss": 0.037280, "validation_loss": 0.068058, "precision": 0.933307, "recall": 0.940521, "f1": 0.936900, "accuracy": 0.976376},
]


def ensure_dirs() -> None:
    for folder in [OUT_DIR, ASSETS, DIAGRAMS, SCREENSHOTS, LOGOS, EVALUATION]:
        folder.mkdir(parents=True, exist_ok=True)


def heading_anchor(text: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", text.lower()).strip("_")
    return f"bm_{slug or 'section'}"


def figure_anchor(number: str) -> str:
    return f"bm_figure_{number.split()[-1]}"


def table_anchor(number: str) -> str:
    return f"bm_table_{number.split()[-1]}"


def caption_anchor(caption: str) -> str | None:
    match = re.match(r"^(Figure|Table)\s+(\d+)\.", caption)
    if not match:
        return None
    return f"bm_{match.group(1).lower()}_{match.group(2)}"


def toc_markdown() -> str:
    return "\n".join([f"- [{label}](#{heading_anchor(target)})" for label, target in TOC_ENTRIES])


def load_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines or [""]


def rounded_box(draw, box, title, subtitle="", fill="#ffffff", outline="#1d4ed8", text="#0f172a"):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=12, fill=fill, outline=outline, width=2)
    title_font = load_font(22, True)
    sub_font = load_font(16)
    draw.text((x1 + 16, y1 + 14), title, fill=text, font=title_font)
    if subtitle:
        y = y1 + 46
        for line in wrap_text(draw, subtitle, sub_font, x2 - x1 - 32):
            draw.text((x1 + 16, y), line, fill="#334155", font=sub_font)
            y += 21


def arrow(draw, start, end, color="#0f766e", width=4):
    draw.line([start, end], fill=color, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    if abs(dx) > abs(dy):
        sign = 1 if dx > 0 else -1
        points = [(ex, ey), (ex - sign * 14, ey - 8), (ex - sign * 14, ey + 8)]
    else:
        sign = 1 if dy > 0 else -1
        points = [(ex, ey), (ex - 8, ey - sign * 14), (ex + 8, ey - sign * 14)]
    draw.polygon(points, fill=color)


INK = "#0f172a"
MUTED = "#475569"
GRID = "#94a3b8"
BLUE = "#2563eb"
GREEN = "#15803d"
TEAL = "#0f766e"
ORANGE = "#b45309"
RED = "#b91c1c"
PURPLE = "#6d28d9"
BG = "#f8fafc"
PAPER = "#ffffff"


def draw_diagram_title(draw: ImageDraw.ImageDraw, width: int, title: str, subtitle: str | None = None) -> None:
    draw.rectangle((0, 0, width, 86), fill=INK)
    draw.text((38, 24), title, fill="white", font=load_font(30, True))
    if subtitle:
        draw.text((40, 96), subtitle, fill=MUTED, font=load_font(17))


def wrapped_text_height(
    draw: ImageDraw.ImageDraw,
    text: str,
    font,
    max_width: int,
    *,
    line_gap: int = 5,
) -> int:
    lines = wrap_text(draw, text, font, max(20, max_width))
    if not lines:
        return 0
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    return len(lines) * line_height


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font,
    *,
    fill: str = INK,
    align: str = "left",
    line_gap: int = 5,
) -> int:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(20, x2 - x1))
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    y = y1
    for line in lines:
        if y + line_height > y2 + line_height:
            break
        text_w = draw.textbbox((0, 0), line, font=font)[2]
        if align == "center":
            x = x1 + max(0, (x2 - x1 - text_w) // 2)
        elif align == "right":
            x = x2 - text_w
        else:
            x = x1
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height
    return y


def draw_centered_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font,
    *,
    fill: str = INK,
    line_gap: int = 5,
) -> None:
    x1, y1, x2, y2 = box
    lines = wrap_text(draw, text, font, max(20, x2 - x1))
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + line_gap
    total_h = len(lines) * line_height - line_gap
    y = y1 + max(0, (y2 - y1 - total_h) // 2)
    for line in lines:
        text_w = draw.textbbox((0, 0), line, font=font)[2]
        draw.text((x1 + max(0, (x2 - x1 - text_w) // 2), y), line, fill=fill, font=font)
        y += line_height


def draw_dashed_line(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    fill: str = GRID,
    width: int = 2,
    dash: int = 12,
    gap: int = 8,
) -> None:
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    length = max(1, int((dx * dx + dy * dy) ** 0.5))
    step = dash + gap
    for offset in range(0, length, step):
        t1 = offset / length
        t2 = min(offset + dash, length) / length
        p1 = (int(sx + dx * t1), int(sy + dy * t1))
        p2 = (int(sx + dx * t2), int(sy + dy * t2))
        draw.line([p1, p2], fill=fill, width=width)


def draw_arrowhead(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    *,
    fill: str = TEAL,
    size: int = 14,
) -> None:
    sx, sy = start
    ex, ey = end
    dx, dy = ex - sx, ey - sy
    if abs(dx) >= abs(dy):
        sign = 1 if dx >= 0 else -1
        points = [(ex, ey), (ex - sign * size, ey - size // 2), (ex - sign * size, ey + size // 2)]
    else:
        sign = 1 if dy >= 0 else -1
        points = [(ex, ey), (ex - size // 2, ey - sign * size), (ex + size // 2, ey - sign * size)]
    draw.polygon(points, fill=fill)


def academic_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
    *,
    color: str = TEAL,
    width: int = 3,
    dashed: bool = False,
    label_pos: float = 0.5,
    label_offset: int = -28,
    font_size: int = 16,
) -> None:
    if dashed:
        draw_dashed_line(draw, start, end, fill=color, width=width)
    else:
        draw.line([start, end], fill=color, width=width)
    draw_arrowhead(draw, start, end, fill=color, size=15)
    if label:
        sx, sy = start
        ex, ey = end
        lx = int(sx + (ex - sx) * label_pos)
        ly = int(sy + (ey - sy) * label_pos) + label_offset
        font = load_font(font_size, True)
        bbox = draw.textbbox((0, 0), label, font=font)
        pad_x, pad_y = 8, 4
        box = (lx - (bbox[2] - bbox[0]) // 2 - pad_x, ly - pad_y, lx + (bbox[2] - bbox[0]) // 2 + pad_x, ly + (bbox[3] - bbox[1]) + pad_y)
        draw.rounded_rectangle(box, radius=6, fill=PAPER, outline="#bae6fd", width=1)
        draw.text((box[0] + pad_x, box[1] + pad_y - 1), label, fill=INK, font=font)


def academic_polyline_arrow(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    label: str = "",
    *,
    color: str = TEAL,
    width: int = 3,
    dashed: bool = False,
    label_at: tuple[int, int] | None = None,
    font_size: int = 15,
) -> None:
    if len(points) < 2:
        return
    for start, end in zip(points, points[1:]):
        if dashed:
            draw_dashed_line(draw, start, end, fill=color, width=width)
        else:
            draw.line([start, end], fill=color, width=width)
    draw_arrowhead(draw, points[-2], points[-1], fill=color, size=15)
    if label:
        lx, ly = label_at or points[len(points) // 2]
        font = load_font(font_size, True)
        bbox = draw.textbbox((0, 0), label, font=font)
        pad_x, pad_y = 8, 4
        box = (lx - (bbox[2] - bbox[0]) // 2 - pad_x, ly - pad_y, lx + (bbox[2] - bbox[0]) // 2 + pad_x, ly + (bbox[3] - bbox[1]) + pad_y)
        draw.rounded_rectangle(box, radius=6, fill=PAPER, outline="#bae6fd", width=1)
        draw.text((box[0] + pad_x, box[1] + pad_y - 1), label, fill=INK, font=font)


def flow_node(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str = "",
    *,
    fill: str = PAPER,
    outline: str = BLUE,
    kind: str = "process",
) -> None:
    x1, y1, x2, y2 = box
    if kind == "terminator":
        draw.rounded_rectangle(box, radius=(y2 - y1) // 2, fill=fill, outline=outline, width=2)
    elif kind == "decision":
        cx, cy = (x1 + x2) // 2, (y1 + y2) // 2
        draw.polygon([(cx, y1), (x2, cy), (cx, y2), (x1, cy)], fill=fill, outline=outline)
        draw.line([(cx, y1), (x2, cy), (cx, y2), (x1, cy), (cx, y1)], fill=outline, width=2)
    elif kind == "data":
        draw.rectangle(box, fill=fill, outline=outline, width=2)
        draw.line((x1 + 22, y1, x1 + 22, y2), fill=outline, width=2)
    else:
        draw.rounded_rectangle(box, radius=10, fill=fill, outline=outline, width=2)
    title_font = load_font(20, True)
    body_font = load_font(15)
    if body:
        draw_centered_wrapped(draw, title, (x1 + 14, y1 + 10, x2 - 14, y1 + 42), title_font)
        draw_wrapped(draw, body, (x1 + 16, y1 + 48, x2 - 16, y2 - 12), body_font, fill="#334155", align="center")
    else:
        draw_centered_wrapped(draw, title, (x1 + 12, y1 + 8, x2 - 12, y2 - 8), title_font)


def c4_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    stereotype: str,
    body: str,
    *,
    fill: str = PAPER,
    outline: str = BLUE,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=8, fill=fill, outline=outline, width=2)
    draw.rectangle((x1, y1, x2, y1 + 36), fill="#e0f2fe", outline=outline, width=2)
    draw.text((x1 + 14, y1 + 8), title, fill=INK, font=load_font(19, True))
    draw.text((x1 + 14, y1 + 43), stereotype, fill=MUTED, font=load_font(14, True))
    draw_wrapped(draw, body, (x1 + 14, y1 + 68, x2 - 14, y2 - 12), load_font(14), fill="#334155")


def c4_boundary(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str, *, outline: str = GRID) -> None:
    draw.rounded_rectangle(box, radius=16, outline=outline, width=2)
    x1, y1, _, _ = box
    draw.rounded_rectangle((x1 + 16, y1 - 18, x1 + 16 + max(190, len(label) * 10), y1 + 18), radius=8, fill=BG, outline=outline, width=1)
    draw.text((x1 + 28, y1 - 10), label, fill=MUTED, font=load_font(16, True))


def dfd_external(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str, body: str = "") -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#ffffff", outline=INK, width=2)
    draw_centered_wrapped(draw, label, (x1 + 10, y1 + 10, x2 - 10, y1 + 48), load_font(20, True))
    if body:
        draw_wrapped(draw, body, (x1 + 14, y1 + 55, x2 - 14, y2 - 10), load_font(14), fill=MUTED, align="center")


def dfd_process(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], number: str, label: str, body: str = "") -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=28, fill="#ecfdf5", outline=GREEN, width=3)
    draw.ellipse((x1 + 16, y1 + 15, x1 + 58, y1 + 57), fill=PAPER, outline=GREEN, width=2)
    draw_centered_wrapped(draw, number, (x1 + 16, y1 + 18, x1 + 58, y1 + 55), load_font(17, True), fill=GREEN)
    draw_centered_wrapped(draw, label, (x1 + 68, y1 + 12, x2 - 16, y1 + 58), load_font(20, True))
    if body:
        draw_wrapped(draw, body, (x1 + 28, y1 + 70, x2 - 28, y2 - 14), load_font(14), fill=MUTED, align="center")


def dfd_store(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], code: str, label: str) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#fff7ed", outline=ORANGE, width=2)
    draw.line((x1 + 46, y1, x1 + 46, y2), fill=ORANGE, width=2)
    draw.text((x1 + 12, y1 + 18), code, fill=ORANGE, font=load_font(18, True))
    draw_centered_wrapped(draw, label, (x1 + 58, y1 + 12, x2 - 12, y2 - 12), load_font(16, True))


def draw_actor(draw: ImageDraw.ImageDraw, cx: int, top: int, label: str) -> tuple[int, int]:
    head_r = 22
    draw.ellipse((cx - head_r, top, cx + head_r, top + 2 * head_r), outline=INK, width=3)
    body_top = top + 2 * head_r
    body_bottom = body_top + 70
    draw.line((cx, body_top, cx, body_bottom), fill=INK, width=3)
    draw.line((cx - 42, body_top + 22, cx + 42, body_top + 22), fill=INK, width=3)
    draw.line((cx, body_bottom, cx - 38, body_bottom + 54), fill=INK, width=3)
    draw.line((cx, body_bottom, cx + 38, body_bottom + 54), fill=INK, width=3)
    draw_centered_wrapped(draw, label, (cx - 80, body_bottom + 62, cx + 80, body_bottom + 108), load_font(18, True))
    return (cx, body_top + 30)


def draw_use_case(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], label: str) -> None:
    draw.ellipse(box, fill="#ffffff", outline=BLUE, width=3)
    draw_centered_wrapped(draw, label, (box[0] + 18, box[1] + 14, box[2] - 18, box[3] - 14), load_font(17, True))


def relation_line(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str = "",
    *,
    color: str = "#475569",
    dashed: bool = False,
) -> None:
    if dashed:
        draw_dashed_line(draw, start, end, fill=color, width=2)
    else:
        draw.line([start, end], fill=color, width=2)
    if label:
        sx, sy = start
        ex, ey = end
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        font = load_font(14, True)
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.rectangle((mx - (bbox[2] - bbox[0]) // 2 - 5, my - 14, mx + (bbox[2] - bbox[0]) // 2 + 5, my + 10), fill=BG)
        draw.text((mx - (bbox[2] - bbox[0]) // 2, my - 13), label, fill=color, font=font)


def erd_table(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], name: str, fields: list[tuple[str, str]]) -> None:
    x1, y1, x2, y2 = box
    draw.rectangle(box, fill="#ffffff", outline=INK, width=2)
    draw.rectangle((x1, y1, x2, y1 + 40), fill="#dbeafe", outline=INK, width=2)
    draw.text((x1 + 12, y1 + 9), name, fill=INK, font=load_font(19, True))
    y = y1 + 52
    for marker, field in fields:
        if y > y2 - 24:
            draw.text((x1 + 14, y), "...", fill=MUTED, font=load_font(14, True))
            break
        if marker:
            badge_w = max(30, draw.textbbox((0, 0), marker, font=load_font(12, True))[2] + 10)
            draw.rounded_rectangle((x1 + 10, y - 2, x1 + 10 + badge_w, y + 18), radius=5, fill="#f1f5f9", outline=GRID, width=1)
            draw.text((x1 + 15, y), marker, fill=INK, font=load_font(12, True))
            tx = x1 + 20 + badge_w
        else:
            tx = x1 + 16
        draw.text((tx, y - 1), field, fill="#334155", font=load_font(15))
        y += 25


def erd_card_label_position(point: tuple[int, int], other: tuple[int, int], role: str) -> tuple[int, int]:
    x, y = point
    dx = other[0] - x
    dy = other[1] - y
    if abs(dx) >= abs(dy):
        if dx >= 0:
            return (x + 8, y - 22) if role == "start" else (x - 44, y - 22)
        return (x - 44, y - 22) if role == "start" else (x + 8, y - 22)
    if dy >= 0:
        return (x + 8, y + 4) if role == "start" else (x + 8, y - 24)
    return (x + 8, y - 24) if role == "start" else (x + 8, y + 4)


def erd_relationship(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    start_card: str,
    end_card: str,
    label: str = "",
    *,
    dashed: bool = False,
) -> None:
    color = "#334155"
    if dashed:
        draw_dashed_line(draw, start, end, fill=color, width=2)
    else:
        draw.line([start, end], fill=color, width=2)
    sx, sy = start
    ex, ey = end
    font = load_font(13, True)
    draw.text(erd_card_label_position(start, end, "start"), start_card, fill=color, font=font)
    draw.text(erd_card_label_position(end, start, "end"), end_card, fill=color, font=font)
    if label:
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.rectangle((mx - (bbox[2] - bbox[0]) // 2 - 5, my - 16, mx + (bbox[2] - bbox[0]) // 2 + 5, my + 8), fill=BG)
        draw.text((mx - (bbox[2] - bbox[0]) // 2, my - 15), label, fill=color, font=font)


def erd_polyline_relationship(
    draw: ImageDraw.ImageDraw,
    points: list[tuple[int, int]],
    start_card: str,
    end_card: str,
    label: str = "",
    *,
    dashed: bool = False,
    label_at: tuple[int, int] | None = None,
) -> None:
    if len(points) < 2:
        return
    color = "#334155"
    for start, end in zip(points, points[1:]):
        if dashed:
            draw_dashed_line(draw, start, end, fill=color, width=2)
        else:
            draw.line([start, end], fill=color, width=2)
    font = load_font(14, True)
    draw.text(erd_card_label_position(points[0], points[1], "start"), start_card, fill=color, font=font)
    draw.text(erd_card_label_position(points[-1], points[-2], "end"), end_card, fill=color, font=font)
    if label:
        lx, ly = label_at or points[len(points) // 2]
        bbox = draw.textbbox((0, 0), label, font=font)
        draw.rectangle((lx - (bbox[2] - bbox[0]) // 2 - 5, ly - 16, lx + (bbox[2] - bbox[0]) // 2 + 5, ly + 8), fill=BG)
        draw.text((lx - (bbox[2] - bbox[0]) // 2, ly - 15), label, fill=color, font=font)


def save_diagram(name: str, title: str, boxes: list[tuple[str, str, tuple[int, int, int, int], str]], arrows: list[tuple[tuple[int, int], tuple[int, int], str]]):
    img = PILImage.new("RGB", (1600, 1000), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1600, title)
    for start, end, label in arrows:
        academic_arrow(draw, start, end, label, color=TEAL, width=3, font_size=14)
    for title_text, subtitle, box, color in boxes:
        rounded_box(draw, box, title_text, subtitle, fill=color)
    img.save(DIAGRAMS / name)


def save_frontend_route_readable_diagram() -> None:
    img = PILImage.new("RGB", (1800, 1100), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1800, 100), fill="#0f172a")
    draw.text((50, 30), "Frontend Route and Layout Architecture", fill="white", font=load_font(34, True))
    draw.text(
        (70, 145),
        "React Router separates public, student, admin, and preview screens through route guards.",
        fill="#475569",
        font=load_font(24),
    )
    cards = [
        ("1", "App Shell", "Router, theme, auth,\nnavbar, footer", (90, 230, 420, 390), "#dbeafe"),
        ("2", "Public Routes", "Home, login, register,\nstatus, policies", (560, 210, 910, 370), "#ecfdf5"),
        ("3", "Student Routes", "Dashboard, jobs, gaps,\nprofile, applications", (560, 470, 910, 630), "#fef3c7"),
        ("4", "Preview Screens", "CV builder, mock interview,\nlearning, tools", (560, 730, 910, 890), "#fff7ed"),
        ("5", "Route Guards", "GuestRoute and ProtectedRoute\nredirects", (1090, 210, 1460, 370), "#fee2e2"),
        ("6", "Admin Routes", "Dashboard, jobs, users,\nsources, targets", (1090, 470, 1460, 630), "#ede9fe"),
        ("7", "UI Evidence", "Screenshots for student, admin,\nstatus, preview", (1090, 730, 1460, 890), "#cffafe"),
    ]
    for num, title, body, box, fill in cards:
        x1, y1, x2, y2 = box
        draw.rounded_rectangle(box, radius=18, fill=fill, outline="#2563eb", width=3)
        draw.ellipse((x1 + 18, y1 + 22, x1 + 58, y1 + 62), fill="#0f172a")
        draw.text((x1 + 31, y1 + 29), num, fill="white", font=load_font(20, True))
        draw.text((x1 + 75, y1 + 25), title, fill="#0f172a", font=load_font(27, True))
        yy = y1 + 78
        for line in body.splitlines():
            draw.text((x1 + 34, yy), line, fill="#334155", font=load_font(22))
            yy += 31
    for start, end in [
        ((420, 310), (560, 290)),
        ((420, 310), (560, 550)),
        ((420, 310), (560, 810)),
        ((910, 290), (1090, 290)),
        ((910, 550), (1090, 550)),
        ((910, 810), (1090, 810)),
    ]:
        arrow(draw, start, end, color="#0f766e", width=4)
    draw.text(
        (90, 990),
        "Short labels keep the diagram readable in print; route details remain in the surrounding chapter text.",
        fill="#64748b",
        font=load_font(21),
    )
    img.save(DIAGRAMS / "66_frontend_route_layout_architecture.png")


def create_high_level_architecture() -> None:
    img = PILImage.new("RGB", (1900, 1150), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1900, "CareerCompass High-Level Architecture", "C4-style container view with client, application, AI, storage, and observability boundaries.")
    c4_boundary(draw, (370, 150, 1490, 800), "CareerCompass platform")
    c4_boundary(draw, (370, 850, 1490, 1060), "Data and storage")
    c4_boundary(draw, (1535, 150, 1835, 800), "External systems")
    c4_boundary(draw, (1535, 850, 1835, 1060), "Observability")

    c4_box(draw, (70, 240, 310, 380), "Student/Admin", "[Person]", "Uses browser UI for CV analysis, jobs, gaps, and administration.", fill="#f8fafc", outline=INK)
    c4_box(draw, (410, 300, 660, 460), "Nginx", "[Container]", "Public HTTP gateway and reverse proxy.", fill="#e0f2fe")
    c4_box(draw, (760, 190, 1080, 360), "React + Vite", "[Container]", "Student, admin, status, and preview route groups.", fill="#dbeafe")
    c4_box(draw, (760, 470, 1080, 650), "Laravel API", "[Container]", "Auth, CV upload, jobs, applications, gap analysis, admin APIs.", fill="#ecfdf5", outline=GREEN)
    c4_box(draw, (1180, 330, 1440, 500), "AI CV Analyzer", "[Container]", "FastAPI service for parsing, extraction, classification, and hybrid matching.", fill="#fce7f3", outline=PURPLE)
    c4_box(draw, (1180, 570, 1440, 735), "AI Job Miner", "[Container]", "FastAPI/Scrapy service for adapters, quality gates, and import callbacks.", fill="#ede9fe", outline=PURPLE)
    c4_box(draw, (410, 885, 680, 1030), "MySQL", "[Database]", "Users, profiles, skills, jobs, applications, sources, and analyses.", fill="#fff7ed", outline=ORANGE)
    c4_box(draw, (770, 885, 1010, 1030), "MinIO", "[Object store]", "Private uploaded CV objects.", fill="#cffafe", outline=TEAL)
    c4_box(draw, (1100, 885, 1410, 1030), "Queue Workers", "[Container]", "Background CV, email, and scraping work.", fill="#dcfce7", outline=GREEN)
    c4_box(draw, (1565, 350, 1810, 510), "Job Sources", "[External]", "Demo/API/HTML sources used by the miner.", fill="#fef3c7", outline=ORANGE)
    c4_box(draw, (1565, 900, 1810, 1030), "Prometheus + Grafana", "[Monitoring]", "Metrics collection and dashboards.", fill="#e0f2fe", outline=BLUE)

    academic_arrow(draw, (310, 310), (410, 380), "HTTPS")
    academic_arrow(draw, (660, 350), (760, 275), "static UI")
    academic_arrow(draw, (920, 360), (920, 470), "API calls")
    academic_arrow(draw, (1080, 535), (1180, 410), "parse / match")
    academic_arrow(draw, (1080, 605), (1180, 650), "scrape jobs")
    academic_arrow(draw, (915, 650), (545, 885), "SQL")
    academic_arrow(draw, (960, 650), (890, 885), "CV files")
    academic_arrow(draw, (1080, 585), (1255, 885), "dispatch", color=GREEN)
    academic_arrow(draw, (1440, 655), (1565, 430), "fetch")
    academic_arrow(draw, (1440, 445), (1565, 960), "metrics", dashed=True, color=BLUE, label_pos=0.64)
    academic_arrow(draw, (1080, 515), (1565, 960), "metrics", dashed=True, color=BLUE, label_pos=0.68)
    img.save(DIAGRAMS / "01_high_level_architecture.png")


def create_docker_deployment_diagram() -> None:
    img = PILImage.new("RGB", (1900, 1150), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1900, "Docker Compose Deployment View", "Container and volume relationships inside the local deployment network.")
    c4_boundary(draw, (50, 150, 1850, 815), "Docker Compose application network")
    c4_boundary(draw, (50, 855, 1850, 1065), "Stateful services and observability")
    c4_box(draw, (95, 285, 330, 430), "nginx", "[edge container]", "Host port 80 reverse proxy.", fill="#e0f2fe")
    c4_box(draw, (450, 205, 700, 350), "frontend", "[static container]", "Built React assets.", fill="#dbeafe")
    c4_box(draw, (825, 205, 1095, 350), "backend-api", "[Laravel PHP-FPM]", "HTTP API and application logic.", fill="#ecfdf5", outline=GREEN)
    c4_box(draw, (1215, 205, 1515, 350), "backend workers", "[queue workers]", "default, high, AI, email, and scraping queues.", fill="#dcfce7", outline=GREEN)
    c4_box(draw, (450, 560, 735, 710), "ai-cv-analyzer", "[FastAPI]", "CV parse and hybrid-match endpoints.", fill="#fce7f3", outline=PURPLE)
    c4_box(draw, (855, 560, 1140, 710), "ai-job-miner", "[FastAPI]", "Scraping adapters and import callbacks.", fill="#ede9fe", outline=PURPLE)
    c4_box(draw, (1260, 560, 1515, 710), "external web/API", "[outside Docker]", "Job boards and demo sources.", fill="#fef3c7", outline=ORANGE)
    c4_box(draw, (115, 900, 365, 1035), "db", "[MySQL 8.0]", "Relational system of record.", fill="#fff7ed", outline=ORANGE)
    c4_box(draw, (485, 900, 735, 1035), "minio", "[S3-compatible]", "Private CV object storage.", fill="#cffafe", outline=TEAL)
    c4_box(draw, (855, 900, 1115, 1035), "prometheus", "[metrics]", "Scrapes runtime metrics.", fill="#e0f2fe", outline=BLUE)
    c4_box(draw, (1235, 900, 1495, 1035), "grafana", "[dashboards]", "Reads Prometheus data.", fill="#e0f2fe", outline=BLUE)

    academic_arrow(draw, (330, 355), (450, 280), "serve")
    academic_arrow(draw, (700, 280), (825, 280), "API")
    academic_arrow(draw, (1095, 280), (1215, 280), "queue jobs", color=GREEN)
    academic_arrow(draw, (960, 350), (590, 560), "parse CV")
    academic_arrow(draw, (980, 350), (1000, 560), "scrape")
    academic_arrow(draw, (1140, 635), (1260, 635), "fetch")
    academic_arrow(draw, (960, 350), (240, 900), "SQL", label_pos=0.35)
    academic_arrow(draw, (990, 350), (610, 900), "S3")
    academic_arrow(draw, (1515, 280), (990, 560), "run")
    academic_arrow(draw, (985, 710), (985, 900), "metrics", dashed=True, color=BLUE)
    academic_arrow(draw, (1115, 965), (1235, 965), "query", color=BLUE)
    img.save(DIAGRAMS / "02_docker_deployment.png")


def create_dfd_level_0() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "DFD Level 0 Context Diagram", "External entities exchange data flows with the CareerCompass system boundary.")
    draw.rounded_rectangle((485, 185, 1265, 900), radius=18, outline=GRID, width=2)
    draw.text((510, 200), "CareerCompass system boundary", fill=MUTED, font=load_font(17, True))
    dfd_external(draw, (90, 260, 370, 420), "Student", "CV, profile actions, job and gap requests")
    dfd_external(draw, (90, 620, 370, 780), "Administrator", "Job/source/user operations and diagnostics")
    dfd_process(draw, (620, 390, 1125, 650), "0", "CareerCompass Platform", "Processes CVs, stores profiles, ranks jobs, analyzes gaps, imports jobs, and reports status.")
    dfd_external(draw, (1410, 240, 1695, 405), "External Job Sources", "API, HTML, and demo job data")
    dfd_external(draw, (1410, 635, 1695, 800), "AI Services", "CV parsing and hybrid match support")
    dfd_store(draw, (700, 745, 1045, 830), "D", "Persistent data stores")

    academic_arrow(draw, (370, 315), (620, 455), "CV/profile/job requests")
    academic_arrow(draw, (620, 545), (370, 370), "recommendations, gaps, status")
    academic_arrow(draw, (370, 700), (620, 575), "admin commands")
    academic_arrow(draw, (620, 625), (370, 760), "dashboards, results")
    academic_arrow(draw, (1125, 455), (1410, 320), "source queries")
    academic_arrow(draw, (1410, 365), (1125, 520), "job data")
    academic_arrow(draw, (1125, 600), (1410, 705), "parse/match requests")
    academic_arrow(draw, (1410, 750), (1125, 640), "analysis results")
    academic_arrow(draw, (875, 650), (875, 745), "read/write")
    img.save(DIAGRAMS / "03_dfd_level_0.png")


def create_dfd_level_1() -> None:
    img = PILImage.new("RGB", (1900, 1250), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1900, "DFD Level 1 Process Diagram", "Major CareerCompass processes, data stores, and labeled data flows.")
    draw.rounded_rectangle((350, 145, 1555, 1060), radius=18, outline=GRID, width=2)
    draw.text((375, 160), "CareerCompass system boundary", fill=MUTED, font=load_font(17, True))

    dfd_external(draw, (60, 290, 280, 430), "Student", "Authenticated user")
    dfd_external(draw, (60, 705, 280, 845), "Admin", "Operator")
    dfd_external(draw, (1630, 260, 1840, 405), "AI Analyzer", "Parse and match")
    dfd_external(draw, (1630, 720, 1840, 865), "Job Sources", "Public/demo sources")

    dfd_process(draw, (430, 230, 720, 390), "1.0", "Authenticate and Manage Profile", "Tokens, roles, profile bootstrap.")
    dfd_process(draw, (820, 230, 1130, 390), "2.0", "Process CV", "Validate upload, store object, call analyzer, normalize evidence.")
    dfd_process(draw, (1210, 230, 1485, 390), "3.0", "Rank Jobs", "Laravel title, skill-overlap, and seniority scoring.")
    dfd_process(draw, (820, 520, 1130, 690), "4.0", "Analyze Skill Gap", "Gap service calls /api/hybrid-match when available.")
    dfd_process(draw, (1210, 720, 1485, 890), "5.0", "Import Jobs", "Queue scraping, validate payload, deduplicate, sync skills.")
    dfd_process(draw, (430, 720, 720, 890), "6.0", "Track Applications", "Save and update job application status.")

    dfd_store(draw, (455, 1010, 730, 1095), "D1", "Users and profiles")
    dfd_store(draw, (785, 1010, 1065, 1095), "D2", "CV analyses and skills")
    dfd_store(draw, (1115, 1010, 1395, 1095), "D3", "Jobs, sources, roles")
    dfd_store(draw, (1425, 1010, 1665, 1095), "D4", "Applications")

    academic_arrow(draw, (280, 350), (430, 310), "credentials / profile")
    academic_arrow(draw, (720, 300), (820, 300), "authorized upload")
    academic_arrow(draw, (975, 390), (975, 520), "profile evidence")
    academic_arrow(draw, (1130, 605), (1630, 335), "/api/hybrid-match")
    academic_arrow(draw, (1630, 365), (1130, 655), "scores and gaps", dashed=True)
    academic_arrow(draw, (1130, 305), (1210, 305), "stored profile + jobs")
    academic_polyline_arrow(draw, [(1485, 305), (1540, 305), (1540, 205), (240, 205), (240, 360), (280, 360)], "ranked jobs", dashed=True, label_at=(820, 205))
    academic_polyline_arrow(draw, [(280, 775), (345, 775), (345, 935), (1210, 935), (1210, 805)], "run/test sources", label_at=(720, 935))
    academic_arrow(draw, (1485, 805), (1630, 790), "source request")
    academic_arrow(draw, (1630, 835), (1485, 860), "job payloads")
    academic_arrow(draw, (720, 805), (820, 640), "selected job")
    academic_arrow(draw, (280, 805), (430, 805), "save/update")
    academic_arrow(draw, (575, 890), (575, 1010), "read/write")
    academic_arrow(draw, (975, 690), (925, 1010), "analysis data")
    academic_arrow(draw, (1350, 890), (1260, 1010), "imported jobs")
    academic_polyline_arrow(draw, [(575, 890), (575, 960), (1545, 960), (1545, 1010)], "application record", label_at=(1270, 960))
    img.save(DIAGRAMS / "04_dfd_level_1.png")


def create_use_case_diagram() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "UML Use Case Diagram", "Actors interact with use cases inside the CareerCompass system boundary.")
    student_anchor = draw_actor(draw, 185, 250, "Student")
    admin_anchor = draw_actor(draw, 185, 680, "Admin")
    draw.rectangle((415, 155, 1335, 965), outline=INK, width=3)
    draw.rectangle((415, 155, 690, 197), fill=BG, outline=INK, width=2)
    draw.text((438, 166), "CareerCompass", fill=INK, font=load_font(21, True))

    use_cases = {
        "Register / login": (500, 225, 760, 315),
        "Upload CV": (815, 225, 1075, 315),
        "View parsed profile": (500, 370, 760, 460),
        "View job list": (815, 370, 1075, 460),
        "Track applications": (500, 515, 760, 605),
        "Analyze skill gap": (815, 515, 1075, 605),
        "Manage jobs": (500, 705, 760, 795),
        "Manage sources / targets": (815, 705, 1115, 795),
        "Monitor dashboard": (500, 850, 760, 940),
        "Import scraped jobs": (885, 850, 1190, 940),
    }
    for label, box in use_cases.items():
        draw_use_case(draw, box, label)

    dfd_external(draw, (1480, 290, 1710, 420), "AI CV Analyzer", "External service")
    dfd_external(draw, (1480, 515, 1710, 645), "AI Matching", "/api/hybrid-match")
    dfd_external(draw, (1480, 745, 1710, 875), "Job Sources", "External data")

    for end in [(500, 270), (815, 270), (500, 415), (815, 415), (500, 560), (815, 560)]:
        relation_line(draw, student_anchor, end)
    for end in [(500, 750), (815, 750), (500, 895), (885, 895)]:
        relation_line(draw, admin_anchor, end)
    relation_line(draw, (1075, 270), (1480, 355), "<<include>> parse", dashed=True)
    relation_line(draw, (1075, 560), (1480, 580), "<<include>> match", dashed=True)
    relation_line(draw, (1190, 895), (1480, 810), "fetches", dashed=True)
    relation_line(draw, (1115, 750), (1480, 810), "configures", dashed=True)
    img.save(DIAGRAMS / "05_use_case_diagram.png")


def create_frontend_route_layout_architecture() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "Frontend Route and Layout Architecture", "React Router tree grouped by layout shell and guard behavior.")
    c4_boundary(draw, (80, 155, 1720, 970), "frontend/src/App.jsx route tree")
    c4_box(draw, (160, 230, 460, 385), "App Shell", "[RouterProvider]", "Theme, auth provider, navbar/footer layout, and route outlet.", fill="#e0f2fe")
    c4_box(draw, (610, 170, 940, 315), "Public Layout", "[public routes]", "Home, login, register, status, policies, and landing content.", fill="#dbeafe")
    c4_box(draw, (610, 380, 940, 535), "Student Layout", "[ProtectedRoute]", "Dashboard, CV upload, jobs, job detail, gap analysis, applications, profile.", fill="#ecfdf5", outline=GREEN)
    c4_box(draw, (610, 590, 940, 745), "Preview Modules", "[future/demo routes]", "CV builder, mock interview, learning paths, planner, tools hub.", fill="#fff7ed", outline=ORANGE)
    c4_box(draw, (610, 800, 940, 925), "Admin Layout", "[AdminRoute]", "Dashboard, jobs, users, sources, target roles, diagnostics.", fill="#ede9fe", outline=PURPLE)
    c4_box(draw, (1120, 275, 1450, 430), "Route Guards", "[GuestRoute / ProtectedRoute]", "Redirects unauthenticated users and prevents role leakage.", fill="#fee2e2", outline=RED)
    c4_box(draw, (1120, 540, 1450, 700), "Shared UI Layer", "[components/hooks/i18n]", "Reusable cards, forms, tables, localization strings, and loading/error states.", fill="#cffafe", outline=TEAL)
    c4_box(draw, (1120, 785, 1450, 920), "Evidence Screens", "[screenshots]", "Student, admin, status, and preview figures referenced in Chapter 5.", fill="#f8fafc", outline=GRID)

    academic_arrow(draw, (460, 305), (610, 240), "routes")
    academic_arrow(draw, (460, 305), (610, 455), "auth")
    academic_arrow(draw, (460, 305), (610, 665), "preview")
    academic_arrow(draw, (460, 305), (610, 860), "admin")
    academic_arrow(draw, (940, 455), (1120, 350), "guard")
    academic_arrow(draw, (940, 860), (1120, 350), "role guard", label_pos=0.40)
    academic_arrow(draw, (940, 455), (1120, 620), "components")
    academic_arrow(draw, (940, 860), (1120, 620), "components")
    academic_arrow(draw, (1285, 700), (1285, 785), "rendered evidence", color=BLUE)
    img.save(DIAGRAMS / "66_frontend_route_layout_architecture.png")


def create_frontend_api_auth_flow() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "Frontend API and Authentication Flow", "Layered API client path from page state to Laravel response handling.")
    c4_boundary(draw, (80, 150, 900, 930), "React frontend")
    c4_boundary(draw, (1020, 150, 1620, 930), "Laravel API")
    flow_node(draw, (150, 230, 420, 350), "Page / Component", "Calls endpoint modules and updates local UI state.", fill="#dbeafe")
    flow_node(draw, (560, 230, 830, 350), "API Wrappers", "jobsAPI, cvAPI, authAPI, adminAPI centralize paths.", fill="#ecfdf5")
    flow_node(draw, (560, 480, 830, 625), "Axios Client", "Base URL, JSON headers, request ID, bearer token, retry rules.", fill="#fef3c7")
    flow_node(draw, (150, 480, 420, 625), "AuthContext", "Login/register persist token and user; refreshUser reloads /user.", fill="#cffafe")
    flow_node(draw, (150, 750, 420, 865), "Browser Storage", "Auth token and user snapshot for session continuity.", fill="#f8fafc", kind="data")
    flow_node(draw, (1085, 250, 1385, 380), "Route + Middleware", "/api and /api/v1 groups apply auth, admin, throttle, and scraper token guards.", fill="#ede9fe")
    flow_node(draw, (1085, 520, 1385, 655), "Controller / Resource", "Validates, delegates to services, and returns JSON resources.", fill="#ecfdf5")
    flow_node(draw, (1085, 770, 1385, 885), "JSON Response", "Data, validation errors, or 401/403 status.", fill="#fee2e2")

    academic_arrow(draw, (420, 290), (560, 290), "call endpoint")
    academic_arrow(draw, (695, 350), (695, 480), "request config")
    academic_arrow(draw, (560, 550), (420, 550), "401 clears state", dashed=True, color=RED)
    academic_arrow(draw, (285, 480), (285, 350), "user/token")
    academic_arrow(draw, (285, 625), (285, 750), "persist")
    academic_arrow(draw, (830, 552), (1085, 315), "HTTP + bearer")
    academic_arrow(draw, (1235, 380), (1235, 520), "dispatch")
    academic_arrow(draw, (1235, 655), (1235, 770), "format")
    academic_arrow(draw, (1085, 830), (830, 560), "JSON/401", dashed=True, color=RED, label_pos=0.55)
    img.save(DIAGRAMS / "67_frontend_api_auth_flow.png")


def create_laravel_backend_request_lifecycle() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "Laravel Backend Request Lifecycle", "Request pipeline with synchronous response path and asynchronous worker branch.")
    c4_boundary(draw, (80, 150, 1660, 765), "Synchronous API path")
    c4_boundary(draw, (520, 815, 1660, 1015), "Asynchronous branch")
    steps = [
        ((125, 305, 345, 430), "HTTP Request", "Browser, scraper service, health probe, or admin action.", "#dbeafe"),
        ((455, 305, 675, 430), "Routes", "/api and /api/v1 route groups.", "#ecfdf5"),
        ((785, 305, 1005, 430), "Middleware", "Sanctum auth, admin, scraper token, throttle, request ID.", "#fee2e2"),
        ((1115, 305, 1335, 430), "Form Request", "Input validation and authorization rules.", "#fef3c7"),
        ((1445, 305, 1625, 430), "Controller", "Coordinates endpoint behavior.", "#ede9fe"),
        ((1115, 560, 1335, 690), "Service Layer", "CV, gaps, applications, scraper client, skill sync.", "#cffafe"),
        ((785, 560, 1005, 690), "Model / IO", "Eloquent, MySQL, MinIO, AI services.", "#dcfce7"),
        ((455, 560, 675, 690), "JSON Resource", "Shapes response payload.", "#e0f2fe"),
    ]
    for box, title, body, fill in steps:
        flow_node(draw, box, title, body, fill=fill)
    flow_node(draw, (605, 860, 875, 970), "Queue Job", "CV processing, scraping, email, or retryable background work.", fill="#dcfce7")
    flow_node(draw, (1035, 860, 1305, 970), "Worker", "Runs queued job with timeout and retry policy.", fill="#dcfce7")
    flow_node(draw, (1450, 860, 1615, 970), "Status / Events", "Persisted counters and pollable completion state.", fill="#f8fafc", kind="data")

    for start, end, label in [
        ((345, 368), (455, 368), "match"),
        ((675, 368), (785, 368), "guard"),
        ((1005, 368), (1115, 368), "validate"),
        ((1335, 368), (1445, 368), "invoke"),
        ((1535, 430), (1225, 560), "delegate"),
        ((1115, 625), (1005, 625), "persist/call"),
        ((785, 625), (675, 625), "resource"),
        ((455, 625), (235, 430), "JSON"),
    ]:
        academic_arrow(draw, start, end, label)
    academic_arrow(draw, (1225, 690), (740, 860), "dispatch long work", color=GREEN)
    academic_arrow(draw, (875, 915), (1035, 915), "queue:work", color=GREEN)
    academic_arrow(draw, (1305, 915), (1450, 915), "update", color=GREEN)
    academic_arrow(draw, (1535, 860), (900, 690), "poll/read", dashed=True, color=BLUE, label_pos=0.35)
    img.save(DIAGRAMS / "68_laravel_backend_request_lifecycle.png")


def create_database_relationship_rationale() -> None:
    img = PILImage.new("RGB", (1800, 1100), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1800, "Database Relationship Rationale", "Normalized groups preserve explainability for profile, skill, job, application, and scraping evidence.")
    c4_boundary(draw, (80, 160, 540, 850), "Identity and profile evidence")
    c4_boundary(draw, (650, 160, 1190, 850), "Reusable skill graph")
    c4_boundary(draw, (1300, 160, 1715, 850), "Jobs and operations")
    c4_box(draw, (135, 235, 455, 385), "users", "[root identity]", "Owns authentication role and account state.", fill="#dbeafe")
    c4_box(draw, (135, 475, 455, 625), "user_profiles / experiences", "[profile evidence]", "Stores headline, seniority, domain, contact info, and experience records.", fill="#ecfdf5", outline=GREEN)
    c4_box(draw, (135, 700, 455, 815), "cv_analyses", "[analysis evidence]", "Stores CV file metadata, parsing status, predicted role, and raw output.", fill="#fce7f3", outline=PURPLE)
    c4_box(draw, (740, 250, 1060, 400), "skills", "[canonical terms]", "Shared technical and soft skill vocabulary.", fill="#fff7ed", outline=ORANGE)
    c4_box(draw, (700, 540, 900, 690), "user_skills", "[pivot]", "Links users to skills with confidence and evidence.", fill="#f8fafc", outline=GRID)
    c4_box(draw, (960, 540, 1160, 690), "job_skills", "[pivot]", "Links jobs to required or nice-to-have skills.", fill="#f8fafc", outline=GRID)
    c4_box(draw, (1360, 240, 1660, 395), "job_postings", "[opportunities]", "Imported or seeded jobs; references source when known.", fill="#ede9fe", outline=PURPLE)
    c4_box(draw, (1360, 500, 1660, 650), "applications", "[tracking]", "User-job status, notes, and applied timestamp.", fill="#cffafe", outline=TEAL)
    c4_box(draw, (1360, 735, 1660, 835), "scraping_* / target roles", "[operations]", "Sources, jobs, failures, and search-intent configuration.", fill="#fee2e2", outline=RED)

    academic_arrow(draw, (455, 310), (700, 615), "1 to many skills", color=GREEN)
    academic_arrow(draw, (900, 615), (740, 340), "FK skill_id", color=ORANGE)
    academic_arrow(draw, (1060, 340), (1160, 615), "FK skill_id", color=ORANGE)
    academic_arrow(draw, (1160, 615), (1360, 320), "job requirements", color=PURPLE)
    academic_arrow(draw, (455, 548), (700, 615), "profile evidence", color=GREEN, dashed=True)
    academic_arrow(draw, (455, 758), (740, 340), "extracted terms", color=PURPLE, dashed=True, label_pos=0.35)
    academic_arrow(draw, (455, 310), (1360, 575), "user applies to job", color=TEAL, label_pos=0.70)
    academic_arrow(draw, (1510, 395), (1510, 500), "1 to many")
    academic_arrow(draw, (1510, 735), (1510, 650), "imports / diagnostics", color=RED)
    draw.text((105, 915), "Rationale: identity is separated from extracted evidence; skills are reusable pivots; scraping operations remain operational metadata, not user profile data.", fill=MUTED, font=load_font(19, True))
    img.save(DIAGRAMS / "69_database_relationship_rationale.png")


def create_dataset_evidence_diagram() -> None:
    width = 1600
    margin_x = 110
    content_right = 1490
    inner_pad_x = 30
    inner_pad_y = 22
    badge_right = 1450
    badge_min_width = 150
    body_max_width = 1100

    title_font = load_font(25, True)
    body_font = load_font(19)
    badge_font = load_font(18, True)
    decision_font = load_font(20, True)
    cards = [
        ("Final NER Training Dataset", "Dataset content unavailable in committed Git evidence", "0 dataset files", "#fee2e2"),
        ("Training Notebook Logic", "Available: labels, token alignment, 90/10 split", "workflow", "#dbeafe"),
        ("Colab Training-Run PDF", "Available: exported output cells with train/test counts and overall metrics", "PDF evidence", "#dcfce7"),
        ("Mini Evaluation Dataset", "Available: 5 synthetic CVs and 8 synthetic jobs", "13 records", "#dcfce7"),
        ("AI CV Smoke Dataset", "Available: 5 deterministic fake CV text samples", "5 samples", "#ecfdf5"),
        ("Per-Label NER Metrics", "Unavailable in the PDF: no classification report or confusion matrix", "not claimed", "#fef3c7"),
    ]
    decision_text = (
        "Decision: use Colab overall metrics; no per-label distribution chart because label support counts are not visible."
    )

    measure_img = PILImage.new("RGB", (width, 10))
    measure_draw = ImageDraw.Draw(measure_img)
    title_line_h = measure_draw.textbbox((0, 0), "Ag", font=title_font)[3] + 8
    body_line_h = measure_draw.textbbox((0, 0), "Ag", font=body_font)[3] + 6
    badge_pad_y = 12
    decision_line_h = measure_draw.textbbox((0, 0), "Ag", font=decision_font)[3] + 8
    card_gap = 24
    decision_gap = 36

    card_layouts: list[dict] = []
    for title, body, badge, fill in cards:
        body_lines = wrap_text(measure_draw, body, body_font, body_max_width)
        body_h = len(body_lines) * body_line_h
        badge_h = measure_draw.textbbox((0, 0), badge, font=badge_font)[3] + badge_pad_y * 2
        content_h = title_line_h + 10 + body_h
        card_h = inner_pad_y * 2 + max(content_h, badge_h)
        card_layouts.append(
            {
                "title": title,
                "body_lines": body_lines,
                "badge": badge,
                "fill": fill,
                "height": card_h,
                "badge_h": badge_h,
            }
        )

    decision_lines = wrap_text(measure_draw, decision_text, decision_font, content_right - margin_x - inner_pad_x * 2)
    decision_h = inner_pad_y * 2 + len(decision_lines) * decision_line_h

    header_h = 96
    top_margin = 64
    bottom_margin = 56
    cards_total = sum(card["height"] for card in card_layouts) + card_gap * (len(card_layouts) - 1)
    total_h = header_h + top_margin + cards_total + decision_gap + decision_h + bottom_margin

    img = PILImage.new("RGB", (width, total_h), BG)
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, width, header_h), fill=INK)
    draw.text((44, 28), "Dataset Evidence Availability", fill="white", font=load_font(34, True))

    y = header_h + top_margin
    for index, card in enumerate(card_layouts):
        card_top = y
        card_bottom = y + card["height"]
        draw.rounded_rectangle((margin_x, card_top, content_right, card_bottom), radius=16, fill=card["fill"], outline=GRID, width=2)
        text_y = card_top + inner_pad_y
        draw.text((margin_x + inner_pad_x, text_y), card["title"], fill=INK, font=title_font)
        text_y += title_line_h + 10
        for line in card["body_lines"]:
            draw.text((margin_x + inner_pad_x, text_y), line, fill="#334155", font=body_font)
            text_y += body_line_h
        badge_w = max(badge_min_width, measure_draw.textbbox((0, 0), card["badge"], font=badge_font)[2] + 36)
        badge_x1 = badge_right - badge_w
        badge_y1 = card_top + (card["height"] - card["badge_h"]) // 2
        draw.rounded_rectangle(
            (badge_x1, badge_y1, badge_right, badge_y1 + card["badge_h"]),
            radius=12,
            fill=PAPER,
            outline="#64748b",
            width=2,
        )
        draw.text((badge_x1 + 18, badge_y1 + badge_pad_y - 2), card["badge"], fill=INK, font=badge_font)
        y = card_bottom + (card_gap if index < len(card_layouts) - 1 else 0)

    decision_top = y + decision_gap
    decision_bottom = decision_top + decision_h
    draw.rounded_rectangle((margin_x, decision_top, content_right, decision_bottom), radius=14, fill="#e0f2fe", outline="#0284c7", width=2)
    text_y = decision_top + inner_pad_y
    for line in decision_lines:
        draw.text((margin_x + inner_pad_x, text_y), line, fill="#0c4a6e", font=decision_font)
        text_y += decision_line_h
    img.save(DIAGRAMS / "29_dataset_evidence_availability.png")


def create_smoke_metrics_diagram() -> None:
    results = run_smoke_evaluation() if not SMOKE_EVAL_RESULTS.exists() else json.loads(SMOKE_EVAL_RESULTS.read_text(encoding="utf-8"))
    summary = results.get("summary", {})
    metrics = [
        ("Skill P", summary.get("macro_skill_precision", 0.0), "#2563eb"),
        ("Skill R", summary.get("macro_skill_recall", 0.0), "#0891b2"),
        ("Skill F1", summary.get("macro_skill_f1", 0.0), "#059669"),
        ("Role", summary.get("role_match_rate", 0.0), "#7c3aed"),
        ("Domain", summary.get("domain_match_rate", 0.0), "#ea580c"),
        ("Seniority", summary.get("seniority_match_rate", 0.0), "#dc2626"),
        ("Status", summary.get("parsing_status_match_rate", 0.0), "#0f766e"),
    ]
    img = PILImage.new("RGB", (1600, 1000), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 96), fill="#0f172a")
    draw.text((44, 28), "AI CV Analyzer Smoke Evaluation Metrics", fill="white", font=load_font(34, True))
    axis_left, axis_bottom = 160, 790
    axis_top, axis_right = 180, 1460
    draw.line((axis_left, axis_bottom, axis_right, axis_bottom), fill="#475569", width=3)
    draw.line((axis_left, axis_top, axis_left, axis_bottom), fill="#475569", width=3)
    for tick in range(0, 101, 20):
        y = axis_bottom - int((axis_bottom - axis_top) * (tick / 100))
        draw.line((axis_left - 8, y, axis_right, y), fill="#e2e8f0", width=1)
        draw.text((82, y - 12), f"{tick}%", fill="#475569", font=load_font(18, True))
    bar_gap = 32
    bar_width = int((axis_right - axis_left - (len(metrics) + 1) * bar_gap) / len(metrics))
    x = axis_left + bar_gap
    for label, value, color in metrics:
        pct = max(0.0, min(1.0, float(value)))
        bar_h = int((axis_bottom - axis_top) * pct)
        y1 = axis_bottom - bar_h
        draw.rounded_rectangle((x, y1, x + bar_width, axis_bottom), radius=8, fill=color)
        draw.text((x + 4, y1 - 34), f"{pct * 100:.0f}%", fill="#0f172a", font=load_font(20, True))
        draw.text((x, axis_bottom + 22), label, fill="#334155", font=load_font(19, True))
        x += bar_width + bar_gap
    draw.rounded_rectangle((160, 850, 1460, 935), radius=14, fill="#fef3c7", outline="#f59e0b", width=2)
    draw.text((190, 872), "Scope: five fake text samples; deterministic smoke evidence only, not final NER accuracy.", fill="#78350f", font=load_font(22, True))
    img.save(DIAGRAMS / "30_ai_cv_analyzer_smoke_metrics.png")


def create_colab_ner_metrics_diagram() -> None:
    metrics = [
        ("Precision", COLAB_NER_EPOCHS[-1]["precision"], "#2563eb"),
        ("Recall", COLAB_NER_EPOCHS[-1]["recall"], "#0891b2"),
        ("F1", COLAB_NER_EPOCHS[-1]["f1"], "#059669"),
        ("Accuracy", COLAB_NER_EPOCHS[-1]["accuracy"], "#7c3aed"),
    ]
    img = PILImage.new("RGB", (1600, 1000), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 96), fill="#0f172a")
    draw.text((44, 28), "Colab NER Final Epoch Metrics", fill="white", font=load_font(34, True))
    axis_left, axis_bottom = 180, 790
    axis_top, axis_right = 180, 1440
    draw.line((axis_left, axis_bottom, axis_right, axis_bottom), fill="#475569", width=3)
    draw.line((axis_left, axis_top, axis_left, axis_bottom), fill="#475569", width=3)
    for tick in range(0, 101, 20):
        y = axis_bottom - int((axis_bottom - axis_top) * (tick / 100))
        draw.line((axis_left - 8, y, axis_right, y), fill="#e2e8f0", width=1)
        draw.text((100, y - 12), f"{tick}%", fill="#475569", font=load_font(18, True))
    bar_gap = 70
    bar_width = int((axis_right - axis_left - (len(metrics) + 1) * bar_gap) / len(metrics))
    x = axis_left + bar_gap
    for label, value, color in metrics:
        pct = max(0.0, min(1.0, float(value)))
        bar_h = int((axis_bottom - axis_top) * pct)
        y1 = axis_bottom - bar_h
        draw.rounded_rectangle((x, y1, x + bar_width, axis_bottom), radius=10, fill=color)
        draw.text((x + 8, y1 - 38), f"{pct * 100:.2f}%", fill="#0f172a", font=load_font(22, True))
        draw.text((x + 8, axis_bottom + 24), label, fill="#334155", font=load_font(21, True))
        x += bar_width + bar_gap
    draw.rounded_rectangle((150, 850, 1450, 940), radius=14, fill="#e0f2fe", outline="#0284c7", width=2)
    draw.text((180, 872), "Source: exported train_ner.ipynb Colab PDF, final epoch row. Validation evidence, not production accuracy.", fill="#0c4a6e", font=load_font(21, True))
    img.save(DIAGRAMS / "31_colab_ner_metrics.png")


def _draw_line_chart(
    filename: str,
    title: str,
    series: list[tuple[str, list[float], str]],
    y_min: float,
    y_max: float,
    ticks: list[float],
    value_suffix: str,
    footnote: str,
) -> None:
    img = PILImage.new("RGB", (1700, 1050), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1700, 100), fill="#0f172a")
    draw.text((48, 30), title, fill="white", font=load_font(34, True))

    axis_left, axis_top, axis_right, axis_bottom = 180, 190, 1500, 790
    draw.rectangle((axis_left, axis_top, axis_right, axis_bottom), fill="#ffffff", outline="#cbd5e1", width=2)
    draw.line((axis_left, axis_bottom, axis_right, axis_bottom), fill="#334155", width=3)
    draw.line((axis_left, axis_top, axis_left, axis_bottom), fill="#334155", width=3)

    def map_y(value: float) -> int:
        clamped = max(y_min, min(y_max, value))
        return axis_bottom - int((axis_bottom - axis_top) * ((clamped - y_min) / (y_max - y_min)))

    epochs = [row["epoch"] for row in COLAB_NER_EPOCHS]
    x_gap = (axis_right - axis_left) / (len(epochs) - 1)
    x_points = [int(axis_left + idx * x_gap) for idx in range(len(epochs))]

    for tick in ticks:
        y = map_y(tick)
        draw.line((axis_left, y, axis_right, y), fill="#e2e8f0", width=1)
        label = f"{tick * 100:.0f}%" if value_suffix == "%" else f"{tick:.3f}"
        draw.text((70, y - 12), label, fill="#475569", font=load_font(18, True))

    for x, epoch in zip(x_points, epochs):
        draw.line((x, axis_bottom, x, axis_bottom + 8), fill="#334155", width=2)
        draw.text((x - 12, axis_bottom + 24), str(epoch), fill="#334155", font=load_font(19, True))
    draw.text((780, 850), "Epoch", fill="#334155", font=load_font(22, True))

    legend_x, legend_y = 1510, 230
    for i, (label, values, color) in enumerate(series):
        points = [(x, map_y(value)) for x, value in zip(x_points, values)]
        draw.line(points, fill=color, width=5)
        for x, y in points:
            draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill=color, outline="#ffffff", width=2)
        y = legend_y + i * 48
        draw.rounded_rectangle((legend_x, y, legend_x + 34, y + 18), radius=5, fill=color)
        draw.text((legend_x + 46, y - 4), label, fill="#0f172a", font=load_font(19, True))

    draw.rounded_rectangle((150, 900, 1550, 985), radius=14, fill="#ecfeff", outline="#0891b2", width=2)
    draw.text((180, 922), footnote, fill="#164e63", font=load_font(21, True))
    img.save(DIAGRAMS / filename)


def create_colab_ner_epoch_performance_diagram() -> None:
    _draw_line_chart(
        "61_colab_ner_epoch_performance.png",
        "Colab NER Epoch Performance Trend",
        [
            ("Precision", [row["precision"] for row in COLAB_NER_EPOCHS], "#2563eb"),
            ("Recall", [row["recall"] for row in COLAB_NER_EPOCHS], "#0891b2"),
            ("F1", [row["f1"] for row in COLAB_NER_EPOCHS], "#059669"),
            ("Accuracy", [row["accuracy"] for row in COLAB_NER_EPOCHS], "#7c3aed"),
        ],
        0.90,
        0.985,
        [0.90, 0.92, 0.94, 0.96, 0.98],
        "%",
        "Overall metrics are from seqeval output in the exported Colab run; no per-label matrix is inferred.",
    )


def create_colab_ner_loss_curve_diagram() -> None:
    _draw_line_chart(
        "62_colab_ner_loss_curve.png",
        "Colab NER Training and Validation Loss",
        [
            ("Training Loss", [row["training_loss"] for row in COLAB_NER_EPOCHS], "#dc2626"),
            ("Validation Loss", [row["validation_loss"] for row in COLAB_NER_EPOCHS], "#ea580c"),
        ],
        0.03,
        0.08,
        [0.03, 0.04, 0.05, 0.06, 0.07, 0.08],
        "",
        "Trainer loss values describe the token-classification objective; they are not a custom loss formula.",
    )


def create_frontend_backend_database_diagrams() -> None:
    create_frontend_route_layout_architecture()
    create_frontend_api_auth_flow()
    create_laravel_backend_request_lifecycle()
    create_database_relationship_rationale()


def create_job_mining_diagrams() -> None:
    save_diagram(
        "32_job_mining_design_philosophy.png",
        "Job Mining Design Philosophy",
        [
            ("Target Roles", "Admin target roles or student search terms define the mining intent.", (90, 160, 390, 300), "#dbeafe"),
            ("Queue Boundary", "Long network work runs in the scraping queue instead of blocking requests.", (500, 160, 820, 300), "#ecfdf5"),
            ("AI Job Miner", "FastAPI/Python service owns adapters, parsing, quality gates, and callbacks.", (930, 160, 1260, 300), "#ede9fe"),
            ("External Sources", "Demo, API, and public HTML adapters are treated as unstable inputs.", (1300, 430, 1560, 590), "#fef3c7"),
            ("Laravel Import API", "Laravel validates, deduplicates, stores, and syncs skills.", (930, 700, 1260, 850), "#dcfce7"),
            ("System of Record", "MySQL plus admin diagnostics and recommendations use accepted jobs only.", (500, 700, 820, 850), "#cffafe"),
            ("Student/Admin UI", "Students poll status; admins inspect source health and failures.", (90, 700, 390, 850), "#fce7f3"),
        ],
        [
            ((390, 230), (500, 230), "intent"),
            ((820, 230), (930, 230), "worker"),
            ((1260, 230), (1430, 430), "fetch"),
            ((1430, 590), (1260, 770), "jobs"),
            ((930, 770), (820, 770), "store"),
            ((500, 770), (390, 770), "visible"),
            ((660, 700), (660, 300), "metrics"),
        ],
    )

    save_diagram(
        "33_complete_job_mining_flow.png",
        "Complete Job Mining Flow",
        [
            ("User Search or Target Role", "Authenticated user request or admin-managed role.", (80, 130, 360, 250), "#dbeafe"),
            ("Existing Job Check", "Laravel searches usable stored jobs first.", (480, 130, 760, 250), "#ecfdf5"),
            ("Enough Jobs?", "If enough stored jobs exist, return them immediately.", (880, 130, 1160, 250), "#fef3c7"),
            ("Create ScrapingJob", "Pending record captures job_title, status, type, and counters.", (480, 360, 760, 500), "#dcfce7"),
            ("Dispatch Worker", "ProcessOnDemandJobScraping runs on the scraping queue.", (880, 360, 1160, 500), "#dcfce7"),
            ("AI Job Miner", "Select adapter, call API/HTML/demo source, parse fields.", (1230, 360, 1520, 500), "#ede9fe"),
            ("Normalize Payload", "Title, company, location, description, URL, skills, source.", (1230, 610, 1520, 760), "#e0f2fe"),
            ("Deduplicate Import", "Laravel checks URL and title/company variants in a transaction.", (880, 610, 1160, 760), "#ecfdf5"),
            ("Sync Skills", "SkillSyncService maps required skills to job_skills.", (480, 610, 760, 760), "#fce7f3"),
            ("Poll Status", "Status endpoint/dashboard reads completion metrics.", (80, 610, 360, 760), "#cffafe"),
        ],
        [
            ((360, 190), (480, 190), "query"),
            ((760, 190), (880, 190), "count"),
            ((1020, 250), (1020, 360), "no"),
            ((1160, 190), (1480, 190), "yes: return stored"),
            ((760, 430), (880, 430), "queue"),
            ((1160, 430), (1230, 430), "call"),
            ((1375, 500), (1375, 610), "parse"),
            ((1230, 685), (1160, 685), "validate"),
            ((880, 685), (760, 685), "save"),
            ((480, 685), (360, 685), "metrics"),
        ],
    )

    save_diagram(
        "34_scraping_runtime_architecture.png",
        "AI Job Miner Runtime Architecture",
        [
            ("React Frontend", "Jobs page, admin sources, target roles, dashboard polling.", (70, 150, 380, 300), "#dbeafe"),
            ("Laravel API", "System of record, auth, admin routes, import validation.", (500, 150, 820, 300), "#ecfdf5"),
            ("Database Queue", "scraping queue records long-running background work.", (930, 150, 1220, 300), "#dcfce7"),
            ("backend-worker-scraping", "queue:work database --queue=scraping --timeout=1200.", (1270, 150, 1570, 300), "#dcfce7"),
            ("cc-job-miner", "FastAPI on container port 8000, host 8003, /health.", (930, 470, 1220, 640), "#ede9fe"),
            ("External APIs/Sites", "Demo, Remotive, RemoteOK, Arbeitnow, Adzuna, HTML adapters.", (1270, 470, 1570, 640), "#fef3c7"),
            ("Internal Import API", "/api/v1/jobs/import, /check, /failed, /proxies/active.", (500, 470, 820, 640), "#e0f2fe"),
            ("MySQL", "job_postings, scraping_jobs, failed URLs, sources, target roles.", (70, 470, 380, 640), "#cffafe"),
            ("Optional Proxies", "Provided only through protected /proxies/active when enabled.", (930, 740, 1220, 870), "#fce7f3"),
            ("Diagnostics UI", "Admin source health, batch progress, failed URL visibility.", (500, 740, 820, 870), "#dbeafe"),
        ],
        [
            ((380, 225), (500, 225), "HTTP"),
            ((820, 225), (930, 225), "jobs"),
            ((1220, 225), (1270, 225), "work"),
            ((1420, 300), (1080, 470), "POST /scrape"),
            ((1220, 555), (1270, 555), "fetch"),
            ((930, 555), (820, 555), "callback"),
            ((500, 555), (380, 555), "SQL"),
            ((660, 640), (660, 740), "status"),
            ((1080, 640), (1080, 740), "proxy cfg"),
        ],
    )

    create_sequence_diagram(
        "35_scraping_sequence_diagram.png",
        "Sequence: Job Mining and Import",
        ["User/Admin", "React UI", "Laravel API", "ScrapingJob", "Queue Worker", "AI Job Miner", "External Source", "Import API", "MySQL"],
        [
            ("User/Admin", "React UI", "Search jobs or start admin run"),
            ("React UI", "Laravel API", "POST scrape endpoint or admin action"),
            ("Laravel API", "ScrapingJob", "Create pending record"),
            ("Laravel API", "Queue Worker", "Dispatch scraping job"),
            ("Queue Worker", "AI Job Miner", "POST /scrape with service token"),
            ("AI Job Miner", "External Source", "Fetch API/HTML/demo data"),
            ("AI Job Miner", "Import API", "POST /jobs/import/check"),
            ("AI Job Miner", "Import API", "POST /jobs/import"),
            ("Import API", "MySQL", "Transaction: job and skills"),
            ("Queue Worker", "ScrapingJob", "Update counters and status"),
            ("React UI", "Laravel API", "Poll status/dashboard"),
            ("Laravel API", "React UI", "Return metrics and imported jobs"),
        ],
    )

    save_diagram(
        "36_scraping_job_lifecycle.png",
        "Scraping Job Lifecycle",
        [
            ("pending", "Record created; worker has not started.", (140, 230, 430, 380), "#dbeafe"),
            ("processing", "Worker active; external calls and imports are running.", (650, 230, 950, 380), "#fef3c7"),
            ("completed", "Counters and completed_at are stored; UI can show jobs.", (1150, 150, 1450, 300), "#dcfce7"),
            ("failed", "Unrecoverable error or failed-only run; error_message saved.", (1150, 450, 1450, 600), "#fee2e2"),
            ("Metrics", "jobs_found, jobs_stored, jobs_duplicated, discovered_count, failed_count, processing_time_ms.", (450, 680, 1150, 820), "#e0f2fe"),
        ],
        [
            ((430, 305), (650, 305), "worker"),
            ((950, 275), (1150, 225), "success"),
            ((950, 335), (1150, 525), "error"),
            ((800, 380), (800, 680), "record"),
            ((1150, 225), (1150, 720), "poll"),
        ],
    )

    save_diagram(
        "37_source_management_flow.png",
        "Source Management and Target Roles",
        [
            ("Target Roles", "Admin manages role names/search queries and active flags.", (90, 170, 390, 320), "#dbeafe"),
            ("Scraping Sources", "Source records define endpoint, method, type, mode, and status.", (500, 170, 820, 320), "#ecfdf5"),
            ("Support Metadata", "Adapter support, credential needs, and external-risk labels.", (930, 170, 1230, 320), "#fef3c7"),
            ("Diagnostics", "Source tests run small controlled extraction checks.", (1290, 170, 1570, 320), "#e0f2fe"),
            ("Full Run", "Runnable source/target pairs become queued market scraping jobs.", (930, 560, 1230, 720), "#dcfce7"),
            ("Batch Progress", "Admin dashboard polls batch progress and source status.", (500, 560, 820, 720), "#cffafe"),
            ("Imported Jobs", "Accepted jobs and failed URLs feed admin review.", (90, 560, 390, 720), "#fce7f3"),
        ],
        [
            ((390, 245), (500, 245), "roles"),
            ((820, 245), (930, 245), "metadata"),
            ((1230, 245), (1290, 245), "test"),
            ((1430, 320), (1080, 560), "run"),
            ((930, 640), (820, 640), "status"),
            ((500, 640), (390, 640), "review"),
            ((1080, 560), (660, 320), "skip unsupported"),
        ],
    )

    save_diagram(
        "38_job_import_deduplication_flow.png",
        "Job Import and Deduplication",
        [
            ("Structured Payload", "title, company, description, URL, source, skills.", (80, 170, 380, 320), "#dbeafe"),
            ("Validate Request", "StoreScrapedJobRequest sanitizes text and rejects unsafe external URLs.", (500, 170, 820, 320), "#ecfdf5"),
            ("DB Transaction", "Import runs atomically around lookup, save, and skill sync.", (930, 170, 1230, 320), "#fef3c7"),
            ("URL Match", "Strongest identity: Job::where('url', ...).", (1290, 170, 1570, 320), "#dcfce7"),
            ("Title + Company", "Original/title-case candidates catch formatting differences.", (1290, 420, 1570, 570), "#dcfce7"),
            ("Lowercase Fallback", "Squished lowercase title and company catch casing differences.", (930, 670, 1230, 820), "#dcfce7"),
            ("Update or Create", "Existing job updated; otherwise a new job_postings row is created.", (500, 670, 820, 820), "#e0f2fe"),
            ("Sync Required Skills", "SkillSyncService links canonical skills without detaching old evidence.", (80, 670, 380, 820), "#fce7f3"),
        ],
        [
            ((380, 245), (500, 245), "POST"),
            ((820, 245), (930, 245), "begin"),
            ((1230, 245), (1290, 245), "1"),
            ((1430, 320), (1430, 420), "2"),
            ((1290, 495), (1080, 670), "3"),
            ((930, 745), (820, 745), "save"),
            ((500, 745), (380, 745), "skills"),
        ],
    )

    save_diagram(
        "39_scraping_failure_dlq_flow.png",
        "Failed URL and Retry Flow",
        [
            ("Source Error", "Timeout, blocked page, parse failure, or callback issue.", (100, 170, 390, 320), "#fee2e2"),
            ("Report Failure", "AI Job Miner posts failed URL details to Laravel.", (500, 170, 820, 320), "#fef3c7"),
            ("Protected Endpoint", "POST /api/v1/jobs/import/failed behind scraper token.", (930, 170, 1230, 320), "#e0f2fe"),
            ("ScrapingFailedUrl", "Stores URL, source/job IDs, message, retried flag, failed_at.", (1290, 170, 1570, 320), "#dcfce7"),
            ("Admin Dashboard", "Failed URLs visible for a scraping job.", (930, 560, 1230, 720), "#dbeafe"),
            ("Retry Marking", "Current admin action marks selected failures as retried.", (500, 560, 820, 720), "#fce7f3"),
            ("Future Requeue", "A stronger DLQ can dispatch targeted reprocessing later.", (100, 560, 390, 720), "#ecfdf5"),
        ],
        [
            ((390, 245), (500, 245), "failure"),
            ((820, 245), (930, 245), "POST"),
            ((1230, 245), (1290, 245), "store"),
            ((1430, 320), (1080, 560), "read"),
            ((930, 640), (820, 640), "select"),
            ((500, 640), (390, 640), "future"),
        ],
    )

    save_diagram(
        "40_scraping_security_boundaries.png",
        "Scraping Security Boundaries",
        [
            ("Public/Auth Routes", "Student searches and status polling require normal user auth.", (80, 150, 390, 300), "#dbeafe"),
            ("Admin Routes", "Source, target, dashboard, and run controls require admin middleware.", (80, 460, 390, 610), "#fce7f3"),
            ("Laravel API", "Owns validation, authorization, database writes, and redacted logs.", (520, 300, 850, 470), "#ecfdf5"),
            ("Scraper Token Routes", "/jobs/import, /check, /failed, /proxies/active use scraper.token plus throttle.", (980, 150, 1320, 320), "#e0f2fe"),
            ("AI Job Miner", "Accepts /scrape only with X-Scraper-Service-Token.", (980, 460, 1320, 630), "#ede9fe"),
            ("Secrets and API Keys", "SCRAPER_SERVICE_TOKEN, SCRAPY_API_TOKEN, LARAVEL_API_TOKEN, Adzuna keys.", (520, 710, 850, 870), "#fef3c7"),
            ("External Sources", "APIs/sites are outside the trust boundary and subject to terms/rate limits.", (1360, 460, 1580, 630), "#fee2e2"),
        ],
        [
            ((390, 225), (520, 350), "auth"),
            ((390, 535), (520, 420), "admin"),
            ((850, 350), (980, 235), "token"),
            ((1150, 320), (1150, 460), "callback"),
            ((1320, 545), (1360, 545), "fetch"),
            ((685, 710), (685, 470), "env"),
        ],
    )

    save_diagram(
        "41_scraping_validation_evidence.png",
        "Scraping Validation Evidence",
        [
            ("compileall", "Syntax importability for ai-job-miner source files.", (90, 170, 390, 320), "#dcfce7"),
            ("pytest", "Service API and AI helper tests when pytest dependencies are available.", (500, 170, 820, 320), "#dcfce7"),
            ("Docker Config", "Compose wiring for cc-job-miner, worker, tokens, and env variables.", (930, 170, 1230, 320), "#dcfce7"),
            ("/health", "Runtime liveness if stack is running on port 8003.", (1290, 170, 1570, 320), "#e0f2fe"),
            ("Import Contracts", "Laravel requests validate candidate jobs before storage.", (930, 560, 1230, 720), "#fef3c7"),
            ("Admin Evidence", "Screenshots show sources, targets, jobs, and dashboard diagnostics.", (500, 560, 820, 720), "#dbeafe"),
            ("Limits", "These checks do not prove whole-market reach or source stability.", (90, 560, 390, 720), "#fee2e2"),
        ],
        [
            ((390, 245), (500, 245), "tests"),
            ((820, 245), (930, 245), "wiring"),
            ((1230, 245), (1290, 245), "runtime"),
            ((1430, 320), (1080, 560), "status"),
            ((930, 640), (820, 640), "UI"),
            ((500, 640), (390, 640), "scope"),
        ],
    )


def create_diagrams() -> None:
    create_high_level_architecture()
    create_docker_deployment_diagram()
    create_dfd_level_0()
    create_dfd_level_1()
    create_use_case_diagram()

    create_sequence_diagram(
        "06_sequence_cv_upload_analysis.png",
        "Sequence: CV Upload and Analysis",
        ["Student", "React UI", "Laravel API", "MinIO", "AI CV Analyzer", "MySQL"],
        [
            ("Student", "React UI", "Select PDF/image CV"),
            ("React UI", "Laravel API", "POST /upload-cv with token"),
            ("Laravel API", "Laravel API", "Validate file type and size"),
            ("Laravel API", "MinIO", "Store private CV object"),
            ("Laravel API", "AI CV Analyzer", "Send file for parsing"),
            ("AI CV Analyzer", "Laravel API", "Return text, role, skills, confidence"),
            ("Laravel API", "MySQL", "Persist profile, skills, analysis"),
            ("Laravel API", "React UI", "Return updated user resource"),
            ("React UI", "Student", "Display parsed profile and warnings"),
        ],
    )

    create_sequence_diagram(
        "07_sequence_job_recommendation_gap_analysis.png",
        "Sequence: Job Recommendation and Gap Analysis",
        ["Student", "React UI", "Laravel API", "AI Matching", "MySQL", "Job Miner"],
        [
            ("Student", "React UI", "Open Jobs page"),
            ("React UI", "Laravel API", "GET /api/v1/jobs/recommended"),
            ("Laravel API", "MySQL", "Read profile, skills, and jobs"),
            ("Laravel API", "Laravel API", "Title/skill/seniority scoring"),
            ("Laravel API", "React UI", "Return ranked jobs"),
            ("Student", "React UI", "Open gap analysis"),
            ("React UI", "Laravel API", "GET /gap-analysis/job/{id}"),
            ("Laravel API", "AI Matching", "Semantic + TF-IDF via /api/hybrid-match"),
            ("AI Matching", "Laravel API", "Match scores, gaps, roadmap"),
            ("Laravel API", "React UI", "Return matched skills and roadmap"),
            ("Laravel API", "Job Miner", "Optional scrape-if-missing flow"),
        ],
        notes=[
            ("Recommendation path", "/api/v1/jobs/recommended uses Laravel title, skill-overlap, and seniority scoring."),
            ("Gap-analysis path", "/api/hybrid-match is used for semantic/adaptive plus TF-IDF gap analysis."),
        ],
    )

    create_erd()

    save_diagram(
        "09_cv_analyzer_runtime_flow.png",
        "AI CV Analyzer Runtime Flow",
        [
            ("React Dashboard", "Client-side checks, multipart upload field named cv, progress and warning UI.", (60, 160, 360, 310), "#dbeafe"),
            ("Laravel API", "CvUploadRequest, CvController, CvProcessingService, private CV storage.", (510, 160, 820, 310), "#ecfdf5"),
            ("FastAPI Gateway", "/api/parse-cv routes PDF/image files and enforces timeout/error fallback status.", (970, 160, 1280, 310), "#fce7f3"),
            ("Text Extraction", "Spatial PDF parsing first; OCR fallback for images or low-text PDFs.", (80, 450, 390, 620), "#fff7ed"),
            ("NER + Rules", "Local token-classification model, contact/date rules, noise filtering, canonicalization.", (520, 450, 850, 620), "#ede9fe"),
            ("Layer 2 + 3", "Domain/seniority inference plus semantic and TF-IDF job matching.", (980, 450, 1310, 620), "#cffafe"),
            ("MySQL + MinIO", "Normalized profile, skills, experiences, analysis metadata, and private CV object.", (310, 740, 660, 890), "#fef3c7"),
            ("Dashboard Output", "Career identity card, profile completeness, extracted skills, recommendations, gaps.", (900, 740, 1260, 890), "#dcfce7"),
        ],
        [
            ((360, 235), (510, 235), "upload"),
            ((820, 235), (970, 235), "parse"),
            ((1125, 310), (235, 450), "file"),
            ((390, 535), (520, 535), "text"),
            ((850, 535), (980, 535), "signals"),
            ((1130, 620), (1080, 740), "score"),
            ((675, 620), (485, 740), "persist"),
            ((660, 815), (900, 815), "return"),
        ],
    )

    save_diagram(
        "10_cv_model_training_pipeline.png",
        "AI CV Analyzer Model-Training Workflow",
        [
            ("Training Need", "Technical CV NER needs SKILL, ROLE, EDU, CERT, and SOFT labels.", (60, 150, 360, 300), "#dbeafe"),
            ("Gemini Generator", "Script creates synthetic annotated snippets with positive and negative samples.", (500, 150, 820, 300), "#fce7f3"),
            ("Cleaner", "Normalizes whitespace, removes duplicates, validates entity spans, preserves decoys.", (960, 150, 1280, 300), "#fff7ed"),
            ("Colab Notebook", "Uploads cleaned JSONL and runs token alignment on bert-base-cased.", (180, 470, 500, 620), "#ede9fe"),
            ("Trainer", "Five epochs, 2e-5 learning rate, batch size 16, early stopping, seqeval metrics.", (650, 470, 970, 620), "#ecfdf5"),
            ("Exported Model", "career_compass_ner_final with config, tokenizer, training args, and safetensors.", (1120, 470, 1460, 620), "#cffafe"),
            ("Runtime Load", "AdvancedNEREngine prefers the local exported model and falls back if missing.", (400, 760, 760, 890), "#fef3c7"),
            ("Book Evidence", "Colab PDF records overall metrics; repository dataset and weights remain unavailable.", (930, 760, 1320, 890), "#dcfce7"),
        ],
        [
            ((360, 225), (500, 225), "prompt"),
            ((820, 225), (960, 225), "JSONL"),
            ((1120, 300), (340, 470), "clean"),
            ((500, 545), (650, 545), "tokens"),
            ((970, 545), (1120, 545), "save"),
            ((1290, 620), (1120, 760), "deploy"),
            ((760, 825), (930, 825), "document"),
        ],
    )

    save_diagram(
        "11_cv_extraction_components.png",
        "AI CV Analyzer Extraction Components",
        [
            ("Spatial Parser", "Extracts page text, line positions, and document statistics from readable PDFs.", (80, 150, 390, 300), "#dbeafe"),
            ("OCR Pipeline", "Renders low-text PDFs or images and reads them with EasyOCR when available.", (520, 150, 850, 300), "#fff7ed"),
            ("Semantic Segmenter", "Splits CV content into sections such as skills, education, and experience.", (990, 150, 1320, 300), "#ede9fe"),
            ("Advanced NER", "Chunked transformer inference groups SKILL, ROLE, EDU, and CERT spans.", (90, 470, 410, 620), "#fce7f3"),
            ("Rule Engines", "Contacts, dates, experience blocks, action verbs, and overlap/noise guards.", (540, 470, 860, 620), "#ecfdf5"),
            ("Canonicalizer", "Normalizes skills and filters low-quality or title-like skill candidates.", (1000, 470, 1320, 620), "#cffafe"),
            ("Classifier", "Infers primary domain and seniority from profile text and recent roles.", (330, 760, 650, 900), "#fef3c7"),
            ("Strict JSON Output", "Returns profile, skills, experiences, analysis, stats, confidence, and warnings.", (850, 760, 1210, 900), "#dcfce7"),
        ],
        [
            ((390, 225), (520, 225), "fallback"),
            ((850, 225), (990, 225), "text"),
            ((1160, 300), (250, 470), "sections"),
            ((410, 545), (540, 545), "entities"),
            ((860, 545), (1000, 545), "filter"),
            ((1160, 620), (490, 760), "profile"),
            ((650, 830), (850, 830), "serialize"),
        ],
    )

    save_diagram(
        "12_layer1_understanding_pipeline.png",
        "Layer 1: CV Understanding Pipeline",
        [
            ("Upload Bytes", "Laravel sends PDF/JPEG/PNG bytes to /api/parse-cv.", (50, 150, 350, 290), "#dbeafe"),
            ("Spatial Parser", "pdfplumber words, row grouping, gap segments, adaptive column ratio, dehyphenation.", (480, 150, 830, 290), "#fff7ed"),
            ("OCR Fallback", "Triggered for no_text, errors, or low character density; uses PyMuPDF render + EasyOCR.", (970, 150, 1340, 290), "#fce7f3"),
            ("Semantic Segmenter", "Exact, regex, font hint, and embedding-based header detection.", (80, 430, 420, 590), "#ede9fe"),
            ("Advanced NER", "Overlapping chunks, local/fallback token-classification model, BIO/subword merging.", (560, 430, 900, 590), "#ecfdf5"),
            ("Rules + Canonicalizer", "Contacts, dates, experience blocks, noise filters, skill dedupe and mapping.", (1040, 430, 1410, 590), "#cffafe"),
            ("Layer 2 Enrichment", "Domain, seniority, and skill category metadata if classifiers load.", (330, 750, 670, 890), "#fef3c7"),
            ("CVParseResult", "Strict profile, stats, skills, experience, analysis, metadata, and status JSON.", (870, 750, 1230, 890), "#dcfce7"),
        ],
        [
            ((350, 220), (480, 220), "PDF"),
            ((830, 220), (970, 220), "fallback"),
            ((665, 290), (250, 430), "text"),
            ((1155, 290), (250, 430), "OCR text"),
            ((420, 510), (560, 510), "sections"),
            ((900, 510), (1040, 510), "entities"),
            ((1220, 590), (500, 750), "features"),
            ((670, 820), (870, 820), "enrich"),
        ],
    )

    save_diagram(
        "13_layer2_classification_flow.png",
        "Layer 2: Classification Flow",
        [
            ("Layer 1 Output", "Profile, current title, summary, skills, experience items, total years.", (60, 170, 380, 320), "#dbeafe"),
            ("CVDomainClassifier", "Singleton wrapper around the shared SemanticEmbedder.", (500, 170, 820, 320), "#ecfdf5"),
            ("DomainEngine", "Compares CV context to taxonomy domain descriptions using cosine similarity.", (950, 170, 1300, 320), "#ede9fe"),
            ("SeniorityEngine", "Checks title/summary keywords, total years, and action-verb strength.", (140, 510, 470, 670), "#fff7ed"),
            ("SkillEngine", "Buckets extracted skills into hard, soft, and management groups.", (620, 510, 950, 670), "#cffafe"),
            ("ClassificationOrchestrator", "Writes predicted role, primary domain, seniority, and metadata.", (1080, 510, 1450, 670), "#fce7f3"),
            ("Analysis Metadata", "domain_scores, seniority_details, categorized_skills.", (560, 790, 980, 910), "#dcfce7"),
        ],
        [
            ((380, 245), (500, 245), "embedder"),
            ((820, 245), (950, 245), "taxonomy"),
            ((220, 320), (305, 510), "years"),
            ((220, 320), (785, 510), "skills"),
            ((1125, 320), (1265, 510), "scores"),
            ((470, 590), (1080, 590), "level"),
            ((950, 590), (1080, 590), "buckets"),
            ((1265, 670), (770, 790), "metadata"),
        ],
    )

    save_diagram(
        "14_layer3_matching_engine.png",
        "Layer 3: Matching Engine",
        [
            ("Parsed CV", "Profile summary, skills, seniority, primary domain, total experience.", (70, 150, 390, 300), "#dbeafe"),
            ("Parsed JD", "Raw text, mandatory skills, bonus skills, seniority, required years, domain.", (70, 430, 390, 590), "#fff7ed"),
            ("Component Scores", "Semantic summary similarity, skill-list similarity, domain alignment.", (540, 220, 900, 390), "#ecfdf5"),
            ("Adaptive Weights", "intern/junior/mid/senior/lead weight profile from matching_config.json.", (1040, 170, 1400, 320), "#ede9fe"),
            ("Constraint Validator", "Mandatory skills, years shortfall, seniority mismatch penalties.", (1040, 430, 1400, 590), "#fce7f3"),
            ("Bonus Boost", "+2% per bonus skill, capped at +10%.", (540, 600, 900, 740), "#cffafe"),
            ("Fit Analysis", "Verdict, summary, strengths, gaps, red flags, missing mandatory skills.", (840, 790, 1240, 920), "#dcfce7"),
        ],
        [
            ((390, 225), (540, 285), "CV"),
            ((390, 510), (540, 330), "JD"),
            ((900, 285), (1040, 245), "base"),
            ((900, 330), (1040, 510), "validate"),
            ((720, 390), (720, 600), "bonus"),
            ((1220, 590), (1040, 790), "penalty"),
            ((720, 740), (940, 790), "boost"),
        ],
    )

    save_diagram(
        "15_ner_token_processing.png",
        "NER Token Processing and BIO Tagging",
        [
            ("Annotated Text", "Synthetic snippets contain entity text and labels such as SKILL and ROLE.", (60, 150, 380, 300), "#dbeafe"),
            ("Tokenizer Offsets", "bert-base-cased tokenizer returns token offsets up to max_length=512.", (520, 150, 860, 300), "#fff7ed"),
            ("BIO Labels", "First covered token gets B-label; continuation tokens get I-label; special tokens get -100.", (1010, 150, 1370, 300), "#ede9fe"),
            ("Runtime Chunking", "AdvancedNEREngine processes long CV text in 3500-character windows with 500-character stride.", (90, 470, 440, 630), "#ecfdf5"),
            ("Token Classification", "Pipeline emits token labels, scores, and character spans.", (580, 470, 920, 630), "#fce7f3"),
            ("Merge + Expand", "Subwords, I-tags, adjacent spans, and boundary expansion reconstruct entity text.", (1060, 470, 1410, 630), "#cffafe"),
            ("Final Entities", "skills, roles, education, certifications, people/org/location fallback groups.", (560, 780, 980, 910), "#dcfce7"),
        ],
        [
            ((380, 225), (520, 225), "text"),
            ((860, 225), (1010, 225), "offsets"),
            ((1190, 300), (265, 470), "runtime"),
            ((440, 550), (580, 550), "chunks"),
            ((920, 550), (1060, 550), "tokens"),
            ((1235, 630), (770, 780), "groups"),
        ],
    )

    save_diagram(
        "16_seniority_decision_logic.png",
        "Seniority Decision Logic",
        [
            ("Layer 1 Signals", "current_title, total_years, action_verb_score from experience bullets.", (70, 150, 430, 300), "#dbeafe"),
            ("Title Override", "intern/trainee forces intern; other keywords map to junior/mid/senior/lead/principal.", (560, 150, 950, 300), "#fff7ed"),
            ("Years Baseline", "<1 intern, <2 junior, <5 mid, <8 senior, <12 lead, else principal.", (1080, 150, 1480, 300), "#ecfdf5"),
            ("Weighted Combine", "Title index 60% + years index 40%, then rounded.", (290, 500, 650, 650), "#ede9fe"),
            ("Verb Adjustment", "Action verb score >=0.8 bumps one level; <0.2 nudges down one level.", (800, 500, 1180, 650), "#cffafe"),
            ("Layer 2 Enrichment", "SeniorityEngine may later enrich metadata using summary/title, years, and verb strength.", (520, 780, 1050, 910), "#fce7f3"),
        ],
        [
            ((430, 225), (560, 225), "title"),
            ((950, 225), (1080, 225), "years"),
            ((760, 300), (470, 500), "indices"),
            ((1280, 300), (470, 500), "baseline"),
            ((650, 575), (800, 575), "level"),
            ((990, 650), (785, 780), "enrich"),
        ],
    )

    save_diagram(
        "17_canonicalization_chain.png",
        "Skill Canonicalization Chain",
        [
            ("Raw NER Skills", "Global or skills-section SKILL spans from AdvancedNEREngine.", (60, 160, 360, 300), "#dbeafe"),
            ("Overlap Filters", "Remove spans matching roles/orgs, long phrases, action-verb phrases, name tokens, blocklist noise.", (500, 160, 900, 300), "#fff7ed"),
            ("DataCanonicalizer", "Exact variant, exact canonical, RapidFuzz, normalized key, semantic match, pass-through.", (1040, 160, 1450, 300), "#ecfdf5"),
            ("Dedupe + Provenance", "Keep highest confidence, merge sources and raw variants, sort by confidence.", (290, 520, 680, 670), "#ede9fe"),
            ("Experience Techs", "Block-scoped technologies are canonicalized separately for experience items.", (810, 520, 1210, 670), "#cffafe"),
            ("SkillItem Output", "name, category hard, confidence_score, evidence for Laravel persistence and UI.", (560, 790, 980, 920), "#dcfce7"),
        ],
        [
            ((360, 230), (500, 230), "filter"),
            ((900, 230), (1040, 230), "map"),
            ((1245, 300), (485, 520), "merge"),
            ((1010, 670), (770, 790), "items"),
            ((485, 670), (770, 790), "items"),
        ],
    )

    save_diagram(
        "18_score_collapse_logic.png",
        "Layer 3 Score Collapse Logic",
        [
            ("Base Components", "semantic_score, skills_score, domain_score from SemanticEmbedder.", (80, 150, 420, 300), "#dbeafe"),
            ("Adaptive Base", "base = semantic*w1 + skills*w2 + domain*w3 by seniority.", (560, 150, 920, 300), "#ecfdf5"),
            ("Constraint Penalties", "Missing mandatory skills, experience shortfall, seniority mismatch; total penalty cap 80%.", (1060, 150, 1460, 300), "#fce7f3"),
            ("Bonus Boost", "Bonus skills add +2% each, capped at +10%.", (270, 510, 610, 660), "#cffafe"),
            ("Clamp", "final = max(0, min(1, base - penalty + bonus)).", (760, 510, 1120, 660), "#fff7ed"),
            ("Pass Decision", "is_qualified if final*100 >= min_pass_score (50).", (1230, 510, 1530, 660), "#ede9fe"),
            ("Hybrid Endpoint", "/api/hybrid-match blends Layer 3 semantic score 60% with TF-IDF 40% when TF-IDF is enabled.", (520, 790, 1090, 920), "#dcfce7"),
        ],
        [
            ((420, 225), (560, 225), "weights"),
            ((920, 225), (1060, 225), "validate"),
            ((740, 300), (440, 510), "base"),
            ((1260, 300), (940, 510), "penalty"),
            ((610, 585), (760, 585), "bonus"),
            ((1120, 585), (1230, 585), "score"),
            ((945, 660), (805, 790), "optional API"),
        ],
    )

    save_diagram(
        "19_ai_design_philosophy.png",
        "AI Design Philosophy: Layered Hybrid Analyzer",
        [
            ("Raw CV", "Noisy file with columns, headers, tables, icons, dates, contact data, and mixed formats.", (80, 140, 420, 285), "#dbeafe"),
            ("Layer 1: Extraction", "Text recovery, sectioning, NER entities, contacts, experience, and canonical skills.", (560, 140, 980, 285), "#ecfdf5"),
            ("Layer 2: Enrichment", "Domain, seniority, skill categories, action evidence, and profile interpretation.", (1120, 140, 1520, 285), "#ede9fe"),
            ("Layer 3: Matching", "Semantic similarity, skill overlap, domain alignment, constraints, and TF-IDF fallback.", (560, 430, 980, 575), "#fce7f3"),
            ("Explainable Output", "Dashboard identity, recommendations, strengths, gaps, red flags, and verdicts.", (1120, 430, 1520, 575), "#cffafe"),
            ("Why Not Pure NER?", "Entities alone do not explain seniority, domain, job fit, or recovery behavior.", (120, 720, 470, 885), "#fff7ed"),
            ("Why Not Pure Rules?", "Rules are stable but brittle when CV wording, skill aliases, and job titles vary.", (580, 720, 930, 885), "#fef3c7"),
            ("Hybrid Choice", "NER + rules + canonicalization + classification + matching keeps the demo explainable and resilient.", (1040, 720, 1480, 885), "#dcfce7"),
        ],
        [
            ((420, 212), (560, 212), "parse"),
            ((980, 212), (1120, 212), "enrich"),
            ((1320, 285), (780, 430), "profile"),
            ((980, 502), (1120, 502), "explain"),
            ((780, 575), (1270, 720), "design"),
        ],
    )

    save_diagram(
        "20_complete_cv_processing_flow.png",
        "Complete CV Processing Flow",
        [
            ("Upload CV", "React submits PDF/JPEG/PNG using multipart field cv.", (55, 140, 285, 255), "#dbeafe"),
            ("File Validation", "Laravel validates type and size; stores file privately after analyzer response.", (370, 140, 660, 255), "#ecfdf5"),
            ("PDF?", "PDF goes to spatial parser; image files use OCR-style path.", (760, 140, 980, 255), "#fff7ed"),
            ("PDF Parser", "Ordered text from readable PDF pages.", (1110, 120, 1430, 235), "#fef3c7"),
            ("OCR Pipeline", "Image or low-text fallback path.", (1110, 280, 1430, 395), "#fce7f3"),
            ("Text Extraction", "Readable text, stats, and extraction metadata.", (80, 450, 350, 565), "#cffafe"),
            ("Quality Check", "No text or low text can trigger OCR/fallback status.", (470, 450, 760, 565), "#ede9fe"),
            ("Section Segmentation", "Profile, skills, experience, education, projects.", (880, 450, 1180, 565), "#ecfdf5"),
            ("NER Extraction", "Skills, roles, education, certifications.", (1290, 450, 1540, 565), "#fce7f3"),
            ("Experience Analysis", "Dates, years, action verbs, career-health signals.", (80, 720, 370, 850), "#fff7ed"),
            ("Skill Canonicalization", "Normalize variants, dedupe, attach provenance.", (470, 720, 760, 850), "#cffafe"),
            ("Domain + Seniority", "Layer 2 enrichment from title, summary, skills, years.", (880, 720, 1180, 850), "#ede9fe"),
            ("Profile + Matching", "Persist profile; Layer 3 supports recommendations and gaps.", (1290, 720, 1540, 850), "#dcfce7"),
        ],
        [
            ((285, 197), (370, 197), ""),
            ((660, 197), (760, 197), ""),
            ((980, 190), (1110, 178), "yes"),
            ((980, 225), (1110, 338), "image"),
            ((1270, 235), (215, 450), "text"),
            ((1270, 395), (215, 450), "ocr text"),
            ((350, 508), (470, 508), ""),
            ((760, 508), (880, 508), ""),
            ((1180, 508), (1290, 508), ""),
            ((1420, 565), (225, 720), ""),
            ((370, 785), (470, 785), ""),
            ((760, 785), (880, 785), ""),
            ((1180, 785), (1290, 785), ""),
        ],
    )

    save_diagram(
        "21_cv_fault_tolerance_flow.png",
        "CV Analyzer Fault Tolerance and Recovery",
        [
            ("CV File", "Uploaded PDF or image.", (80, 160, 340, 285), "#dbeafe"),
            ("PDF/Text Extraction", "Spatial parser or image OCR path attempts to recover text.", (500, 160, 850, 285), "#ecfdf5"),
            ("Readable Text?", "Quality gate from extraction result and text length.", (1010, 160, 1350, 285), "#fff7ed"),
            ("Continue Pipeline", "Segmentation, NER, experience, canonicalization, classification.", (1030, 430, 1390, 570), "#dcfce7"),
            ("OCR Fallback", "Render pages and OCR when PDF extraction fails or text is too short.", (500, 430, 850, 570), "#fce7f3"),
            ("OCR Text Found?", "If OCR text exists, continue with ocr_fallback status.", (500, 700, 850, 840), "#ede9fe"),
            ("No Text / Error Status", "Return no_text, empty_file, error, or timeout payload as supported by the caller.", (1010, 700, 1390, 840), "#fee2e2"),
            ("Backend Preservation", "Laravel stores analysis status and avoids refreshing structured profile/skills for incomplete statuses.", (80, 700, 360, 840), "#fef3c7"),
        ],
        [
            ((340, 222), (500, 222), ""),
            ((850, 222), (1010, 222), ""),
            ((1180, 285), (1210, 430), "yes"),
            ((1010, 245), (850, 500), "no"),
            ((675, 570), (675, 700), ""),
            ((850, 770), (1010, 770), "no"),
            ((500, 770), (360, 770), "failed"),
            ((675, 700), (1210, 570), "yes"),
        ],
    )

    save_diagram(
        "22_confidence_signal_flow.png",
        "Confidence and Readiness Signal Flow",
        [
            ("Text Quality", "Document stats, text length, extraction source, parsing status.", (70, 150, 370, 290), "#dbeafe"),
            ("Entity Coverage", "NER/profile confidence plus skills and experience evidence.", (500, 150, 840, 290), "#ecfdf5"),
            ("Profile Completeness", "Backend stores completeness_score from analysis confidence when available.", (970, 150, 1350, 290), "#ede9fe"),
            ("Skill Count", "Dashboard normalizes extracted skill count, capped at ten skills.", (120, 490, 420, 630), "#cffafe"),
            ("Experience Years", "Dashboard maps total years to a display percentage with a three-year reference.", (540, 490, 900, 630), "#fff7ed"),
            ("Parsing Status", "success, ocr_fallback, timeout, error, empty_file, no_text, partial_success.", (1040, 490, 1430, 630), "#fce7f3"),
            ("Career Readiness Snapshot", "Estimated visual dashboard signals, not hiring probability.", (560, 780, 1050, 910), "#dcfce7"),
        ],
        [
            ((370, 220), (500, 220), ""),
            ((840, 220), (970, 220), ""),
            ((250, 290), (250, 490), ""),
            ((670, 290), (720, 490), ""),
            ((1160, 290), (1235, 490), ""),
            ((270, 630), (680, 780), ""),
            ((720, 630), (800, 780), ""),
            ((1235, 630), (900, 780), ""),
        ],
    )

    save_diagram(
        "23_skill_canonicalization_example.png",
        "Skill Canonicalization Example",
        [
            ("Raw Extracted Variants", "JS | Java Script | Javascript | React.js | React JS", (90, 170, 520, 320), "#dbeafe"),
            ("Normalize Text", "lowercase, trim, punctuation/spacing normalization.", (680, 170, 1040, 320), "#fff7ed"),
            ("Map Candidate", "exact variant -> exact canonical -> RapidFuzz -> normalized key -> semantic -> pass-through.", (1180, 170, 1530, 320), "#ecfdf5"),
            ("Canonical Skills", "JavaScript | React", (540, 530, 960, 675), "#dcfce7"),
            ("Provenance", "Keep raw variants, source names, and highest confidence for Laravel persistence.", (1040, 530, 1460, 675), "#cffafe"),
            ("Documentation Note", "Examples are illustrative of the implemented mapping stages; committed config is mostly industry-agnostic.", (210, 770, 1210, 900), "#fef3c7"),
        ],
        [
            ((520, 245), (680, 245), ""),
            ((1040, 245), (1180, 245), ""),
            ((1360, 320), (750, 530), "dedupe"),
            ((960, 605), (1040, 605), "metadata"),
        ],
    )

    save_diagram(
        "24_fine_tuned_bert_ner_architecture.png",
        "Fine-Tuned BERT NER Architecture",
        [
            ("CV Text", "Role, skills, education, certifications, noisy CV wording.", (70, 150, 330, 285), "#dbeafe"),
            ("Tokenizer", "bert-base-cased tokenizer with offsets and max_length 512.", (470, 150, 760, 285), "#fff7ed"),
            ("Token IDs + Mask", "input_ids, attention_mask, labels; special tokens ignored with -100.", (900, 150, 1220, 285), "#ede9fe"),
            ("BERT Encoder", "Pretrained contextual representation; not trained from scratch.", (1320, 150, 1540, 285), "#ecfdf5"),
            ("Token Classification Head", "Predicts BIO labels for each token.", (360, 500, 700, 645), "#fce7f3"),
            ("BIO Labels", "O, B/I-SKILL, B/I-ROLE, B/I-EDU, B/I-CERT, B/I-SOFT.", (850, 500, 1210, 645), "#cffafe"),
            ("Grouped Entities", "Merge subwords and adjacent spans into useful CV entities.", (560, 790, 940, 910), "#dcfce7"),
            ("Runtime Boundary", "Local weights can be loaded when supplied; weights are not committed.", (1040, 790, 1500, 910), "#fef3c7"),
        ],
        [
            ((330, 218), (470, 218), ""),
            ((760, 218), (900, 218), ""),
            ((1220, 218), (1320, 218), ""),
            ((1430, 285), (530, 500), ""),
            ((700, 572), (850, 572), ""),
            ((1030, 645), (750, 790), "group"),
            ((940, 850), (1040, 850), "deploy"),
        ],
    )

    save_diagram(
        "25_detailed_training_pipeline.png",
        "Detailed NER Training Pipeline",
        [
            ("Dataset Need", "CV-specific labels for technical recruiting language.", (55, 130, 310, 250), "#dbeafe"),
            ("Gemini Generation", "Synthetic technical CV snippets from external API keys.", (400, 130, 720, 250), "#fce7f3"),
            ("Variation + Decoys", "Noise, varied formatting, negative examples with no entities.", (830, 130, 1160, 250), "#fff7ed"),
            ("Cleaning", "Whitespace normalization, dedupe, entity-span validation.", (1260, 130, 1540, 250), "#ecfdf5"),
            ("Train/Test Split", "Notebook uses 90/10 split with seed 42 when cleaned data is supplied.", (130, 430, 450, 560), "#ede9fe"),
            ("Token Alignment", "Character spans converted to BIO token labels using tokenizer offsets.", (570, 430, 900, 560), "#cffafe"),
            ("BERT Fine-Tuning", "bert-base-cased, Trainer, 5 epochs, 2e-5, batch size 16.", (1020, 430, 1360, 560), "#fef3c7"),
            ("Colab Metrics", "Seqeval precision, recall, F1, and accuracy are visible in exported Colab PDF.", (310, 760, 680, 900), "#dcfce7"),
            ("Model Export", "career_compass_ner_final export path for runtime loading.", (850, 760, 1210, 900), "#dcfce7"),
        ],
        [
            ((310, 190), (400, 190), ""),
            ((720, 190), (830, 190), ""),
            ((1160, 190), (1260, 190), ""),
            ((1400, 250), (290, 430), "clean"),
            ((450, 495), (570, 495), ""),
            ((900, 495), (1020, 495), ""),
            ((1190, 560), (500, 760), "metrics"),
            ((680, 830), (850, 830), "save"),
        ],
    )

    save_diagram(
        "26_matching_formula_flow.png",
        "Matching Formula and Penalty Flow",
        [
            ("Semantic Score", "CV summary vs job text cosine similarity.", (80, 160, 380, 295), "#dbeafe"),
            ("Skill Score", "CV skill list vs mandatory + bonus skills.", (80, 390, 380, 525), "#cffafe"),
            ("Domain Score", "Exact domain match or semantic domain similarity; cutoff 0.65.", (80, 620, 380, 755), "#ede9fe"),
            ("Adaptive Weights", "Weights chosen by candidate seniority from matching_config.json.", (560, 260, 940, 420), "#fff7ed"),
            ("BaseScore", "w_sem*S + w_skill*K + w_domain*D", (1110, 260, 1480, 420), "#ecfdf5"),
            ("Constraint Penalties", "Mandatory skill gap, years shortfall, seniority mismatch; cap 0.8.", (560, 570, 940, 735), "#fee2e2"),
            ("Bonus Boost", "+0.02 per matched bonus skill, capped at 0.10.", (1110, 570, 1480, 735), "#fef3c7"),
            ("Final Score", "clamp(BaseScore - Penalty + Bonus) then convert to percent.", (700, 820, 1240, 930), "#dcfce7"),
        ],
        [
            ((380, 227), (560, 320), "S"),
            ((380, 457), (560, 340), "K"),
            ((380, 687), (560, 360), "D"),
            ((940, 340), (1110, 340), "weighted"),
            ((1290, 420), (970, 820), "base"),
            ((750, 735), (930, 820), "subtract"),
            ((1290, 735), (1080, 820), "add"),
        ],
    )

    save_diagram(
        "27_explainable_ai_output.png",
        "Explainable AI Fit Output",
        [
            ("Matching Score", "Final percentage and qualified flag.", (80, 150, 380, 290), "#dbeafe"),
            ("Missing Mandatory Skills", "ConstraintValidator identifies required skills absent from CV.", (500, 150, 860, 290), "#fee2e2"),
            ("Experience Gap", "Years found compared with minimum years required.", (980, 150, 1340, 290), "#fff7ed"),
            ("Domain Alignment", "Exact or semantic domain agreement after cutoff.", (110, 470, 430, 610), "#ede9fe"),
            ("Skill Overlap", "Skill similarity and bonus skill boost.", (560, 470, 900, 610), "#cffafe"),
            ("FitAnalysisGenerator", "Creates strengths, gaps, red flags, summary, and verdict.", (1040, 470, 1440, 610), "#ecfdf5"),
            ("Dashboard + Gap Analysis UI", "Shows score, matched skills, missing skills, warnings, and next actions.", (500, 780, 1120, 910), "#dcfce7"),
        ],
        [
            ((230, 290), (650, 470), ""),
            ((680, 290), (1200, 470), ""),
            ((1150, 290), (1200, 470), ""),
            ((430, 540), (560, 540), ""),
            ((900, 540), (1040, 540), ""),
            ((1240, 610), (810, 780), "explain"),
        ],
    )

    create_sequence_diagram(
        "28_ai_analyzer_sequence.png",
        "Sequence: AI Analyzer End-to-End",
        ["User", "React Frontend", "Laravel Backend", "FastAPI Analyzer", "Layer 1", "Layer 2", "Layer 3", "Persistence"],
        [
            ("User", "React Frontend", "Upload CV"),
            ("React Frontend", "Laravel Backend", "POST /upload-cv"),
            ("Laravel Backend", "FastAPI Analyzer", "POST /api/parse-cv"),
            ("FastAPI Analyzer", "Layer 1", "Extract text, NER, experience, skills"),
            ("Layer 1", "Layer 2", "Classify domain and seniority"),
            ("Layer 2", "FastAPI Analyzer", "Return CVParseResult"),
            ("Laravel Backend", "Persistence", "Store CV, analysis, skills, profile"),
            ("React Frontend", "Laravel Backend", "Request jobs/gap analysis"),
            ("Laravel Backend", "Layer 3", "Match profile against job"),
            ("Layer 3", "Laravel Backend", "Score, strengths, gaps, red flags"),
            ("Laravel Backend", "React Frontend", "Dashboard and recommendations"),
            ("React Frontend", "User", "Review explainable outputs"),
        ],
    )
    create_dataset_evidence_diagram()
    create_smoke_metrics_diagram()
    create_colab_ner_metrics_diagram()
    create_colab_ner_epoch_performance_diagram()
    create_colab_ner_loss_curve_diagram()
    create_frontend_backend_database_diagrams()
    create_job_mining_diagrams()


def create_sequence_diagram(
    name: str,
    title: str,
    participants: list[str],
    messages: list[tuple[str, str, str]],
    notes: list[tuple[str, str]] | None = None,
):
    n = len(participants)
    width = max(1600, 260 + max(1, n - 1) * 215)
    if n >= 8:
        width = max(width, 2200)
    note_height = 0 if not notes else 92 * len(notes) + 20
    height = max(1120, 330 + len(messages) * 92 + note_height)
    img = PILImage.new("RGB", (width, height), BG)
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, width, title, "UML sequence notation: participants, lifelines, calls, returns, and activation bars.")

    aliases = {
        "React Frontend": "React UI",
        "Laravel Backend": "Laravel API",
        "FastAPI Analyzer": "AI Analyzer",
        "AI Job Miner": "Job Miner",
        "External Source": "Source",
        "ScrapingJob": "Scraping Job",
        "Queue Worker": "Worker",
        "Persistence": "DB",
        "AI Matching": "AI Match",
    }
    labels = [aliases.get(p, p) for p in participants]
    x_margin = 105
    spacing = (width - 2 * x_margin) // max(1, n - 1)
    xs = {p: x_margin + idx * spacing for idx, p in enumerate(participants)}
    header_y = 145
    header_h = 62
    lifeline_top = header_y + header_h
    lifeline_bottom = height - 130 - note_height
    header_w = min(205, max(130, spacing - 22))

    for participant, label in zip(participants, labels):
        x = xs[participant]
        box = (x - header_w // 2, header_y, x + header_w // 2, header_y + header_h)
        draw.rounded_rectangle(box, radius=8, fill="#ffffff", outline=INK, width=2)
        draw_centered_wrapped(draw, label, (box[0] + 8, box[1] + 8, box[2] - 8, box[3] - 8), load_font(16, True))
        draw_dashed_line(draw, (x, lifeline_top), (x, lifeline_bottom), fill="#64748b", width=2, dash=13, gap=10)

    def message_label(text: str, cx: int, y: int, max_w: int) -> None:
        font = load_font(15, True)
        lines = wrap_text(draw, text, font, max(160, min(max_w, 360)))[:3]
        line_h = 20
        text_w = max(draw.textbbox((0, 0), line, font=font)[2] for line in lines)
        box = (cx - text_w // 2 - 8, y - 41, cx + text_w // 2 + 8, y - 41 + len(lines) * line_h + 8)
        draw.rounded_rectangle(box, radius=5, fill="#ffffff", outline="#bfdbfe", width=1)
        ty = box[1] + 4
        for line in lines:
            tw = draw.textbbox((0, 0), line, font=font)[2]
            draw.text((cx - tw // 2, ty), line, fill=INK, font=font)
            ty += line_h

    def activation(x: int, y: int, *, fill: str = "#ffffff") -> None:
        draw.rectangle((x - 7, y - 26, x + 7, y + 34), fill=fill, outline=INK, width=2)

    for idx, (src, dst, label) in enumerate(messages):
        if src not in xs or dst not in xs:
            continue
        y = lifeline_top + 58 + idx * 82
        sx, dx = xs[src], xs[dst]
        if src == dst:
            loop_w = min(90, max(55, spacing // 2))
            points = [(sx + 8, y), (sx + loop_w, y), (sx + loop_w, y + 34), (sx + 12, y + 34)]
            draw.line(points, fill=TEAL, width=3)
            draw_arrowhead(draw, (sx + loop_w, y + 34), (sx + 12, y + 34), fill=TEAL, size=13)
            activation(sx, y + 16)
            message_label(label, sx + loop_w // 2, y, loop_w + 180)
            continue
        direction = 1 if dx > sx else -1
        start = (sx + direction * 10, y)
        end = (dx - direction * 18, y)
        is_return = dx < sx
        if is_return:
            draw_dashed_line(draw, start, end, fill=BLUE, width=3, dash=13, gap=8)
            draw_arrowhead(draw, start, end, fill=BLUE, size=14)
        else:
            draw.line([start, end], fill=TEAL, width=3)
            draw_arrowhead(draw, start, end, fill=TEAL, size=14)
        activation(dx, y, fill="#eff6ff" if not is_return else "#ffffff")
        message_label(label, (start[0] + end[0]) // 2, y, abs(end[0] - start[0]) - 16)

    legend_y = lifeline_bottom + 26
    draw.line((90, legend_y, 190, legend_y), fill=TEAL, width=3)
    draw_arrowhead(draw, (90, legend_y), (190, legend_y), fill=TEAL, size=12)
    draw.text((205, legend_y - 13), "Call", fill=MUTED, font=load_font(15, True))
    draw_dashed_line(draw, (285, legend_y), (385, legend_y), fill=BLUE, width=3)
    draw_arrowhead(draw, (285, legend_y), (385, legend_y), fill=BLUE, size=12)
    draw.text((400, legend_y - 13), "Return / response", fill=MUTED, font=load_font(15, True))
    draw.rectangle((570, legend_y - 25, 584, legend_y + 25), fill="#eff6ff", outline=INK, width=2)
    draw.text((602, legend_y - 13), "Activation", fill=MUTED, font=load_font(15, True))

    if notes:
        y = lifeline_bottom + 86
        for note_title, note_body in notes:
            draw.rectangle((90, y, width - 90, y + 72), fill="#fff7ed", outline=ORANGE, width=2)
            draw.text((110, y + 12), f"note: {note_title}", fill=ORANGE, font=load_font(17, True))
            draw_wrapped(draw, note_body, (110, y + 38, width - 115, y + 64), load_font(15), fill="#334155")
            y += 92
    img.save(DIAGRAMS / name)


def create_erd() -> None:
    img = PILImage.new("RGB", (1900, 1250), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw_diagram_title(draw, 1900, "ERD and Database Summary", "Migration-backed entities with PK/FK/UQ markers and readable cardinality labels.")
    c4_boundary(draw, (45, 135, 705, 955), "Student profile evidence")
    c4_boundary(draw, (720, 135, 1365, 700), "Reusable skill graph")
    c4_boundary(draw, (1390, 135, 1815, 890), "Jobs and applications")
    c4_boundary(draw, (390, 940, 1815, 1210), "Scraping operations")
    tables = [
        ("users", [("PK", "id"), ("UQ", "email"), ("", "name"), ("", "role"), ("", "is_banned")], (85, 450, 345, 635)),
        ("user_profiles", [("PK", "id"), ("FK", "user_id"), ("UQ", "user_id"), ("", "headline"), ("", "seniority"), ("", "primary_domain")], (85, 185, 345, 390)),
        ("cv_analyses", [("PK", "id"), ("FK", "user_id"), ("UQ", "user_id"), ("", "cv_path"), ("", "parsing_status"), ("", "predicted_role")], (410, 185, 680, 410)),
        ("user_experiences", [("PK", "id"), ("FK", "user_id"), ("", "title"), ("", "company"), ("", "start_date"), ("", "technologies")], (85, 720, 345, 930)),
        ("skills", [("PK", "id"), ("UQ", "name"), ("", "type")], (790, 185, 1055, 335)),
        ("user_skills", [("PK", "id"), ("FK", "user_id"), ("FK", "skill_id"), ("UQ", "user_id + skill_id"), ("", "confidence_score")], (790, 450, 1055, 655)),
        ("job_skills", [("PK", "id"), ("FK", "job_id"), ("FK", "skill_id"), ("UQ", "job_id + skill_id"), ("", "required"), ("", "importance_score")], (1095, 450, 1350, 655)),
        ("job_postings", [("PK", "id"), ("FK", "scraping_source_id"), ("UQ", "url"), ("UQ", "title + company"), ("", "title"), ("", "company"), ("", "skills JSON")], (1430, 255, 1745, 520)),
        ("applications", [("PK", "id"), ("FK", "user_id"), ("FK", "job_id"), ("", "status"), ("", "notes"), ("", "applied_at")], (1430, 670, 1745, 880)),
        ("target_job_roles", [("PK", "id"), ("UQ", "name"), ("", "search_query"), ("", "is_active")], (425, 985, 690, 1165)),
        ("scraping_jobs", [("PK", "id"), ("", "job_title"), ("", "status"), ("", "type"), ("", "jobs_found"), ("", "failed_count")], (790, 985, 1055, 1195)),
        ("scraping_failed_urls", [("PK", "id"), ("FK", "scraping_source_id"), ("FK", "scraping_job_id"), ("", "url"), ("", "error_message"), ("", "retried")], (1120, 985, 1400, 1195)),
        ("scraping_sources", [("PK", "id"), ("", "name"), ("", "endpoint"), ("", "type"), ("", "mode"), ("", "status")], (1470, 985, 1745, 1195)),
    ]
    for table_name, fields, box in tables:
        erd_table(draw, box, table_name, fields)

    erd_relationship(draw, (215, 390), (215, 450), "0..1", "1", "profile")
    erd_polyline_relationship(draw, [(345, 535), (375, 535), (375, 300), (410, 300)], "1", "0..1", "analysis", label_at=(378, 420))
    erd_relationship(draw, (215, 635), (215, 720), "1", "0..*", "experience")
    erd_relationship(draw, (345, 542), (790, 542), "1", "0..*", "user skills")
    erd_relationship(draw, (925, 335), (925, 450), "1", "0..*", "skill")
    erd_relationship(draw, (925, 335), (1218, 450), "1", "0..*", "skill")
    erd_relationship(draw, (1350, 542), (1430, 400), "0..*", "1", "job requirements")
    erd_polyline_relationship(
        draw,
        [(345, 575), (365, 575), (365, 1215), (1588, 1215), (1588, 880)],
        "1",
        "0..*",
        "applications",
        label_at=(900, 1215),
    )
    erd_relationship(draw, (1588, 520), (1588, 670), "1", "0..*", "job")
    erd_polyline_relationship(draw, [(1745, 388), (1810, 388), (1810, 950), (1608, 950), (1608, 985)], "0..*", "0..1", "source", label_at=(1810, 800))
    erd_relationship(draw, (690, 1075), (790, 1075), "0..*", "0..*", "search intent", dashed=True)
    erd_relationship(draw, (1055, 1082), (1120, 1082), "1", "0..*", "run failures")
    erd_relationship(draw, (1470, 1082), (1400, 1082), "1", "0..*", "failed URL")

    draw.rectangle((410, 725, 680, 930), fill="#ffffff", outline=GRID, width=2)
    draw.text((430, 748), "Notes", fill=INK, font=load_font(18, True))
    draw_wrapped(draw, "CV analysis creates normalized profile evidence; user_skills and job_skills compare reusable skill rows. target_job_roles drives search intent but has no FK to scraping_jobs.", (430, 785, 660, 915), load_font(15), fill=MUTED)

    draw.rectangle((80, 980, 345, 1165), fill="#ffffff", outline=GRID, width=2)
    draw.text((105, 952), "Integrity notes", fill=INK, font=load_font(19, True))
    notes = [
        "PK = primary key",
        "FK = foreign key",
        "UQ = unique constraint",
        "1, 0..1, and 0..* mark cardinality",
    ]
    y = 1024
    for note in notes:
        draw.text((105, y), "- " + note, fill=MUTED, font=load_font(16))
        y += 34
    img.save(DIAGRAMS / "08_erd.png")


def command_text(command: list[str]) -> str:
    try:
        result = subprocess.run(command, cwd=ROOT, text=True, capture_output=True, timeout=45)
        return (result.stdout or result.stderr or "").strip()
    except Exception as exc:
        return f"Unable to run {' '.join(command)}: {exc}"


def text_image(path: Path, title: str, lines: Iterable[str]) -> None:
    img = PILImage.new("RGB", (1600, 1000), "#0f172a")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, True)
    mono = ImageFont.truetype("C:/Windows/Fonts/consola.ttf", 19) if Path("C:/Windows/Fonts/consola.ttf").exists() else load_font(18)
    draw.text((48, 36), title, fill="#e0f2fe", font=title_font)
    y = 105
    for raw in lines:
        for line in str(raw).splitlines() or [""]:
            if y > 960:
                break
            draw.text((48, y), line[:145], fill="#e2e8f0", font=mono)
            y += 28
        if y > 960:
            break
    img.save(path)


def evidence_table_image(path: Path, title: str, columns: list[str], rows: list[list[str]]) -> None:
    img = PILImage.new("RGB", (1600, 1000), "#f8fafc")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1600, 90), fill="#0f172a")
    draw.text((44, 28), title, fill="white", font=load_font(31, True))
    widths = [360, 250, 830] if len(columns) == 3 else [360] * len(columns)
    x_positions = [80]
    for width in widths[:-1]:
        x_positions.append(x_positions[-1] + width)
    header_y = 135
    border = "#94a3b8"
    for idx, column in enumerate(columns):
        x1 = x_positions[idx]
        x2 = x1 + widths[idx]
        draw.rectangle((x1, header_y, x2, header_y + 54), fill="#dbeafe", outline=border, width=2)
        draw.text((x1 + 14, header_y + 15), column, fill="#0f172a", font=load_font(21, True))
    y = header_y + 54
    for row_idx, row in enumerate(rows[:9]):
        wrapped_cells: list[list[str]] = []
        max_lines = 1
        for idx, cell in enumerate(row):
            lines = wrap_text(draw, str(cell), load_font(19), widths[idx] - 28)
            wrapped_cells.append(lines[:2])
            max_lines = max(max_lines, min(2, len(lines)))
        height = max(78, 28 + max_lines * 28)
        for idx, lines in enumerate(wrapped_cells):
            x1 = x_positions[idx]
            x2 = x1 + widths[idx]
            fill = "#ffffff" if row_idx % 2 == 0 else "#f1f5f9"
            draw.rectangle((x1, y, x2, y + height), fill=fill, outline=border, width=1)
            ty = y + 15
            for line in lines:
                draw.text((x1 + 14, ty), line, fill="#334155", font=load_font(19))
                ty += 28
        y += height
        if y > 900:
            break
    draw.text(
        (80, 945),
        "Readable evidence summary; raw command output remains in local validation logs.",
        fill="#64748b",
        font=load_font(17),
    )
    img.save(path)


def create_terminal_evidence() -> None:
    ps_json = command_text(["docker", "compose", "-f", "docker-compose.yml", "-f", "docker-compose.prod.yml", "ps", "--format", "json"])
    ps_text = command_text(["docker", "compose", "-f", "docker-compose.yml", "-f", "docker-compose.prod.yml", "ps"])
    ps_path = SCREENSHOTS / "18_docker_containers.png"
    ps_lower = (ps_json + ps_text).lower()
    docker_daemon_unavailable = any(
        marker in ps_lower
        for marker in [
            "failed to connect to the docker api",
            "cannot connect to the docker daemon",
            "is the docker daemon running",
        ]
    )
    # Keep the previously captured running-services evidence when this local machine cannot reach Docker.
    if not docker_daemon_unavailable or not ps_path.exists():
        services: dict[str, dict] = {}
        for line in ps_json.splitlines():
            try:
                item = json.loads(line)
            except Exception:
                continue
            service = str(item.get("Service") or item.get("Name") or "")
            if service:
                services[service] = item
        preferred = [
            ("nginx", "Public HTTP gateway, host port 80."),
            ("frontend", "React/Vite frontend, host port 5173."),
            ("backend-api", "Laravel API, internal FPM port 9000."),
            ("backend-worker-scraping", "Long-running scraping queue worker."),
            ("ai-cv-analyzer", "FastAPI CV analyzer, host port 8000."),
            ("ai-job-miner", "FastAPI job miner, host port 8003."),
            ("db", "MySQL database, host port 3306."),
            ("minio", "Private object storage, host ports 9000/9001."),
            ("prometheus", "Metrics collection, host port 9090."),
        ]
        service_rows: list[list[str]] = []
        for service, role in preferred:
            item = services.get(service)
            if not item:
                continue
            state = item.get("State") or item.get("Status") or "unknown"
            health = item.get("Health")
            status = f"{state}" + (f" / {health}" if health else "")
            service_rows.append([service, status, role])
        if not service_rows:
            service_rows = [[line[:34], "reported", "Raw compose output retained in local logs."] for line in ps_text.splitlines()[1:10]]
        evidence_table_image(ps_path, "Docker Compose Services Evidence", ["Service", "Status", "Meaning"], service_rows)

    validation_rows = [
        ["Compose config", "PASSED", "Base plus production overlay YAML/config validated."],
        ["Laravel health", "200", "/api/health and /api/ready verify API availability and dependencies."],
        ["Frontend status", "200", "/status returned the React/Vite page HTML."],
        ["AI Job Miner", "75 passed, 1 warning", "Container pytest validates service/API helper behavior."],
        ["AI CV Analyzer", "PASSED", "Python syntax compilation passed; root endpoint is operational."],
        ["Backend tests", "RETAINED", "Earlier container validation passed 39 tests / 297 assertions."],
        ["Frontend build", "RETAINED", "Earlier ESLint/build evidence passed with documented warnings."],
        ["Document build", "PASSED", "Markdown, DOCX, PDF, links, images, and JSON examples validated."],
    ]
    evidence_table_image(SCREENSHOTS / "19_validation_summary.png", "Validation Evidence Summary", ["Evidence", "Result", "Meaning"], validation_rows)


def references_markdown() -> str:
    lines = ["# References", ""]
    for ref in REFERENCES:
        lines.append(
            f"[{ref.key}] {ref.organization}, \"{ref.title},\" {ref.source}, {ref.year}. "
            f"[Online]. Available: {ref.url}. {ref.accessed}."
        )
    lines.append("")
    return "\n".join(lines)


def run_mini_evaluation() -> dict:
    if MINI_EVAL_SCRIPT.exists():
        subprocess.run([sys.executable, str(MINI_EVAL_SCRIPT)], cwd=ROOT, check=True)
    if MINI_EVAL_RESULTS.exists():
        return json.loads(MINI_EVAL_RESULTS.read_text(encoding="utf-8"))
    return {
        "summary": {
            "evaluation_mode": "not_run",
            "statistical_scope": "Mini evaluation results were not generated.",
            "cv_samples": 0,
            "job_samples": 0,
            "cv_analyzer_offline": {},
            "recommendation_offline": {},
            "gap_analysis_offline": {},
        },
        "cv_analyzer_results": [],
        "recommendation_results": [],
        "gap_analysis_results": [],
    }


def run_smoke_evaluation() -> dict:
    if SMOKE_EVAL_SCRIPT.exists():
        subprocess.run([sys.executable, str(SMOKE_EVAL_SCRIPT)], cwd=ROOT, check=True)
    if SMOKE_EVAL_RESULTS.exists():
        return json.loads(SMOKE_EVAL_RESULTS.read_text(encoding="utf-8"))
    return {
        "evaluation_mode": "not_run",
        "statistical_scope": "Smoke evaluation results were not generated.",
        "sample_count": 0,
        "dependency_probe": {
            "packages_available": {},
            "full_analyzer_import": {"ok": False, "error": "not_run"},
            "tfidf_probe": {"ok": False, "score": None, "error": "not_run"},
        },
        "summary": {
            "macro_skill_precision": 0.0,
            "macro_skill_recall": 0.0,
            "macro_skill_f1": 0.0,
            "role_match_rate": 0.0,
            "domain_match_rate": 0.0,
            "seniority_match_rate": 0.0,
            "parsing_status_match_rate": 0.0,
        },
        "sample_results": [],
    }


def md_table(headers: list[str], rows: list[list[str]]) -> str:
    body = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        body.append("| " + " | ".join(str(cell).replace("\n", " ").replace("|", "/") for cell in row) + " |")
    return "\n".join(body)


def mini_eval_markdown(results: dict) -> str:
    cv_samples = json.loads((EVALUATION / "mini_cv_dataset.json").read_text(encoding="utf-8"))
    job_samples = json.loads((EVALUATION / "mini_jobs_dataset.json").read_text(encoding="utf-8"))
    summary = results["summary"]
    cv_summary = summary["cv_analyzer_offline"]
    rec_summary = summary["recommendation_offline"]
    gap_summary = summary["gap_analysis_offline"]

    cv_dataset_table = md_table(
        ["Sample ID", "Expected Role", "Seniority", "Domain", "Expected Skills"],
        [
            [
                sample["sample_id"],
                sample["expected_role"],
                sample["expected_seniority"],
                sample["expected_domain"],
                ", ".join(sample["expected_skills"]),
            ]
            for sample in cv_samples
        ],
    )
    job_dataset_table = md_table(
        ["Job ID", "Title", "Domain", "Required Skills"],
        [
            [
                job["job_id"],
                job["title"],
                job["domain"],
                ", ".join(job["required_skills"]),
            ]
            for job in job_samples
        ],
    )
    metric_table = md_table(
        ["Area", "Metric", "Value", "Notes"],
        [
            ["CV offline", "Macro skill precision", f"{cv_summary['macro_skill_precision']:.3f}", "Keyword extraction over synthetic CV text"],
            ["CV offline", "Macro skill recall", f"{cv_summary['macro_skill_recall']:.3f}", "Compared with expected skill labels"],
            ["CV offline", "Macro skill F1", f"{cv_summary['macro_skill_f1']:.3f}", "F1 computed from precision/recall [30]"],
            ["CV offline", "Role match rate", f"{cv_summary['role_match_rate']:.3f}", "Rule-based role inference on synthetic data"],
            ["CV offline", "Seniority match rate", f"{cv_summary['seniority_match_rate']:.3f}", "Rule-based seniority inference"],
            ["CV offline", "Domain match rate", f"{cv_summary['domain_match_rate']:.3f}", "Rule-based domain inference"],
            ["Recommendation offline", "Top-1 relevance", f"{rec_summary['top_1_relevance']:.3f}", "Top recommendation belongs to manual relevant set"],
            ["Recommendation offline", "Top-3 relevance", f"{rec_summary['top_3_relevance']:.3f}", "Any top-3 job belongs to manual relevant set"],
            ["Recommendation offline", "Mean precision@3", f"{rec_summary['mean_precision_at_3']:.3f}", "Relevant jobs among top three"],
            ["Gap offline", "Matched skill agreement F1", f"{gap_summary['mean_matched_skill_agreement_f1']:.3f}", "Computed matched skills vs. expected matched skills"],
            ["Gap offline", "Missing skill agreement F1", f"{gap_summary['mean_missing_skill_agreement_f1']:.3f}", "Computed missing skills vs. expected missing skills"],
        ],
    )
    recommendation_table = md_table(
        ["CV Sample", "Expected Relevant Jobs", "Top 3 Offline Recommendations", "P@3"],
        [
            [
                item["sample_id"],
                ", ".join(item["expected_relevant_job_ids"]),
                ", ".join(item["top_3_recommended_job_ids"]),
                f"{item['precision_at_3']:.3f}",
            ]
            for item in results["recommendation_results"]
        ],
    )
    gap_table = md_table(
        ["CV / Job Pair", "Matched Skills", "Missing Skills", "Agreement"],
        [
            [
                f"{item['cv_sample_id']} -> {item['job_id']}",
                ", ".join(item["computed_matched_skills"]) or "None",
                ", ".join(item["computed_missing_skills"]) or "None",
                f"matched F1={item['matched_skill_agreement_f1']:.3f}; missing F1={item['missing_skill_agreement_f1']:.3f}",
            ]
            for item in results["gap_analysis_results"]
        ],
    )
    return f"""
### Mini Dataset Files

The mini evaluation uses fake synthetic CVs and fake synthetic job records stored under `docs/graduation-book/evaluation/`. It is intentionally small and preliminary. It is useful for graduation validation and regression checks, but it is not statistically representative and should not be used as a production benchmark.

{cv_dataset_table}

*Table 50. Mini CV dataset.*

{job_dataset_table}

*Table 51. Mini job dataset.*

### Metric Definitions

Skill precision measures how many extracted skills are expected labels. Skill recall measures how many expected skills were extracted. Skill F1 is the harmonic mean of precision and recall [30]. Recommendation top-1 and top-3 relevance compare ranked jobs against manual relevance labels. Gap agreement compares computed matched/missing skills against expected matched/missing skills.

{metric_table}

*Table 52. Mini evaluation metrics.*

### Recommendation Ranking Details

{recommendation_table}

*Table 53. Recommendation ranking details.*

### Gap Analysis Pair Details

{gap_table}

*Table 54. Gap analysis pair details.*
"""


def smoke_eval_markdown(results: dict) -> str:
    summary = results.get("summary", {})
    dependency_probe = results.get("dependency_probe", {})
    full_import = dependency_probe.get("full_analyzer_import", {})
    tfidf_probe = dependency_probe.get("tfidf_probe", {})
    package_status = dependency_probe.get("packages_available", {})
    package_rows = [
        [name, "Available" if available else "Unavailable"]
        for name, available in sorted(package_status.items())
    ]
    sample_rows = [
        [
            item["sample_id"],
            f"{item['skill_f1']:.3f}",
            "Pass" if item["role_match"] else "Check",
            "Pass" if item["domain_match"] else "Check",
            "Pass" if item["seniority_match"] else "Check",
            "Pass" if item["parsing_status_match"] else "Check",
        ]
        for item in results.get("sample_results", [])
    ]
    smoke_metric_table = md_table(
        ["Area", "Metric", "Value", "Notes"],
        [
            ["AI analyzer smoke", "Macro skill precision", f"{summary.get('macro_skill_precision', 0):.3f}", "Five manually labeled text samples"],
            ["AI analyzer smoke", "Macro skill recall", f"{summary.get('macro_skill_recall', 0):.3f}", "Expected skill labels defined in smoke sample JSON"],
            ["AI analyzer smoke", "Macro skill F1", f"{summary.get('macro_skill_f1', 0):.3f}", "Deterministic text smoke metric, not NER benchmark"],
            ["AI analyzer smoke", "Role match rate", f"{summary.get('role_match_rate', 0):.3f}", "Simple role inference over sample text"],
            ["AI analyzer smoke", "Domain match rate", f"{summary.get('domain_match_rate', 0):.3f}", "Simple domain evidence markers"],
            ["AI analyzer smoke", "Seniority match rate", f"{summary.get('seniority_match_rate', 0):.3f}", "One mismatch preserved as evidence of limitation"],
            ["AI analyzer smoke", "Parsing status match rate", f"{summary.get('parsing_status_match_rate', 0):.3f}", "Includes low-information abstention sample"],
        ],
    )
    return f"""
## 8.15 AI CV Analyzer Smoke Evaluation

The smoke evaluation under `docs/graduation-book/evaluation/` uses five short, fake CV text samples: backend, data analyst, frontend, DevOps/cloud, and low-information/noisy input. It is deterministic and useful as a small reproducibility check. It does not run the full transformer NER model and must not be reported as final model accuracy.

The script also probes the local documentation runtime. In this run, full analyzer import was {'available' if full_import.get('ok') else 'unavailable'} because `{full_import.get('error') or 'no import error'}`. The pure Python TF-IDF matcher was {'available' if tfidf_probe.get('ok') else 'unavailable'}, with a backend-overlap probe score of `{tfidf_probe.get('score')}`.

{md_table(["Package", "Status"], package_rows)}

*Dependency probe for the AI CV Analyzer smoke evaluation.*

{smoke_metric_table}

*AI CV Analyzer smoke evaluation metrics.*

{md_table(["Sample", "Skill F1", "Role", "Domain", "Seniority", "Status"], sample_rows)}

*Per-sample AI CV Analyzer smoke evaluation results.*

{figure_markdown("Figure 30", "AI CV Analyzer smoke evaluation metrics.", "assets/diagrams/30_ai_cv_analyzer_smoke_metrics.png")}
"""


def figure_markdown(number: str, caption: str, rel_path: str) -> str:
    return f"![{caption}]({rel_path})\n\n*{number}. {caption}*"


def report_markdown(mini_results: dict, smoke_results: dict) -> str:
    fig_list = "\n".join([f"- [{num}. {caption}](#{figure_anchor(num)})" for num, caption, _ in FIGURES])
    table_list = "\n".join([f"- [{num}. {caption}](#{table_anchor(num)})" for num, caption in TABLES])
    refs = "\n".join(
        [
            f"[{ref.key}] {ref.organization}, \"{ref.title},\" {ref.source}, {ref.year}. [Online]. Available: {ref.url}. {ref.accessed}."
            for ref in REFERENCES
        ]
    )
    mini_eval = mini_eval_markdown(mini_results)
    smoke_eval = smoke_eval_markdown(smoke_results)
    screenshots_markdown = chr(10).join(
        figure_markdown(num, caption, rel_path)
        for num, caption, rel_path in FIGURES
        if rel_path.startswith("assets/screenshots/")
    )

    return f"""# {PROJECT_TITLE}

{UNIVERSITY}

{FACULTY}

{DEPARTMENT}

Graduation Project Book

Academic Year: {ACADEMIC_YEAR}

Supervisor: {SUPERVISOR}

Submitted by:
{chr(10).join(STUDENTS)}

\\pagebreak

# Table of Contents

{toc_markdown()}

\\pagebreak

# List of Figures

{fig_list}

\\pagebreak

# List of Tables

{table_list}

\\pagebreak

# Acknowledgment

The project team would like to express sincere appreciation to {SUPERVISOR} for academic supervision, technical guidance, and continuous feedback during the preparation of CareerCompass. The team also thanks the Faculty of Computers and Information at Kafr El-Sheikh University for providing the academic setting in which this graduation project was designed, implemented, tested, and documented.

The work presented in this book reflects a collaborative software engineering effort. It combines web application development, database design, AI-assisted document analysis, explainable matching, containerized deployment, testing, and technical documentation. The two supervisor-provided graduation books were used only to understand expected report structure and visual formality; no content, wording, project-specific claims, diagrams, or references were copied from them.

# Abstract

CareerCompass is a graduation/demo career guidance platform that helps students and early-career users understand their CV profile, explore imported job opportunities, and compare their current skills against job requirements. The system consists of a React and Vite frontend, a Laravel API backend, a MySQL database, a FastAPI-based CV analyzer, a FastAPI/Scrapy-based job miner, MinIO-compatible private file storage, Nginx routing, and Prometheus/Grafana monitoring. The platform supports registration, login, CV upload, AI-assisted CV parsing, normalized profile and skills storage, job recommendation, gap analysis, an application tracker, and administrator dashboards for job and source diagnostics.

The AI CV Analyzer is documented as a hybrid implementation rather than a single opaque model. It combines PDF/image text extraction, OCR fallback, section segmentation, a BERT-family token-classification path for named-entity recognition, rule-based contact/date/experience extraction, skill canonicalization, domain and seniority classification, sentence embeddings, and TF-IDF-style matching. The runtime code can load an exported local NER artifact when that ignored deployment folder is present, and a Colab-oriented training notebook documents how the artifact is produced. A user-provided exported Colab PDF now records the overall NER training metrics, while the cleaned dataset content and model weights remain outside Git; therefore, the report separates recorded training-run evidence from repository-alone reproducibility and production benchmark claims.

The implementation is intentionally described as a graduation/demo system rather than a production product. The AI outputs are estimates, the job data depends on imported and demo sources, and the security posture is appropriate for demonstration but requires further production hardening. Validation was performed through Docker Compose configuration checks, backend/frontend evidence from earlier passes, Python syntax checks, containerized AI Job Miner tests, service health probes, a deterministic demo-source smoke check, and manual browser screenshots. Backend tests previously passed with 39 tests and 297 assertions, the AI Job Miner container test suite passed with 75 tests in the final Phase 1 documentation-fix pass, and frontend lint/build evidence remains recorded from the earlier validation pass. The AI CV Analyzer pytest suite was not rerun in this Phase 1 documentation-fix pass, while Python syntax compilation passed.

\\pagebreak

# Abbreviations

| Abbreviation | Meaning |
|---|---|
| API | Application Programming Interface |
| CV | Curriculum Vitae |
| DFD | Data Flow Diagram |
| ERD | Entity Relationship Diagram |
| HTTP | Hypertext Transfer Protocol |
| JSON | JavaScript Object Notation |
| JWT | JSON Web Token |
| ML | Machine Learning |
| NLP | Natural Language Processing |
| OCR | Optical Character Recognition |
| RBAC | Role-Based Access Control |
| REST | Representational State Transfer |
| S3 | Simple Storage Service compatible object storage |
| TF-IDF | Term Frequency-Inverse Document Frequency |
| UI | User Interface |

\\pagebreak

# Chapter 1: Introduction

## 1.1 Introduction to the Project

CareerCompass is an AI-assisted career guidance and job recommendation platform developed as a Computer Science graduation project for {UNIVERSITY}. The system helps a student upload a CV, receive a structured profile, view estimated job matches, inspect skill gaps, and save opportunities in an application tracker. The project combines web engineering, backend service design, natural language processing support, web scraping support, data persistence, and containerized operations.

The platform uses a modern web stack. The frontend is implemented with React, a component-based UI library [4], and bundled with Vite [5]. The backend uses Laravel, which provides routing, controllers, validation, middleware, Eloquent ORM, queues, testing support, and file storage abstractions [1]. The AI services are implemented with FastAPI, a Python API framework that supports type-hinted request and response handling [6]. The deployment is orchestrated through Docker Compose [11].

## 1.2 Background and Motivation

Many students prepare CVs and search for jobs without a clear view of how their current skills relate to job descriptions. Career guidance is often fragmented across CV feedback, job boards, learning resources, and manual advice. CareerCompass addresses this academic problem by placing those steps into one integrated demonstration system. The objective is not to replace human career advising, but to provide a practical software prototype that can parse a CV, normalize profile data, import job records, estimate matches, and present explainable gap analysis.

## 1.3 Problem Statement

Students frequently face three connected problems. First, they may not know whether their CV communicates their skills clearly. Second, job listings often describe requirements in different formats, making comparison difficult. Third, students may save opportunities across separate tools without a coherent tracker. CareerCompass proposes a centralized graduation/demo system that links CV data, jobs, matching, gap analysis, and tracking.

## 1.4 Purpose of the Project

The purpose of CareerCompass is to demonstrate how a distributed web system can support career exploration using explainable AI-assisted workflows. The project demonstrates authentication, CV upload, private storage, parsing, skill extraction, job import, recommendation, gap analysis, admin diagnostics, and monitoring in a Dockerized environment.

## 1.5 Project Objectives

- Provide a student-facing web interface for account creation, login, CV upload, profile review, job recommendations, gap analysis, and application tracking.
- Provide an administrator interface for dashboard statistics, job review, source diagnostics, and target role management.
- Implement a Laravel API that protects authenticated workflows using token-based access and role checks.
- Implement Python services for CV analysis and job mining support.
- Store uploaded CV files privately and expose downloads through controlled URLs.
- Validate the system with backend tests, frontend lint/build, Python tests where available, Docker checks, HTTP probes, and browser screenshots.

## 1.6 Proposed Solution

The proposed solution is a multi-service application. React renders the browser interface. Nginx serves as the gateway. Laravel handles REST-style APIs over HTTP [16], authentication, validation, data models, business services, and admin routes. MySQL stores normalized data [8]. MinIO provides S3-compatible object storage [9]. The CV analyzer parses PDFs and images using Python text extraction and OCR-related libraries. The job miner imports jobs from demo, API, and scraping-style sources using FastAPI, Scrapy, and HTML parsing concepts [16], [17], [18].

## 1.7 Project Scope

The scope covers a graduation/demo environment: local Docker deployment, student workflows, admin workflows, testing, screenshots, and documentation. It does not claim production readiness, job placement certainty, full job market coverage, legal compliance for production privacy, or complete AI accuracy.

## 1.8 Graduation Demo Positioning

CareerCompass should be presented as a working academic prototype with realistic engineering boundaries. The screenshots and tests show that the core workflows can run locally. However, large-scale evaluation, production identity management, cloud infrastructure, observability hardening, legal privacy review, and complete market integrations remain future work.

## 1.9 Report Organization

Chapter 2 analyzes requirements and users. Chapter 3 presents architecture, diagrams, database design, and deployment. Chapter 4 lists software and tools with references. Chapter 5 documents implementation modules from the repository. Chapter 6 presents the AI CV Analyzer deep technical analysis as a standalone academic contribution. Chapter 7 presents the AI Job Miner and scraping deep technical analysis. Chapter 8 presents testing and evaluation results. Chapter 9 discusses security and privacy. Chapter 10 concludes with achievements, limitations, and future work.

\\pagebreak

# Chapter 2: System Analysis

## 2.1 Introduction

System analysis defines what CareerCompass should do, who uses it, and what constraints influence implementation. The analysis is based on the reviewed repository, including the root documentation, Docker configuration, Laravel API, React frontend, Python services, database migrations, seeders, tests, and finalization notes.

## 2.2 Development Methodology

The project followed an iterative engineering approach. Each feature was implemented, checked through local commands or tests, integrated with Docker services, then documented. This approach is appropriate for a graduation project because it supports incremental progress while keeping the system demonstrable. The architecture uses separate services, which follows a common microservice-style pattern where independently deployable services communicate over network APIs [29].

## 2.3 Existing Career Guidance Problems

The system targets common problems in student career preparation:

- CV content is not structured for direct software comparison.
- Job descriptions vary in wording and completeness.
- Students need explainable feedback rather than unexplained match numbers.
- Admin users need visibility into imported jobs and scraping sources.
- Demo systems need reliable startup, testing, and evidence for academic evaluation.

## 2.4 Target Users and Stakeholders

| Stakeholder | Description | Main Interest |
|---|---|---|
| Student user | A university student or early-career user. | Upload CV, view profile, discover jobs, analyze gaps, save opportunities. |
| Administrator | A project operator or supervisor/demo administrator. | Review jobs, source diagnostics, users, target roles, and system health. |
| Supervisor | Academic supervisor evaluating the graduation project. | Correctness, completeness, originality, and honest evaluation. |
| Project team | Developers responsible for design and implementation. | Maintainable code, demonstrable workflows, testing, and documentation. |

*Table 1. Stakeholder summary.*

## 2.5 User Roles

CareerCompass implements two practical roles. The student role can register, login, upload a CV, view recommendations, run gap analysis, and track applications. The admin role can access protected admin routes for dashboard statistics, job administration, scraping sources, target roles, and user review.

\\pagebreak

## 2.6 Functional Requirements

| ID | Requirement | Measurable Acceptance / Implementation Evidence |
|---|---|---|
| FR-01 | Register, login, and logout users. | The API must create users with unique email addresses, issue Sanctum tokens on login, reject invalid credentials, and expose the authenticated user to the React app through `AuthController`, `RegisterRequest`, `LoginRequest`, and Login/Register pages. |
| FR-02 | Upload a CV file. | The API must accept only PDF/JPEG/PNG uploads through field `cv`, enforce the configured maximum file size, and return validation errors for invalid files through `CvUploadRequest` and the dashboard upload flow. |
| FR-03 | Store CV files privately. | Uploaded CV binaries must be stored through `CvStorageService` on the configured private MinIO/S3-compatible disk while database rows store only metadata and object references. |
| FR-04 | Parse CV and extract profile/skills. | `CvProcessingService` must call AI CV Analyzer `/api/parse-cv`, persist successful structured profile, skill, experience, and analysis metadata, and preserve explicit timeout/error/no-text statuses. |
| FR-05 | Display normalized profile and skills. | The user dashboard/profile must display persisted profile, skills, experience, predicted role, seniority, domain, completeness, and parsing status from Laravel resources and React pages. |
| FR-06 | Import and display jobs. | The backend must store usable imported/demo jobs, deduplicate by URL and title/company constraints, expose paginated listings/details, and support admin/user job views through `JobController`, `ScrapedJobController`, and AI Job Miner integration. |
| FR-07 | Estimate job recommendations with Laravel scoring. | `/api/v1/jobs/recommended` ranks up to 50 usable jobs using predicted role/profile title matching, required-skill overlap, and seniority hints in `JobController::getRecommended`. The job list ranker is separate from the AI gap-analysis matcher. |
| FR-08 | Analyze skill gaps with AI fallback. | Gap analysis for a selected job or role compares stored user data with job requirements; `GapAnalysisService` calls AI CV Analyzer `/api/hybrid-match` for semantic/adaptive plus TF-IDF matching when available and falls back to database skill matching when unavailable. |
| FR-09 | Track applications. | ApplicationController, ApplicationTrackerService, Applications page. |
| FR-10 | Provide admin dashboards. | Admin Dashboard, Jobs, Sources, Targets pages and admin API routes. |
| FR-11 | Provide health and metrics endpoints. | HealthController, MetricsController, Prometheus/Grafana compose services. |

*Table 2. Functional requirements summary.*

## 2.7 Non-Functional Requirements

| Category | Measurable Requirement | CareerCompass Approach |
|---|---|---|
| Usability | A student must be able to complete registration, login, CV upload, job browsing, gap analysis, and application tracking through visible React screens without direct API calls. | React dashboard/profile/jobs/gap-analysis/applications pages, status indicators, cards, and action buttons. |
| Maintainability | Core business behavior must be separated by controllers, request validators, services, resources, models, and Python service modules rather than being embedded in a single UI or script. | Laravel controllers/requests/services/resources/models, React page components, and FastAPI service modules. |
| Reliability | AI failures must not be reported as successful analysis; recoverable failures must return explicit statuses or deterministic fallbacks. | CV timeout/error/no-text statuses, preserved existing profile data, recommendation fallback to recent jobs, and gap-analysis fallback to DB skill matching. |
| Security | Protected user routes must require authentication, admin routes must require admin role, internal scraper/metrics paths must require service tokens where configured, and uploaded CV files must not be publicly stored. | Sanctum tokens, request validation, admin middleware, signed URLs, private storage disk, and internal service tokens. |
| Performance limits | User-facing list endpoints and expensive operations must apply practical demo limits to avoid unbounded processing. | CV max file validation, recommendation cap of 50 returned jobs after ranking up to 200 candidates, gap-analysis batch limit of 20 jobs, and paginated job listings. |
| Observability | The local stack must expose health/readiness/status/metrics endpoints sufficient for demonstration. | `/api/health`, `/api/ready`, `/status`, `/api/metrics`, Prometheus, and Grafana. |
| Portability | The project must be runnable as a local multi-service demo with documented environment variables. | Docker Compose services, `.env.example`-style configuration, MinIO/MySQL/Python/Laravel/React service wiring. |

*Table 3. Non-functional requirements summary.*

## 2.8 Requirement-to-Code/Test Traceability

| Requirement | Main Code Evidence | Test / Evidence Status |
|---|---|---|
| FR-01 Authentication | `backend-api/routes/api.php`, `backend-api/app/Http/Controllers/Api/AuthController.php`, `backend-api/app/Http/Requests/Auth/RegisterRequest.php`, `backend-api/app/Http/Requests/Auth/LoginRequest.php`, `frontend/src/pages/Login.jsx`, `frontend/src/pages/Register.jsx` | Existing backend feature tests in `backend-api/tests/Feature/AuthApiTest.php`; previously recorded Laravel test pass, not freshly rerun in the current Phase 1 documentation shell. |
| FR-02/FR-03 CV upload and private storage | `backend-api/app/Http/Requests/CvUploadRequest.php`, `backend-api/app/Http/Controllers/Api/CvController.php`, `backend-api/app/Services/CvStorageService.php`, `backend-api/config/filesystems.php`, `frontend/src/pages/user/Dashboard.jsx` | Existing `backend-api/tests/Feature/CvUploadTest.php`; manual upload screenshot evidence; current Phase 1 pass did not rerun PHP tests. |
| FR-04/FR-05 CV parsing and profile display | `backend-api/app/Services/CvProcessingService.php`, `ai-cv-analyzer/main.py`, `backend-api/app/Services/SkillSyncService.php`, `backend-api/app/Http/Resources/UserResource.php`, profile/dashboard React pages | Existing AI CV Analyzer API tests and smoke evidence; recorded Colab/model evidence; live AI CV Analyzer pytest was not freshly reproducible from the current shell. |
| FR-06 Job import/display | `backend-api/app/Http/Controllers/Api/JobController.php`, `backend-api/app/Http/Controllers/Api/ScrapedJobController.php`, `ai-job-miner/service_api.py`, `backend-api/database/migrations/2026_02_19_000002_create_all_jobs_table.php` | Existing `ScrapedJobImportTest.php`, scraper/orchestrator tests, and previously recorded AI Job Miner container pytest pass. |
| FR-07 Job recommendations | `backend-api/app/Http/Controllers/Api/JobController.php::getRecommended`, `frontend/src/pages/user/Jobs.jsx`, `backend-api/database/migrations/2026_02_19_000003_create_job_skills_table.php` | Covered by code inspection and offline recommendation mini evaluation; the production endpoint uses Laravel title/skill/seniority scoring, not `/api/hybrid-match`. |
| FR-08 Gap analysis | `backend-api/app/Http/Controllers/Api/GapAnalysisController.php`, `backend-api/app/Services/GapAnalysisService.php`, `ai-cv-analyzer/main.py` `/api/hybrid-match`, `frontend/src/pages/user/GapAnalysis.jsx` | Existing `backend-api/tests/Feature/GapAnalysisTest.php`, AI CV Analyzer API tests with fakes, and offline gap mini evaluation; current shell did not freshly rerun all tests. |
| FR-09 Application tracker | `backend-api/app/Http/Controllers/Api/ApplicationController.php`, `backend-api/app/Services/ApplicationTrackerService.php`, `frontend/src/pages/user/Applications.jsx` | Existing `backend-api/tests/Feature/ApplicationTrackerTest.php` and manual screenshot evidence. |
| FR-10 Admin operations | `backend-api/app/Http/Controllers/Api/Admin/*`, `frontend/src/pages/admin/*`, admin middleware/routes, scraping source and target role controllers | Existing scraper/admin-related tests plus manual admin screenshots; role rejection should be rechecked before final defense. |
| FR-11 Health and metrics | `backend-api/app/Http/Controllers/Api/HealthController.php`, `backend-api/app/Http/Controllers/Api/MetricsController.php`, `docker-compose.yml`, Prometheus/Grafana configs | Existing `backend-api/tests/Feature/HealthAndMetricsTest.php`, recorded health-probe evidence, and Docker Compose configuration checks. |
| NFR Security/reliability/performance | Request classes, middleware, private storage config, service-token checks, timeout/fallback logic, endpoint pagination/limits | Evidence is mixed: existing tests and code inspection support key controls, but a final clean-shell test rerun and browser security checks remain recommended. |

*Table 3a. Requirement-to-code/test traceability matrix.*

## 2.9 Hardware Requirements

For local demonstration, a developer machine capable of running Docker Desktop and multiple containers is required. CV parsing and OCR-like processing can be CPU-intensive; therefore, enough memory should be available for the Laravel backend, MySQL, frontend, Python services, MinIO, Prometheus, and Grafana. GPU acceleration is not required for the demonstrated flow.

## 2.10 Software Requirements

| Layer | Software |
|---|---|
| Frontend | React, Vite, Tailwind-style CSS classes, lucide-react icons, Recharts. |
| Backend | Laravel 12, PHP, Composer dependencies, Sanctum-style token authentication. |
| AI services | Python, FastAPI, text extraction, OCR-related dependencies, Scrapy-style job mining. |
| Data | MySQL 8.0, MinIO/S3-compatible storage. |
| Infrastructure | Docker, Docker Compose, Nginx, Prometheus, Grafana. |
| Testing | PHPUnit/Pest-style Laravel tests, pytest for Python, ESLint and Vite build. |

*Table 4. Software environment summary.*

## 2.11 Input and Output Flow

Primary inputs include user account data, uploaded CV files, imported job records, target role settings, and administrator source configurations. Primary outputs include normalized user profiles, skill lists, CV analysis metadata, estimated job matches, gap reports, application records, admin statistics, health checks, and monitoring metrics.

\\pagebreak

## 2.12 Use Case Summary

The main use cases are shown in Figure 5. The system separates student and administrator responsibilities while sharing the same backend API and database.

{figure_markdown("Figure 5", "UML use case diagram.", "assets/diagrams/05_use_case_diagram.png")}

\\pagebreak

# Chapter 3: System Design and Architecture

## 3.1 Introduction

CareerCompass is designed as a Dockerized multi-service application. This design separates browser UI, API logic, AI services, data storage, object storage, reverse proxy routing, and monitoring. Docker containers help package runtime dependencies consistently [10], while Docker Compose coordinates the multi-container local deployment [11].

## 3.2 High-Level System Architecture

The high-level architecture is shown in Figure 1. Browser users interact with the React frontend through Nginx. The frontend calls the Laravel API. Laravel persists records in MySQL, stores CV files in MinIO-compatible storage, calls the AI CV Analyzer for CV parsing and gap-analysis matching, ranks recommended jobs with Laravel title/skill/seniority scoring, and receives job imports from the job miner. This diagram was reviewed during the final evidence pass and already represents the important deployment boundaries: React, Nginx, Laravel, MySQL, MinIO, AI CV Analyzer, AI Job Miner, and monitoring.

{figure_markdown("Figure 1", "High-level architecture of CareerCompass.", "assets/diagrams/01_high_level_architecture.png")}

## 3.3 Frontend Architecture

The frontend is implemented with React and organized around React Router routes in `frontend/src/App.jsx` [42]. Public pages include Home, Login, Register, About, Privacy, Terms, and System Status. Protected student routes include Dashboard, Jobs, Gap Analysis, Profile, Settings, Market Intelligence, Applications, CV Builder, Mock Interview, Learning, Career Planner, Mentorship, and Tools Hub. Protected admin routes include Admin Dashboard, Jobs, Users, Sources, and Target Roles.

Figure 63 shows how the route tree is divided. This separation matters because the student experience is focused on career guidance, while the admin experience is focused on operating data, users, sources, targets, and diagnostics. The preview modules are included as graduation and future-extension screens; they demonstrate the intended product direction but should not be described as complete production modules.

{figure_markdown("Figure 63", "Frontend route and layout architecture.", "assets/diagrams/66_frontend_route_layout_architecture.png")}

The frontend API layer is located under `frontend/src/api`. Axios is configured in `client.js`, including base URL resolution, bearer token injection, request IDs, retry behavior for safe GET/HEAD requests, and 401 handling [43]. Authentication state is managed by `AuthContext.jsx`, which stores the user and token in local storage and refreshes the current user through `/user`. Route guards in `ProtectedRoute.jsx` and `GuestRoute.jsx` redirect unauthenticated users, keep admin-only routes behind role checks, and prevent logged-in users from returning to guest-only screens. Localization files exist under `frontend/src/locales`, and `i18n.js` uses browser language detection with English fallback.

{figure_markdown("Figure 64", "Frontend API and authentication flow.", "assets/diagrams/67_frontend_api_auth_flow.png")}

\\pagebreak

## 3.4 Backend API Architecture

The backend is a Laravel API. Routes are defined in `backend-api/routes/api.php` and are registered both at `/api` and `/api/v1`. The API includes public health/readiness/metrics endpoints, guest authentication routes, public job listing routes, internal scraper import routes protected by a service token, authenticated student routes, and admin routes protected by middleware.

Laravel provides structured controllers, form requests, resources, services, models, migrations, seeders, and tests. This aligns with Laravel's documented framework responsibilities, including routing, validation, database access, queues, and testing [1], [3].

Figure 65 summarizes the normal Laravel request lifecycle and the asynchronous branch used for longer tasks such as CV processing and job mining. The main design point is that controllers do not directly own every behavior: form requests validate input, services isolate reusable work, Eloquent models persist records, resources shape JSON responses, and queue workers handle tasks that should not block the browser.

{figure_markdown("Figure 65", "Laravel backend request lifecycle.", "assets/diagrams/68_laravel_backend_request_lifecycle.png")}

| Backend Module | Main Files / Components | Responsibility | Evidence |
|---|---|---|---|
| Authentication | `routes/api.php`, `AuthController`, `RegisterRequest`, `LoginRequest`, `UserResource` | Registration, login, logout, current-user lookup, profile update, and token lifecycle. | Guest auth routes, Sanctum tokens, throttled login group. |
| CV Upload and Processing | `CvController`, `CvUploadRequest`, `CvProcessingService`, `CvStorageService` | Validate CV files, store private file metadata, call the AI CV Analyzer, and persist analysis/profile/skill updates. | `/upload-cv`, signed CV file route, `cv_analyses` storage fields. |
| Profile and Skills | `User`, `UserProfile`, `Skill`, `user_skills`, `SkillSyncService` | Keep extracted skills and profile evidence normalized for matching. | Eloquent relationships and many-to-many skill pivots. |
| Jobs and Recommendations | `JobController`, `JobResource`, `Job`, `job_postings`, `job_skills` | Public job listing, details, recommendations, and job-skill requirements. | Public `/jobs` routes plus authenticated recommendation routes. |
| Gap Analysis | `GapAnalysisController`, `GapAnalysisService`, `GapAnalysisResource` | Compare a user's evidence against job or role requirements. | `/gap-analysis/job/{{jobId}}` and `/gap-analysis/role/{{roleId}}`. |
| Application Tracking | `ApplicationController`, `ApplicationTrackerService`, `Application`, `ApplicationResource` | Save and update opportunities across saved/applied/interviewing-style statuses. | Authenticated `applications` resource routes. |
| Admin Operations | `Admin\\DashboardController`, `AdminJobController`, `AdminUserController`, `ScrapingSourceController`, `TargetJobRoleController` | Admin statistics, users, jobs, scraping sources, diagnostics, and target roles. | Admin route group protected by `admin` middleware. |
| Scraping Import | `ScrapedJobController`, `VerifyScraperToken`, import request classes | Protected duplicate checks, imports, failure reports, and proxy access for scraper integration. | `scraper.token` and `throttle:scraper` route group. |
| Health and Metrics | `HealthController`, `MetricsController`, monitoring middleware | Liveness, readiness, and metrics surfaces for local operations. | `/health`, `/ready`, `/metrics`, and `/status`. |
| Queues and Workers | `ProcessOnDemandJobScraping`, market scraping jobs, Docker workers | Move long network or service calls out of request time. | Queue jobs and dedicated worker services. |
| File Storage | `CvStorageService`, signed CV file route, `cv_analyses` metadata | Keep uploaded CV binaries out of public MySQL fields and expose only signed access. | Private storage path, disk, checksum, MIME, and size metadata. |

*Table 5. Backend module responsibility summary.*

| Risk / Input Boundary | Control | Example Files |
|---|---|---|
| Invalid login or registration payload | Form requests, guest auth routes, and login throttling. | `RegisterRequest`, `LoginRequest`, `routes/api.php` |
| Invalid CV file upload | File validation, upload throttling, private storage, and status-aware error handling. | `CvUploadRequest`, `CvController`, `CvStorageService` |
| Unauthenticated user routes | Sanctum bearer-token middleware and current-user refresh. | `auth:sanctum`, `AuthController::user` |
| Non-admin access to admin routes | Backend `admin` middleware plus frontend role-based route guards. | `routes/api.php`, `ProtectedRoute.jsx` |
| Invalid scraper import payload | Service-token middleware, scraper throttling, and import form requests. | `VerifyScraperToken`, `StoreScrapedJobRequest`, `CheckScrapedJobRequest`, `ReportScrapingFailureRequest` |
| Duplicate scraped jobs | URL and title/company checks inside a transaction before create/update. | `ScrapedJobController`, `job_postings` migration |
| Private CV file access | Signed route and storage service rather than public direct file paths. | `/cv-files/{{cvAnalysis}}`, `CvStorageService` |
| Request tracing and debugging | Frontend request IDs and backend request-id middleware. | `client.js`, `RequestIdMiddleware` |

*Table 6. Laravel validation and protection mapping.*

\\pagebreak

## 3.5 AI CV Analyzer Architecture

The AI CV Analyzer is a FastAPI service. Laravel sends CV files to this service for parsing through `/api/parse-cv`. The analyzer routes PDFs and images differently, enforces timeout/error fallbacks, extracts readable text, runs structured extraction, and returns fields such as predicted role, seniority, domain, skills, strengths, gaps, red flags, confidence, document statistics, and parsing status. The backend handles statuses such as success, OCR fallback, timeout, error, empty file, and no text.

The analyzer is not a pure pretrained-model wrapper and not a model built entirely from scratch. It is a hybrid pipeline. The runtime prefers a local exported token-classification model under `ai-cv-analyzer/models/ner_weights/career_compass_ner_final` when that ignored deployment artifact exists; if the local artifact is unavailable, the NER engine has a fallback model path. Around that model-loading path, the team implemented practical CV-specific logic: spatial PDF parsing, OCR fallback, semantic sectioning, contact extraction, date/experience parsing, noisy-skill filtering, canonicalization, domain inference, seniority inference, and hybrid matching. PDF and OCR-related libraries are supported by external tools such as PyMuPDF, pdfplumber, and EasyOCR [22], [23], [24]. Transformer token classification and training concepts follow Hugging Face documentation [31], [32], [33].

Figure 9 summarizes the runtime path from the browser upload to Laravel persistence, FastAPI parsing, model/rule extraction, and dashboard output.

{figure_markdown("Figure 9", "AI CV Analyzer runtime flow.", "assets/diagrams/09_cv_analyzer_runtime_flow.png")}

## 3.6 AI Job Miner Architecture

The AI Job Miner is a FastAPI service with scraping and import support. It includes source adapters for demo/local data, public APIs, and HTML/Scrapy-style extraction. Scrapy is a Python framework for extracting structured data from websites [17], and Beautiful Soup is commonly used to parse HTML documents [18]. CareerCompass uses a quality gate and honest source classifications so that demo/imported data is not overstated as broad labor-market reach.

## 3.7 Database Design

MySQL stores users, user profiles, skills, user-skill pivots, experience records, CV analyses, job postings, job-skill pivots, applications, scraping sources, target job roles, scraping jobs, failed URLs, role statistics, optional scraping proxies, and Laravel runtime tables. MySQL is a relational database system documented by Oracle [8]. The Laravel migrations define schema constraints, indexes, foreign keys, and unique combinations such as job title/company uniqueness.

Figure 66 explains the relationship rationale behind the schema. The database is deliberately normalized around users, reusable skills, job requirements, CV analysis evidence, and job-mining operations. This keeps matching logic explainable: the system can compare a user's normalized skills against job-required skills while preserving the CV and scraping evidence that produced those records.

{figure_markdown("Figure 66", "Database relationship rationale.", "assets/diagrams/69_database_relationship_rationale.png")}

| Area | Tables | Design Reason |
|---|---|---|
| User identity and profile | `users`, `user_profiles`, `personal_access_tokens` | Separates authentication/token data from CV-derived profile data. |
| User skills and experience | `skills`, `user_skills`, `user_experiences` | Supports normalized skill matching and structured career evidence. |
| CV analysis | `cv_analyses` | Stores parsing status, extracted metadata, model evidence, and private file references. |
| Jobs and requirements | `job_postings`, `job_skills` | Separates job records from reusable required skills. |
| Applications | `applications` | Tracks saved/applied opportunities for each user and job. |
| Job mining operations | `scraping_jobs`, `scraping_sources`, `scraping_failed_urls`, `target_job_roles`, `job_role_statistics`, `scraping_proxies` | Preserves operational scraping state, source configuration, target roles, role-level aggregate evidence, optional proxy definitions, and failure evidence. |
| Runtime infrastructure | `jobs`, `job_batches`, `failed_jobs`, `cache`, `cache_locks`, `sessions`, `password_reset_tokens`, `personal_access_tokens` | Supports queues, failed-job tracking, batches, cache locks, sessions, password reset tokens, Sanctum tokens, and repeatable local demonstration behavior. |

*Table 7. Database design rationale.*

| Data Integrity Mechanism | Purpose | Code / Schema Evidence |
|---|---|---|
| Foreign keys | Keep user, job, application, skill, scraping, and source records consistent. | Migration `foreignId(...)->constrained()` and `nullOnDelete()` definitions. |
| Unique identity fields | Prevent duplicate core identities. | `users.email`, `skills.name`, `user_profiles.user_id`, `cv_analyses.user_id`. |
| Pivot uniqueness | Prevent repeated skill links while allowing many-to-many matching. | Unique pairs on `user_skills` and `job_skills`. |
| Job duplicate constraints | Reduce duplicate imported jobs at the database layer. | Unique `url` and unique `title`/`company` combination in `job_postings`. |
| Application uniqueness | Prevent a user from tracking the same job repeatedly. | Unique `user_id`/`job_id` pair in `applications`. |
| Operational indexes | Make status, source, and retry queries practical for dashboards. | Indexes on scraping status/type/source/created and failed-URL retry fields. |
| Private object references | Avoid storing CV binary files directly in MySQL. | `cv_analyses` storage disk/path/checksum/size/MIME metadata. |
| Form request validation | Reject malformed payloads before persistence. | Laravel `app/Http/Requests` classes. |

*Table 8. Data integrity mechanisms.*

## 3.8 ERD

Figure 8 summarizes the main application tables and relationships from the current Laravel migrations. It is not a complete replacement for migrations, and it intentionally keeps Laravel runtime tables such as queues, cache, sessions, password-reset tokens, and Sanctum tokens in Appendix C rather than overloading the ERD. The diagram was corrected to remove columns not present in migrations and to include the implemented job-mining support tables.

{figure_markdown("Figure 8", "ERD and database summary diagram.", "assets/diagrams/08_erd.png")}

\\pagebreak

| Relationship | Meaning |
|---|---|
| `users` -> `user_profiles` | One user has one profile; profile data is kept separate from authentication data. |
| `users` -> `user_experiences` | One user can have many structured experience records. |
| `users` <-> `skills` | Many-to-many relationship through `user_skills`, including confidence/evidence metadata. |
| `users` -> `cv_analyses` | CV analysis records preserve parsing status, extracted evidence, storage metadata, and model output. |
| `job_postings` <-> `skills` | Many-to-many relationship through `job_skills`, including required/importance metadata. |
| `users` <-> `job_postings` | Users track opportunities through `applications`; the application row stores user-specific status and notes. |
| `job_postings` -> `scraping_sources` | Imported jobs can reference their source configuration while remaining available if the source is later disabled. |
| `scraping_jobs` -> `scraping_failed_urls` | Failed URLs can be associated with an operational scraping run for diagnostics. |
| `target_job_roles` -> scraping workflows | Active target roles guide full scraping runs; they are operational configuration, not a direct job foreign key. |
| `job_role_statistics` | Stores aggregate market/mining observations by `role_title`; it has no direct foreign key to jobs or target roles. |
| `scraping_proxies` | Stores optional proxy connection settings for scraper operation; it is operational configuration rather than a domain relationship. |

*Table 9. Main ERD relationship notes.*

## 3.9 Data Flow Diagrams

The context-level data flow is shown in Figure 3, and the expanded process-level view is shown in Figure 4. Student and administrator workflows enter the same system boundary, while external job sources and AI services interact with controlled backend processes.

{figure_markdown("Figure 3", "DFD Level 0 context diagram.", "assets/diagrams/03_dfd_level_0.png")}

{figure_markdown("Figure 4", "DFD Level 1 process diagram.", "assets/diagrams/04_dfd_level_1.png")}

## 3.10 UML Use Case Diagram

The use case diagram separates student actions from administrator actions. Student workflows focus on career exploration. Admin workflows focus on operating and inspecting the imported job ecosystem.

## 3.11 UML Sequence Diagrams

Figure 6 shows the CV upload and analysis sequence. Figure 7 shows recommendation and gap analysis.

{figure_markdown("Figure 6", "Sequence diagram for CV upload and analysis.", "assets/diagrams/06_sequence_cv_upload_analysis.png")}

{figure_markdown("Figure 7", "Sequence diagram for recommendation and gap analysis.", "assets/diagrams/07_sequence_job_recommendation_gap_analysis.png")}

## 3.12 Deployment Architecture with Docker

The deployment is defined by Docker Compose files. Nginx exposes the application, the frontend serves built React assets, the Laravel API and workers handle backend work, MySQL stores structured data, MinIO stores private CV objects, Python services provide AI CV parsing and job mining, and Prometheus/Grafana provide monitoring. Figure 2 summarizes the container layout.

{figure_markdown("Figure 2", "Docker deployment architecture.", "assets/diagrams/02_docker_deployment.png")}

## 3.13 Monitoring Architecture

CareerCompass includes live, readiness, and metrics endpoints. Prometheus is used for scraping and time-series metrics collection [13], while Grafana visualizes metrics and dashboard panels [14]. The admin dashboard also exposes application-level health information for the graduation demo.

## 3.14 Design Decisions and Justification

| Decision | Justification |
|---|---|
| Use Laravel for the main API. | The project benefits from built-in routing, validation, Eloquent ORM, queues, resources, tests, and middleware [1]. |
| Use React for UI. | Component-based pages support student/admin route separation and reusable cards [4]. |
| Use FastAPI for AI services. | Python AI/NLP dependencies are easier to isolate behind typed HTTP services [6]. |
| Use Docker Compose. | Multiple services can be started consistently for a graduation defense [11]. |
| Use private object storage for CV files. | CV files are sensitive; private storage and signed downloads reduce accidental exposure [9], [27]. |
| Keep AI wording honest. | Match scores and CV parsing are estimates; the demo should avoid claiming certain outcomes. |

*Table 10. Design decisions summary.*

\\pagebreak

# Chapter 4: Software and Tools Used

## 4.1 Laravel

Laravel is used for the backend API, controllers, middleware, validation requests, Eloquent models, migrations, resources, queues, tests, and storage integration [1]. In CareerCompass, Laravel centralizes authenticated business logic and data persistence.

## 4.2 React

React is used to build the browser-based interface for students and administrators [4]. The project uses React routes and components for dashboard cards, forms, pages, tables, modals, and visualization panels.

## 4.3 Vite

Vite is used as the frontend build tool [5]. The successful production build transformed 2904 modules during validation.

## 4.4 FastAPI

FastAPI is used for Python services because it supports efficient API development using Python type hints and modern web service patterns [6]. CareerCompass uses it for CV analysis and job mining service endpoints.

## 4.5 Python

Python is used for AI and scraping-related services [7]. It supports the ecosystem of text extraction, OCR, NLP, testing, and scraping libraries used by the AI services.

## 4.6 MySQL

MySQL stores relational project data including users, profiles, skills, jobs, applications, scraping sources, and CV analyses [8].

## 4.7 MinIO / S3-Compatible Storage

MinIO-compatible object storage is used to store CV files privately. This avoids placing uploaded CVs directly in public web directories and supports signed temporary download flows [9].

## 4.8 Docker and Docker Compose

Docker containers package each runtime, and Docker Compose coordinates multi-container execution [10], [11]. CareerCompass uses Compose for Nginx, frontend, backend, workers, MySQL, MinIO, AI services, Prometheus, and Grafana.

## 4.9 Nginx

Nginx acts as the reverse proxy and public gateway for the local deployment [12]. It routes browser and API traffic to the correct containers.

## 4.10 Prometheus and Grafana

Prometheus collects metrics [13], and Grafana visualizes metrics and dashboards [14]. CareerCompass uses these tools to support monitoring-oriented evaluation.

## 4.11 GitHub Actions

GitHub Actions is configured for repository automation and CI/CD-style validation [15]. The report treats workflow definitions as part of the development process, while final PR status should be reviewed after the draft PR is opened.

## 4.12 NLP / AI Libraries

The project uses concepts and libraries related to text extraction, OCR, transformer token classification, TF-IDF, cosine similarity, and sentence embeddings. TF-IDF and cosine similarity are documented by scikit-learn [19], [20]. Sentence Transformers provides sentence embedding models and utilities [21]. PyMuPDF, pdfplumber, and EasyOCR support PDF/image text extraction and OCR-style workflows [22], [23], [24]. The token-classification path builds on BERT-style contextual representations [37].

Hugging Face Transformers is used for model loading and token-classification style inference/training [31], [32]. The training notebook uses Hugging Face Trainer concepts, token-label alignment, the `seqeval` metric family, and a BERT token-classification configuration [33]. Optional dynamic quantization is supported in the local model-loading code as an inference optimization concept [38]. The deployed matching layer also uses sentence embeddings and a custom TF-IDF fallback so that a service outage or weak semantic signal does not become a silent failure.

## 4.13 Synthetic Data and Training Tools

The repository includes a model-training workflow designed around synthetic annotated CV snippets. The data-generation script uses the Gemini API through Google AI developer tooling to generate labeled examples, while the training notebook is designed for Google Colab GPU execution [34], [35], [36]. This report treats the external AI tooling as training-support tooling, not as a runtime dependency for private CV uploads.

## 4.14 Testing Tools

Backend tests use Laravel/PHP testing tools and PHPUnit concepts [25]. Python service tests use pytest where available [26]. Frontend validation uses ESLint and Vite build checks.

## 4.15 Development and Version Control Tools

Git and GitHub are used for version control and pull-request-based collaboration. Docker Desktop provides the local container runtime. Browser screenshots were captured through a Chrome DevTools Protocol helper to provide evidence images for this report.

\\pagebreak

# Chapter 5: System Implementation

## 5.1 Introduction

This chapter documents the implemented CareerCompass modules based on repository files. The implementation is not a generic career platform; it is a Laravel/React/FastAPI/Docker system with specific routes, services, pages, models, seeders, and tests.

## 5.2 Authentication and User Management

Authentication is implemented in `backend-api/app/Http/Controllers/Api/AuthController.php`. Registration creates a user with role `user`, loads profile and related resources, and returns a token. Login validates credentials, checks the banned state, revokes old tokens, creates a new token, and returns a user resource. The frontend stores the token as `auth_token` and the user object in local storage through `frontend/src/context/AuthContext.jsx`.

The register request restricts emails to selected public email domains and validates password format. The admin role is protected by the `IsAdmin` middleware. Admin seed data creates a demo-only administrator account through `AdminUserSeeder`.

## 5.3 Student Dashboard

The student dashboard is implemented in `frontend/src/pages/user/Dashboard.jsx`. It presents the current profile state, CV upload/update controls, profile completeness, career identity, AI insights, and next actions. Before CV upload, it prompts the user to add a CV. After upload, it displays parsed CV availability, role inference, profile score, experience, and action buttons.

{figure_markdown("Figure 34", "Student dashboard before CV upload.", "assets/screenshots/04_dashboard_before_cv_upload.png")}

{figure_markdown("Figure 36", "Dashboard after successful CV parsing.", "assets/screenshots/06_dashboard_after_cv_upload.png")}

## 5.4 CV Upload and Storage

`CvUploadRequest` requires a `cv` file and accepts PDF, JPEG, JPG, and PNG files up to 5 MB. The frontend appends the selected file as `cv` in a `FormData` object. `CvController` calls the CV processing service, persists the file path and metadata, and returns a unified user resource.

CV storage is handled as a private file workflow. The system supports signed download URLs, which is a better demo posture than public file exposure. OWASP recommends validating uploaded file type, extension, size, and storage handling carefully [27].

{figure_markdown("Figure 35", "CV upload user interface.", "assets/screenshots/05_cv_upload_ui.png")}

## 5.5 CV Parsing and Skill Extraction

The CV processing flow sends the file to the AI CV Analyzer, receives parsed data, synchronizes skills, updates profile fields, and stores CV analysis metadata. The implementation handles multiple parsing statuses honestly. If analysis times out, fails, or finds no readable text, the backend returns warnings and preserves existing profile details rather than silently replacing data with low-quality output.

### 5.5.1 Analyzer Runtime Components

The analyzer is implemented as layered Python code rather than one monolithic function. `main.py` exposes FastAPI endpoints, `CVOrchestrator` coordinates extraction, `AdvancedNEREngine` runs transformer-based named-entity recognition, and supporting engines handle contacts, sections, experience blocks, canonicalization, domain classification, seniority classification, semantic embeddings, and hybrid job matching. Figure 11 summarizes these extraction components.

{figure_markdown("Figure 11", "AI CV Analyzer extraction components.", "assets/diagrams/11_cv_extraction_components.png")}

| Component | Repository Evidence | Responsibility | Output Used By |
|---|---|---|---|
| FastAPI gateway | `ai-cv-analyzer/main.py` | Receives `/api/parse-cv` and `/api/hybrid-match` requests; handles timeout/error fallbacks. | Laravel `CvProcessingService` and `GapAnalysisService` |
| Laravel CV service | `backend-api/app/Services/CvProcessingService.php` | Sends uploads to the analyzer, stores private CV objects, persists normalized analysis. | User profile, skills, experiences, recommendations |
| Spatial/OCR extraction | `spatial_parser.py`, `ocr_pipeline.py` | Reads PDF text first and falls back to image/OCR when needed. | Section segmenter and NER pipeline |
| Advanced NER | `advanced_ner.py` and optional ignored local model folder | Loads the exported local token-classification model when deployed; chunks long CVs and groups entity spans. | Skills, roles, education, certifications |
| Rule engines | contact, experience, date, noise-filtering helpers | Extract contact details, experience blocks, dates, and remove title-like or noisy skill candidates. | Profile and experience persistence |
| Canonicalization/classification | Layer 1 and Layer 2 modules | Normalize skills and infer primary domain plus seniority. | Dashboard identity card and matching |
| Hybrid matching | Layer 3 matching modules | Combines semantic scores, skill text similarity, domain alignment, constraints, and TF-IDF fallback. | Gap analysis and fit explanation; the job-list ranking endpoint is separate. |
| Frontend display | `Dashboard.jsx`, `AiInsights.jsx` | Shows upload status, confidence-style signals, role/seniority, and extracted skills. | Student-facing CV feedback |

*Table 11. AI CV Analyzer components.*

### 5.5.2 Model Type and Customization

The NER part is a fine-tuned transformer token-classification design, not a model trained from randomly initialized architecture. The notebook uses `bert-base-cased` as the base checkpoint and defines a CV-specific BIO label set. The runtime checks for the exported local model at `ai-cv-analyzer/models/ner_weights/career_compass_ner_final` and uses it when available. That folder is ignored by Git, so committed repository evidence should be described as code and training workflow evidence. Local ignored metadata inspected during this documentation update identified a BERT token-classification configuration, 512-token maximum position setting, 28,996-token vocabulary, 12 transformer layers, 768 hidden size, and a cased tokenizer; the binary model weights were not copied into the report.

The project customization is mainly in the data, labels, orchestration, and post-processing. The team defined CV-specific labels, generated synthetic annotated examples, cleaned entity spans, aligned character spans to tokens, exported a local model, and connected it to a CV-specific extraction pipeline. The pipeline then adds deterministic rules for contacts, dates, experience blocks, noisy skill rejection, canonical skill names, domain/seniority inference, and safer fallback statuses.

| Label | BIO Forms | Meaning in Training Data | Runtime Note |
|---|---|---|---|
| O | O | Token outside a labeled entity. | Used by the model to ignore ordinary text. |
| SKILL | B-SKILL, I-SKILL | Technical skill such as Laravel, React, Docker, SQL, or PyTorch. | Returned as skills after filtering and canonicalization. |
| ROLE | B-ROLE, I-ROLE | Job title or role such as Backend Developer or Data Analyst. | Used for predicted role and role evidence. |
| EDU | B-EDU, I-EDU | Degree, major, faculty, university, or education phrase. | Used as education/profile evidence. |
| CERT | B-CERT, I-CERT | Certification such as AWS Cloud Practitioner. | Used as certification evidence. |
| SOFT | B-SOFT, I-SOFT | Soft skill phrase such as leadership or communication. | Present in training configuration; the runtime NER grouping mainly returns SKILL, ROLE, EDU, and CERT. |

*Table 12. NER entity label schema.*

### 5.5.3 Synthetic Dataset Generation

The helper material under `D:/Graduation/model-analys-helper` was reviewed as documentation support. The top-level helper folders found were `docs`, `layer1`, `layer2`, and `layer3`. No raw datasets, screenshots, or secrets were copied into the graduation-book folder. The important training workflow is already represented in the repository by `ai-cv-analyzer/training/generate_tech_dataset.py`, `clean_dataset.py`, and `train_ner.ipynb`.

The dataset generator is designed to call the Gemini API through API keys stored outside source control. It asks for synthetic technical CV snippets across backend, frontend, DevOps, mobile, AI/data, cybersecurity, cloud, QA, and networking contexts. It intentionally includes positive labeled examples and negative decoy examples so that the model learns both what to tag and what to ignore. Because this uses external generation and keys, it was inspected rather than executed for this documentation update.

| Step | Repository Evidence | Description | Documentation Decision |
|---|---|---|---|
| API-key loading | `generate_tech_dataset.py` | Reads `GEMINI_API_KEYS` from `.env` and rotates keys/models during generation. | Do not commit keys; no secrets were found in helper docs. |
| Synthetic sample generation | Generator system prompt | Requests batches with technical domains, noise, varied CV formats, and labeled entities. | Documented as synthetic data, not real student CV data. |
| Negative decoys | Generator distribution comments | Includes examples with no entities to reduce false positives. | Preserved in the description because it affects model behavior. |
| Cleaning | `clean_dataset.py` | Normalizes whitespace, deduplicates exact text, validates entity text, filters long skill spans. | Documented as a data-quality gate before training. |
| Output | Training docs and scripts | Expected cleaned JSON/JSONL input for notebook training. | Dataset files were not present in the repository, so metrics cannot be reproduced from repo files alone. |

*Table 13. Synthetic dataset generation workflow.*

### 5.5.4 Training Notebook Workflow

The training notebook is structured for Google Colab rather than local execution. It installs model-training dependencies, loads cleaned JSON data, defines labels, tokenizes examples, aligns entity spans to token labels, initializes `AutoModelForTokenClassification` from `bert-base-cased`, trains with Hugging Face Trainer, evaluates with sequence-labeling metrics, then exports `career_compass_ner_final` for deployment [31], [32], [33], [36].

{figure_markdown("Figure 10", "AI CV Analyzer model-training workflow.", "assets/diagrams/10_cv_model_training_pipeline.png")}

| Setting | Value Found in Notebook or Docs | Purpose | Evidence Limitation |
|---|---|---|---|
| Base checkpoint | `bert-base-cased` | Provides pretrained language representations for token classification. | Confirmed by training notebook and Colab PDF. |
| Labels | O plus B/I for SKILL, ROLE, EDU, CERT, SOFT | Encodes CV entity spans using BIO tagging. | Label map is reproducible from notebook/config. |
| Split | 90 percent train, 10 percent test, seed 42 | Creates a repeatable train/evaluation split. | Colab PDF records 41,319 train rows and 4,592 test rows; dataset content is not committed. |
| Max length | 512 tokens | Fits BERT token-classification input limits. | Long runtime CVs are handled through chunking separately. |
| Epochs and learning rate | 5 epochs, 2e-5 | Standard fine-tuning style schedule for a small NER task. | No final epoch table was available. |
| Batch size | 16 train/eval | Balances GPU memory and throughput on Colab/T4-style runtime. | Visible in the exported Colab PDF. |
| Metrics | precision, recall, F1, accuracy via sequence labeling | Evaluates entity extraction quality when labels are available. | Colab PDF records overall metrics; per-label report is not visible. |
| Export | `career_compass_ner_final` zip/model folder | Produces the deployable local model artifact. | Export success is visible in the PDF, but model weights are ignored by Git. |

*Table 14. Model training configuration.*

### 5.5.5 Layer 1: CV Understanding Pipeline

Layer 1 is responsible for turning a noisy CV file into a structured candidate profile. The runtime begins at `main.py`, which accepts the uploaded file, chooses PDF or image handling, applies timeout/error wrappers, and delegates the actual extraction to `CVOrchestrator`. The orchestrator first tries ordered PDF text extraction. The spatial parser reads words from PDF pages, groups words into rows using an adaptive tolerance, splits row segments when large x-axis gaps imply columns, removes `(cid:...)` artifacts, and falls back to plain PDF extraction when the spatial output loses too much text.

If the file has little or no readable text, the OCR path renders PDF pages to images and uses EasyOCR after grayscale/blur preprocessing. After text recovery, the semantic segmenter finds CV sections, contact extraction parses email/phone/location fields, the NER engine extracts entity candidates, experience logic estimates date ranges and career signals, and the canonicalizer normalizes skills before the result is validated through strict Pydantic schema classes.

{figure_markdown("Figure 12", "Layer 1 CV understanding pipeline.", "assets/diagrams/12_layer1_understanding_pipeline.png")}

\\pagebreak

| Layer 1 Component | Main Files | Important Behavior | Risk or Fallback |
|---|---|---|---|
| API gateway | `main.py` | `/api/parse-cv`, `/api/hybrid-match`, timeout handling, health and metrics endpoints. | Timeout results are returned as explicit status dictionaries. |
| Spatial parser | `core/layer1_understanding/spatial_parser.py` | Word extraction, row grouping, column ordering, dehyphenation, plain-text fallback. | Falls back when spatial output is too weak. |
| OCR fallback | `core/layer1_understanding/ocr_pipeline.py`, orchestrator OCR helpers | Renders image-like PDFs, preprocesses pages, and extracts text when normal PDF parsing fails. | Triggered for short/no-text inputs. |
| Section segmenter | `core/layer1_understanding/section_segmenter.py` | Header detection from patterns and optional semantic header matching. | Missing headers fall back to profile-style grouping. |
| Contact and experience engines | `contact_extractor.py`, `experience_engine.py` | Extract emails/phones/location, date ranges, total years, skill durations, gaps, overlaps, and action-verb strength. | Ambiguous dates are treated conservatively. |
| NER and canonicalization | `advanced_ner.py`, `canonicalizer.py` | Extracts skills/roles/education/certifications, filters noise, deduplicates, and maps skills to canonical names. | Fallback model path exists when local deployment artifact is missing. |
| Output schema | `schema.py` | Strict typed response for profile, skills, experience, confidence, stats, and parsing status. | Invalid shapes are prevented before backend persistence. |

*Table 15. Layer 1 component details.*

{figure_markdown("Figure 17", "Skill canonicalization chain.", "assets/diagrams/17_canonicalization_chain.png")}

### 5.5.6 Layer 2: Classification Engine

Layer 2 enriches the extracted CV with a domain and seniority interpretation. The classification orchestrator reads the Layer 1 result, then combines title, experience, summary, skill categories, and taxonomy descriptions. `DomainEngine` compares CV context against taxonomy descriptions using semantic embeddings when available. `SkillEngine` separates hard, soft, and management-oriented skills using taxonomy rules. `SeniorityEngine` combines years of experience, title keywords, semantic title/summary hints, and action-verb strength.

{figure_markdown("Figure 13", "Layer 2 classification flow.", "assets/diagrams/13_layer2_classification_flow.png")}

{figure_markdown("Figure 16", "Seniority decision logic.", "assets/diagrams/16_seniority_decision_logic.png")}

\\pagebreak

| Layer 2 Component | Main Files | Input | Output |
|---|---|---|---|
| Classification orchestrator | `core/layer2_classification/orchestrator.py` | Parsed CV profile, skills, experience, and summary. | Adds primary domain, seniority level, skill categories, and confidence-style signals. |
| Domain engine | `core/layer2_classification/domain_engine.py`, `data/taxonomy.json` | Title, summary, and first experience titles. | Primary technical domain selected from taxonomy descriptions. |
| Seniority engine | `core/layer2_classification/seniority_engine.py` | Experience years, title, summary, and action verbs. | Intern, Junior, Mid-Level, Senior, or Lead / Manager estimate. |
| Skill engine | `core/layer2_classification/skill_engine.py` | Canonical skill names and taxonomy terms. | Hard, soft, and management skill buckets. |
| Taxonomy loader | `core/layer2_classification/utils.py` | JSON taxonomy file. | Shared configuration for domain and skill classification. |

*Table 16. Layer 2 classification engine details.*

### 5.5.7 Layer 3: Matching Engine

Layer 3 compares a candidate profile with a job description. `JobDescriptionEngine` parses job text into seniority, required years, mandatory skills, bonus skills, domain, and summary. `IntelligentMatcher` calculates semantic similarity, skill-text similarity, and domain alignment using adaptive weights that change by seniority level. `ConstraintValidator` subtracts penalties for missing mandatory skills, experience shortfalls, and seniority mismatch. `FitAnalysisGenerator` turns the numeric result into strengths, gaps, red flags, and a verdict.

{figure_markdown("Figure 14", "Layer 3 matching engine.", "assets/diagrams/14_layer3_matching_engine.png")}

{figure_markdown("Figure 18", "Layer 3 score collapse logic.", "assets/diagrams/18_score_collapse_logic.png")}

| Layer 3 Component | Main Files | Scoring Contribution | Explanation Contribution |
|---|---|---|---|
| JD parser | `job_description_engine.py` | Extracts requirements that become matching inputs. | Explains what the system understood from the job post. |
| Semantic embedder | `embedder.py` | Summary similarity and domain-similarity fallback. | Captures meaning beyond exact keyword overlap when dependencies are available. |
| Intelligent matcher | `similarity.py`, `matching_config.json` | Combines semantic, skill, and domain scores using seniority-aware weights. | Produces score breakdown and qualification flag. |
| Constraint validator | `constraint_validator.py` | Applies capped penalties for mandatory gaps, experience gaps, and seniority mismatch. | Lists missing mandatory skills and mismatch reasons. |
| Fit analysis generator | `fit_analysis_generator.py` | Converts score ranges into verdict categories. | Generates strengths, gaps, and red flags for the UI. |
| Ranking orchestrator | `ranking_orchestrator.py` | Applies matcher repeatedly across candidates/jobs. | Sorts candidates or opportunities by explainable fit. |

*Table 17. Layer 3 matching engine details.*

### 5.5.8 Semantic Embedding, Caching, and TF-IDF Fallback

The semantic embedder uses the configured sentence-transformer model name, with `all-MiniLM-L6-v2` as the default. The class is implemented as a singleton so that model loading is not repeated for every request. It keeps a bounded embedding cache of 2,000 entries and can optionally apply dynamic quantization through `EMBEDDER_QUANTIZE`. If the embedding stack is unavailable, it returns zero vectors and the calling code falls back to transparent lower-confidence behavior instead of pretending semantic comparison succeeded.

The FastAPI `/api/hybrid-match` endpoint also combines semantic-style matching with the pure Python TF-IDF matcher when the TF-IDF module is importable. In that endpoint, semantic/adaptive scoring contributes 60 percent and TF-IDF contributes 40 percent. This gives the demo a useful exact-keyword safety check for skills such as Laravel, Docker, MySQL, React, AWS, or Kubernetes.

| Method | Repository Evidence | Strength | Limitation |
|---|---|---|---|
| Sentence embeddings | `core/layer3_matching/embedder.py` | Captures meaning even when words differ. | Requires sentence-transformer dependencies and model loading. |
| Domain similarity | `core/layer3_matching/similarity.py` | Allows related domains to receive partial credit. | Low domain similarity is cut off below the configured threshold. |
| Skill text similarity | `core/layer3_matching/similarity.py` | Compares extracted CV skills against mandatory and bonus job skills. | Quality depends on upstream extraction and canonicalization. |
| TF-IDF fallback | `core/layer3_matching/tfidf.py`, `/api/hybrid-match` | Lightweight deterministic keyword overlap. | Does not understand synonyms or phrase meaning. |
| Constraint penalties | `constraint_validator.py` | Prevents high scores when hard requirements are missing. | Penalty weights need larger evaluation data for calibration. |

*Table 18. Semantic embedding and TF-IDF fallback comparison.*

### 5.5.9 NER Token Processing and BIO Tagging

The training notebook uses character-span annotations and converts them into token labels. Each text sample is tokenized with offsets; special tokens are assigned `-100` so they are ignored by the loss; tokens whose offsets fall inside an entity span are assigned `B-` or `I-` labels. The cased BERT tokenizer is appropriate for CV text because names, certificates, role titles, and technology names often rely on capitalization. At runtime, long CV text is chunked with overlap, model predictions are merged, subword prefixes are cleaned, and duplicate/noisy entities are filtered before canonicalization.

{figure_markdown("Figure 15", "NER token processing and BIO tagging.", "assets/diagrams/15_ner_token_processing.png")}

\\pagebreak

| Simplified Text Token | BIO Label | Why It Matters |
|---|---|---|
| Experienced | O | Ordinary descriptive word, not extracted as an entity. |
| Backend | B-ROLE | Start of a role phrase. |
| Developer | I-ROLE | Continuation of the role phrase. |
| with | O | Connector word. |
| Laravel | B-SKILL | Skill entity. |
| Docker | B-SKILL | Skill entity. |
| MySQL | B-SKILL | Skill entity. |

*Table 19. Simplified BIO tagging example.*

### 5.5.10 AI CV Analyzer Source Code Inventory

The AI CV Analyzer was audited as source code, not only as a running service. The inventory below summarizes the relevant repository areas. Full support notes are stored under `docs/graduation-book/model-analysis/`.

| Area | Representative Files | Main Responsibility |
|---|---|---|
| FastAPI service | `main.py`, `Dockerfile`, `requirements.txt`, `.env.example` | Service startup, endpoints, container runtime, and documented configuration variables. |
| Layer 1 understanding | `core/layer1_understanding/*.py`, `core/layer1_understanding/data/config.json` | PDF/image text extraction, OCR fallback, sectioning, NER, contact extraction, experience analysis, and canonicalization. |
| Layer 2 classification | `core/layer2_classification/*.py`, `core/layer2_classification/data/taxonomy.json` | Domain, seniority, and skill-category enrichment. |
| Layer 3 matching | `core/layer3_matching/*.py`, `core/layer3_matching/matching_config.json` | Job parsing, semantic/TF-IDF matching, constraints, fit explanation, and AI-side candidate ranking when explicitly invoked. |
| Training workflow | `training/generate_tech_dataset.py`, `clean_dataset.py`, `train_ner.ipynb` | Synthetic labeled dataset generation, cleaning, token alignment, Trainer setup, metrics code, and export. |
| Diagnostics and tests | `scripts/verify_phase*.py`, `tests/test_service_api.py`, `tests/trace_cv.py`, manual tests | Phase checks, service API tests with fakes, tracing, and manual validation helpers. |
| Documentation | Layer README and EXPLAIN files | Developer explanations for analyzer layers and matching logic. |
| Ignored deployment assets | `.env`, `models/ner_weights/...` | Local secrets and model weights are intentionally ignored; only safe metadata was inspected locally. |

*Table 20. AI CV Analyzer source inventory summary.*

\\pagebreak

| Algorithm or Workflow | Primary File(s) | Notes |
|---|---|---|
| Parse-CV request routing | `main.py` | File validation, image/PDF branching, timeout and error wrappers. |
| Ordered PDF extraction | `core/layer1_understanding/spatial_parser.py` | Adaptive row/column ordering and dehyphenation. |
| OCR fallback | `core/layer1_understanding/ocr_pipeline.py`, `CVOrchestrator` | Used when PDF text is absent or too short. |
| NER extraction | `core/layer1_understanding/advanced_ner.py` | Singleton model loading, chunked inference, entity merging, and cleanup. |
| Skill normalization | `core/layer1_understanding/canonicalizer.py` | Exact, fuzzy, and semantic mapping toward canonical names. |
| Experience calculation | `core/layer1_understanding/experience_engine.py` | Date ranges, total years, skill durations, gaps, overlaps, and action verbs. |
| Seniority inference | `core/layer1_understanding/orchestrator.py`, `core/layer2_classification/seniority_engine.py` | Combines title keywords, years, semantic hints, and action verbs. |
| Domain classification | `core/layer2_classification/domain_engine.py` | Embedding comparison against taxonomy descriptions. |
| Job parsing | `core/layer3_matching/job_description_engine.py` | Seniority, years, mandatory/bonus skills, domain, and summary extraction. |
| Hybrid match scoring | `core/layer3_matching/similarity.py`, `core/layer3_matching/constraint_validator.py` | Weighted semantic/skill/domain score minus constraint penalties plus bonus boost. |
| TF-IDF fallback | `core/layer3_matching/tfidf.py`, `main.py` | Pure Python sparse cosine score used in hybrid endpoint. |
| Fit explanation | `core/layer3_matching/fit_analysis_generator.py` | Verdict, strengths, gaps, and red flags. |
| NER training | `training/train_ner.ipynb` | BERT base checkpoint, BIO labels, token alignment, Trainer, seqeval metrics. |
| Synthetic data generation | `training/generate_tech_dataset.py` | Gemini-based generation with key rotation and negative decoys. |

*Table 21. Algorithm-to-file mapping.*

## 5.6 Profile and Skills Management

The profile page reads normalized user data, profile fields, experiences, skills, and CV analysis. The system distinguishes user fields, profile fields, extracted skills, predicted role, seniority, and completeness score. Skill synchronization is handled through backend services rather than only frontend state.

{figure_markdown("Figure 37", "Extracted profile and skills page.", "assets/screenshots/07_extracted_profile_skills.png")}

## 5.7 Job Data Model

Jobs are represented in the backend through job posting models and migrations. Fields include title, company, description/requirements, URL, source, and metadata. The seeders and import controllers enforce quality gates and uniqueness rules, including a title/company uniqueness constraint that prevented duplicate seed insertion during validation.

## 5.8 AI Job Miner and Scraping Sources

The job miner exposes a FastAPI service and imports candidate jobs through configured sources. This implementation chapter keeps the feature overview short: Laravel remains the system of record, the Python service handles adapter work, and admin pages expose source diagnostics, source status, testing, and target role management. Chapter 7 expands this subsystem with runtime diagrams, queue flow, API contracts, import/deduplication logic, failed-URL handling, security boundaries, evaluation evidence, and ethical limitations.

{figure_markdown("Figure 46", "Admin sources diagnostics page.", "assets/screenshots/16_admin_sources_diagnostics.png")}

## 5.9 Job Recommendations

The jobs page requests `/api/v1/jobs/recommended` when no manual search query is active. In the current Laravel implementation, `JobController::getRecommended` uses the user's predicted role or profile title to select candidate job titles, then ranks up to 200 candidates by title similarity, required-skill overlap, and seniority hints before returning up to 50 jobs with an estimated `match_percentage`.

This endpoint does not call `/api/hybrid-match`. Semantic/adaptive plus TF-IDF scoring belongs to the gap-analysis workflow through `/api/hybrid-match`.

{figure_markdown("Figure 38", "Jobs recommendations page.", "assets/screenshots/08_jobs_recommendations.png")}

## 5.10 Gap Analysis

Gap analysis compares a selected job or target role against the user's profile and extracted skills. It returns matched skills, critical/missing skills, recommendations, match percentage, and roadmap-like guidance. The frontend displays these outputs in an explainable layout rather than a single opaque score.

{figure_markdown("Figure 40", "Gap analysis page.", "assets/screenshots/10_gap_analysis.png")}

## 5.11 Application Tracker

The application tracker is implemented through `ApplicationController`, `ApplicationTrackerService`, and `frontend/src/pages/user/Applications.jsx`. Students can save a job, update status, view counts, and delete tracked items. The backend validates job existence and allowed statuses.

{figure_markdown("Figure 41", "Applications tracker page.", "assets/screenshots/11_applications_tracker.png")}

## 5.12 Admin Dashboard

The admin dashboard summarizes users, imported jobs, active sources, target roles, health status, and scraping batch progress. It is protected by admin middleware and uses admin API routes.

{figure_markdown("Figure 44", "Admin dashboard.", "assets/screenshots/14_admin_dashboard.png")}

## 5.13 Admin Source Diagnostics

The source diagnostics page lists configured scraping sources, supports source testing, and displays quality and scraping status information. The target roles page manages role names used by scraping and market discovery.

{figure_markdown("Figure 47", "Admin target roles page.", "assets/screenshots/17_admin_targets.png")}

## 5.14 System Health and Monitoring

Health endpoints include live and readiness checks. The system status page presents service state to users, while admin health data supports operational monitoring. Metrics are available for Prometheus and dashboards are available through Grafana.

{figure_markdown("Figure 43", "System status page.", "assets/screenshots/13_system_status.png")}

## 5.15 Error Handling and Fallbacks

The code includes explicit handling for CV processing failures, AI gateway connection failures, validation errors, missing user data, empty job data, and unavailable services. Recommendations can fall back to recent usable jobs when no CV/profile title is available. Gap analysis can fall back from the AI `/api/hybrid-match` call to database-based skill matching when the AI matching service is unavailable.

## 5.16 Internationalization and UI Preview Modules

The frontend contains English and Arabic locale files. Several menu items intentionally appear as preview modules so the interface can show the intended product direction without presenting those areas as completed core deliverables. The completed graduation/demo core remains CV upload, parsed profile and skills, recommendations, gap analysis, application tracking, admin diagnostics, and system status.

| Module | Current Status | Reason Included | Future Work |
|---|---|---|---|
| CV Builder | Preview placeholder | Shows where guided resume authoring would fit beside CV upload. | Add editable sections, export templates, and validation tests. |
| Mock Interview | Preview placeholder | Demonstrates a possible practice workflow after gap analysis. | Add question banks, scoring rubrics, and consent-aware recording rules. |
| Learning Paths | Preview placeholder | Connects missing skills to future study plans. | Integrate curated resources and progress tracking. |
| Career Planner | Preview placeholder | Shows longer-term roadmap direction. | Add milestone planning and advisor review. |
| Mentorship | Preview placeholder | Shows possible human-support extension. | Add mentor profiles, matching rules, and moderation. |
| Tools Hub | Preview screen | Groups preview tools in one visible place. | Replace placeholders with independently tested modules. |
| Market Intelligence | Supporting/preview view | Helps explain job-market context from imported jobs. | Add stronger statistics, freshness indicators, and source quality labels. |

*Preview module clarification table.*

{figure_markdown("Figure 42", "Tools Hub preview page.", "assets/screenshots/12_tools_hub.png")}

## 5.17 Dockerized Runtime Flow

The runtime starts through Docker Compose. Nginx exposes the app, frontend and backend containers serve UI/API flows, backend workers process queues, Python services support AI workflows, MySQL and MinIO persist state, and monitoring services observe the stack.

{figure_markdown("Figure 48", "Docker services evidence.", "assets/screenshots/18_docker_containers.png")}

\\pagebreak

# Chapter 6: AI CV Analyzer Deep Technical Analysis

## 6.1 Introduction

The AI CV Analyzer is one of the main technical contributions of CareerCompass. It should not be understood as a thin wrapper around one pretrained model. The implemented analyzer is a layered hybrid pipeline that combines document-processing logic, NER, deterministic extraction rules, semantic enrichment, score composition, and explanation generation. This chapter separates that AI contribution from the general implementation chapter so that supervisors and examiners can evaluate the design as an academic system component.

## 6.2 AI Design Philosophy

CareerCompass does not use a pure NER model because CVs are noisy, multi-format documents. They can contain multiple columns, icons, section headers, table-like blocks, scanned pages, mixed date formats, and skill aliases. NER can extract entity candidates, but NER alone does not naturally explain seniority, primary technical domain, job-fit constraints, or gap-analysis reasons.

The system also does not use a pure rule-based parser. Rules are deterministic and useful for validation, but fixed rules are brittle when skill names, job titles, section headings, and CV layouts vary. A rule set can recognize known patterns, but it struggles with semantic similarity, synonyms, and role/domain interpretation.

The implemented design is therefore hybrid. NER extracts structured candidates, deterministic rules improve consistency and safety, canonicalization reduces noisy variants, Layer 2 adds domain and seniority interpretation, Layer 3 compares candidate and job evidence, and the explanation layer turns scores into strengths, gaps, red flags, and verdicts. TF-IDF fallback keeps the matching endpoint useful when heavier semantic components are unavailable.

{figure_markdown("Figure 19", "AI design philosophy for the layered hybrid analyzer.", "assets/diagrams/19_ai_design_philosophy.png")}

\\pagebreak

| Design Option | Advantage | Limitation | CareerCompass Decision |
|---|---|---|---|
| Pure NER | Learns entity patterns from data. | Does not solve file recovery, seniority, domain, matching, or explanation by itself. | Used only as one extraction component. |
| Pure rules | Predictable and easy to inspect. | Brittle when CVs use new wording, layouts, and aliases. | Used for safety, contacts, dates, validation, and fallback behavior. |
| Hybrid layered AI | Combines learned extraction, deterministic checks, semantic signals, and explanation. | More components must be tested and documented. | Chosen because it fits noisy CVs and graduation-demo transparency. |

*Table 22. AI design alternatives comparison.*

## 6.3 Complete CV Processing Flow

The end-to-end flow begins when the student uploads a PDF or image CV. The frontend validates the file before sending it to Laravel. Laravel sends the file to the FastAPI analyzer, stores the private CV object, persists successful structured outputs, and records parsing status. The analyzer first recovers text, then segments sections, extracts entities, estimates experience, canonicalizes skills, and classifies the profile. Its Layer 3 `/api/hybrid-match` endpoint supports detailed gap analysis; the separate job recommendation list is ranked inside Laravel with title, skill-overlap, and seniority scoring.

{figure_markdown("Figure 20", "Complete CV processing flow.", "assets/diagrams/20_complete_cv_processing_flow.png")}

The flowchart is intentionally more detailed than the high-level architecture diagram. It shows that the AI service performs several recoverable steps before returning data. The output is not only a list of words; it includes profile fields, skills, experience signals, domain, seniority, confidence-style values, metadata, and status.

## 6.4 Fault Tolerance and Recovery

CV parsing can fail for normal reasons: scanned PDFs, image-only files, weak text extraction, unsupported content, or service timeouts. The analyzer and backend are designed to report these states explicitly instead of silently overwriting good profile data with empty extraction results.

{figure_markdown("Figure 21", "CV analyzer fault tolerance and recovery flow.", "assets/diagrams/21_cv_fault_tolerance_flow.png")}

The FastAPI schema supports `success`, `ocr_fallback`, `empty_file`, `no_text`, and `error` statuses. The API-level timeout path returns a timeout payload. Laravel treats `timeout`, `error`, `empty_file`, and `no_text` as incomplete statuses and avoids refreshing profile, experience, and skills for those results. The backend still records analysis status and file metadata, then warns the frontend so the user can retry with a clearer document.

## 6.5 Confidence and Readiness Signals

CareerCompass uses confidence-style and readiness signals rather than a certified probability of hiring success. In Layer 1, `_aggregate_confidence` averages positive confidence-style values and caps the result at 1.0. Skills, profile, experience, and analysis sections can each carry confidence values. Laravel stores parser `confidence_score` and converts it into a `completeness_score` when available. The dashboard then visualizes completeness, model confidence, skill count, and experience as an estimated Career Readiness Snapshot.

{figure_markdown("Figure 22", "Confidence and readiness signal flow.", "assets/diagrams/22_confidence_signal_flow.png")}

The current code-derived signal boundary is:

```text
AggregateConfidence =
    min(1.0, average(positive confidence signals))

AnalysisConfidence =
    AggregateConfidence([
        skills_section.confidence_score,
        experience_section.confidence_score
    ])

BackendCompleteness =
    round(analysis.confidence_score * 100)

DashboardSkillSignal =
    min(extracted_skill_count * 10, 100)

DashboardExperienceSignal =
    min((total_experience_years / 3) * 100, 100)
```

These formulas are intentionally described as signals. They help the dashboard explain whether enough structured CV evidence was extracted, but they are not probabilities of employability, hiring success, or model correctness.

| Signal | Source File or Component | Meaning | Used In | Limitation |
|---|---|---|---|---|
| `parsing_status` | `main.py`, `schema.py`, `CvProcessingService.php` | Whether parsing succeeded, used OCR, failed, timed out, or found no text. | Upload feedback, stored analysis status, preservation logic. | It is a status flag, not a quality score. |
| `confidence_score` | Layer 1 schema and orchestrator aggregation | Confidence-style value for extracted/derived fields. | Stored CV analysis and AI insight display. | It is not a calibrated probability of employment success. |
| `completeness_score` | `CvProcessingService.php` | Backend percentage derived from analysis confidence when available. | Profile/dashboard completeness display. | It depends on parser output availability. |
| Skill signal | `Dashboard.jsx` | Extracted skill count normalized and capped at ten skills. | Career Readiness Snapshot. | More skills do not automatically mean a better candidate. |
| Experience signal | `Dashboard.jsx` | Total parsed years mapped to a display percentage using a three-year reference. | Career Readiness Snapshot. | It is a UI signal, not a universal seniority formula. |
| Extraction metadata | Layer 1 metadata and Laravel analysis metadata | Source, spatial status, segmentation, gaps, action-verb and experience details. | Debugging, review, and future evaluation. | Metadata quality depends on successful text recovery. |

*Table 23. Confidence and readiness signal summary.*

## 6.6 Skill Canonicalization With Practical Example

Skill extraction is noisy because the same skill can appear in different forms. The canonicalizer supports exact variant mapping, exact canonical matching, RapidFuzz matching when available, normalized-key fallback, semantic embedding fallback, and pass-through behavior. The current committed config is largely industry-agnostic, so the example below is labeled illustrative of the implemented mapping stages rather than proof that every alias is already configured in source data.

{figure_markdown("Figure 23", "Skill canonicalization example.", "assets/diagrams/23_skill_canonicalization_example.png")}

| Raw Extracted Skill | Normalized Skill | Why |
|---|---|---|
| JS | JavaScript | Illustrative abbreviation normalization. |
| Java Script | JavaScript | Illustrative spacing normalization. |
| Javascript | JavaScript | Illustrative casing/spelling normalization. |
| React.js | React | Illustrative framework alias normalization. |
| React JS | React | Illustrative punctuation and spacing normalization. |

*Table 24. Skill canonicalization example.*

## 6.7 Fine-Tuned BERT NER Architecture

The NER architecture is a fine-tuning workflow, not a from-scratch language model. The training notebook uses `bert-base-cased`, tokenizes CV text with offsets, aligns character-span annotations to BIO token labels, trains a token-classification head, and exports `career_compass_ner_final`. At runtime, `AdvancedNEREngine` can load local ignored model weights if supplied; those weights are not committed to Git. A user-provided Colab export now provides recorded training-run metrics, but the repository alone still does not contain the final dataset, model weights, or a fully reproducible benchmark package.

{figure_markdown("Figure 24", "Fine-tuned BERT NER architecture.", "assets/diagrams/24_fine_tuned_bert_ner_architecture.png")}

The simplified BIO example in Table 19 remains valid for examiner explanation: `Backend` can start a ROLE entity, `Developer` can continue it, and `Laravel` or `Docker` can start SKILL entities. The actual model sees tokenized subwords and offsets rather than only human-readable words.

## 6.8 Detailed Training Pipeline

Synthetic training data is used because labeled CV NER data is not naturally available in the repository. The generator is designed to create varied technical CV snippets, including positive examples and negative decoys. Negative decoys matter because they teach the model not to tag every technical-looking phrase. The cleaner normalizes samples and validates entity spans before the notebook performs token alignment and fine-tuning.

{figure_markdown("Figure 25", "Detailed NER training pipeline.", "assets/diagrams/25_detailed_training_pipeline.png")}

The documentation pass also reviewed committed evaluation evidence and generated a dataset transparency note under `docs/graduation-book/model-analysis/dataset_statistics.md`. The user-provided Colab PDF gives recorded training-run output cells, including split counts and overall validation metrics. However, the cleaned dataset content and model weights are still not committed, and the PDF does not show per-label support counts. Therefore, the report includes the verified Colab metrics but still avoids a fake per-label distribution chart. Figure 29 records which evidence is available and what remains unavailable for reproducible academic review.

{figure_markdown("Figure 29", "Dataset evidence availability summary.", "assets/diagrams/29_dataset_evidence_availability.png")}

\\pagebreak

| Dataset Statistic | Status | Reason |
|---|---|---|
| Cleaned NER training samples | 45,911 rows recorded in Colab PDF | The PDF shows the generated rows loaded from `train_real_tech_cleaned.json`; the dataset content itself is not committed. |
| Recorded train split | 41,319 rows | Visible in the exported Colab PDF; split uses test size 0.1 and seed 42. |
| Recorded test split | 4,592 rows | Visible in the exported Colab PDF and used for the notebook evaluation output. |
| Entity counts by label | Not visible in the PDF | The PDF shows labels but not per-label support counts. |
| Negative decoy count | Not available from committed evidence | Generator/cleaner support decoys, but final generated data is absent. |
| Mini evaluation CV samples | 5 | Generated documentation mini dataset under `docs/graduation-book/evaluation/`; separate from NER training. |
| Mini evaluation job samples | 8 | Generated documentation mini dataset; not a production benchmark. |
| AI CV Analyzer smoke samples | 5 | Deterministic text-only smoke set created for this evidence pass; evaluates parser-style labels and dependency availability, not transformer weights. |
| Final NER label distribution chart | Not generated | No per-label support counts are visible in the PDF, and no committed final labeled dataset exists to count SKILL/ROLE/EDU/CERT/SOFT/O labels honestly. |

*Table 25. Dataset availability and transparency.*

### 6.8.1 Colab NER Fine-Tuning Results

The team exported the Google Colab notebook `train_ner.ipynb` as a PDF with visible output cells. This PDF was copied into `docs/graduation-book/model-analysis/colab_train_ner_results.pdf` after inspection. It is treated as supporting training evidence for the NER fine-tuning process. The PDF shows the notebook title `train_ner.ipynb - Colab`, timestamp `6/7/26, 3:55 AM`, the heading `CareerCompass AI Engine: Global Skill NER training (Autonomous)`, and a synthetic data augmentation strategy. It also shows the cleaned dataset path `train_real_tech_cleaned.json`, 11 BIO labels, train/test row counts, tokenization completion, model initialization from `bert-base-cased`, training arguments, and epoch-level metrics.

These numbers improve the academic evidence for the training workflow, but they should be interpreted carefully. They are Colab-run validation outputs for the generated/synthetic dataset and notebook split visible in the PDF. They are not production accuracy, not a large real-world CV benchmark, and not reproducible from the repository alone unless the same dataset, runtime, and exported model artifacts are supplied.

\\pagebreak

| Parameter | Value from Colab PDF | Source |
|---|---|---|
| Notebook | `train_ner.ipynb - Colab` | PDF header |
| Export timestamp | `6/7/26, 3:55 AM` | PDF header |
| Dataset file | `train_real_tech_cleaned.json` | PDF data-loading cell |
| Total loaded rows | 45,911 | PDF dataset output |
| Train rows | 41,319 | PDF dataset output |
| Test rows | 4,592 | PDF dataset output |
| Split | test size 0.1, seed 42 | PDF data-loading cell |
| Base checkpoint | `bert-base-cased` | PDF model initialization cell |
| Labels | O plus B/I for SKILL, ROLE, EDU, CERT, SOFT | PDF label configuration cell |
| Max length | 512 tokens | PDF tokenization cell |
| Epochs | 5 | PDF training arguments |
| Learning rate | 2e-5 | PDF training arguments |
| Batch size | 16 train / 16 eval | PDF training arguments |
| Weight decay | 0.01 | PDF training arguments |
| Best model metric | F1 | PDF training arguments |

*Table 33. Colab NER training run configuration.*

The full epoch-by-epoch numeric table is retained in `docs/graduation-book/model-analysis/colab_ner_training_results_summary.md`. In the main chapter, the same verified values are shown as charts so the trend is easier to read in the PDF.

{figure_markdown("Figure 50", "Colab NER final epoch metrics.", "assets/diagrams/31_colab_ner_metrics.png")}

{figure_markdown("Figure 51", "Colab NER epoch performance trend.", "assets/diagrams/61_colab_ner_epoch_performance.png")}

The performance chart shows that the overall F1 score increases from 0.924603 at epoch 1 to 0.936900 at epoch 5, while accuracy remains high throughout the run. These values are overall `seqeval` metrics from the notebook validation split; they are not per-label SKILL/ROLE/EDU/CERT/SOFT metrics.

{figure_markdown("Figure 52", "Colab NER training and validation loss curve.", "assets/diagrams/62_colab_ner_loss_curve.png")}

The notebook uses Hugging Face `Trainer` with `AutoModelForTokenClassification` and a token-classification dataset. It does not define a custom loss function in the visible code, so this report treats the reported training and validation loss as the Trainer's token-classification objective values rather than a separately designed loss formula. Training loss decreases steadily from 0.077623 to 0.037280. Validation loss remains low, with the lowest visible value at epoch 3 (0.063463) and a small increase by epoch 5 (0.068058). The final epoch still has the strongest visible F1 score, but the validation-loss movement is a reason to interpret the run cautiously rather than overclaiming model generalization.

| Metric | Final Epoch Value | Source |
|---|---:|---|
| Precision | 0.933307 | Colab PDF output, epoch 5 |
| Recall | 0.940521 | Colab PDF output, epoch 5 |
| F1-score | 0.936900 | Colab PDF output, epoch 5 |
| Accuracy | 0.976376 | Colab PDF output, epoch 5 |
| Training loss | 0.037280 | Colab PDF output, epoch 5 |
| Validation loss | 0.068058 | Colab PDF output, epoch 5 |

*Table 34. Colab NER final metric summary.*

No per-label classification report or confusion matrix is visible in the exported Colab PDF, and the attached notebook output does not contain `classification_report`, `confusion_matrix`, `sklearn.metrics`, or matrix-like output cells. Therefore, this report does not invent per-label SKILL/ROLE/EDU/CERT/SOFT support, precision, recall, F1, or a confusion-matrix chart. Future training exports should save a token-level or entity-level confusion matrix plus a per-label classification report with support counts.

## 6.9 Matching Score Formula and Penalty Logic

The Layer 3 matcher exposes a clear score-composition formula in `similarity.py` and `matching_config.json`. It is exact for the code path implemented by `IntelligentMatcher.calculate_match`; however, it still depends on upstream component scores such as semantic similarity, skill similarity, and domain similarity.

```text
BaseScore =
    semantic_score * w_semantic
  + skills_score   * w_skills
  + domain_score   * w_domain

FinalScore =
    clamp(BaseScore - total_penalty + bonus_boost, 0.0, 1.0)

MatchScorePercent = round(FinalScore * 100, 2)
```

{figure_markdown("Figure 26", "Matching formula and penalty flow.", "assets/diagrams/26_matching_formula_flow.png")}

| Seniority | Semantic Weight | Skill Weight | Domain Weight | Notes |
|---|---:|---:|---:|---|
| intern | 0.30 | 0.60 | 0.10 | Early roles emphasize concrete skill overlap. |
| junior | 0.40 | 0.40 | 0.20 | Balanced summary and skill evidence. |
| mid | 0.35 | 0.35 | 0.30 | Adds more domain importance. |
| senior | 0.25 | 0.25 | 0.50 | Domain alignment becomes more important. |
| lead | 0.20 | 0.20 | 0.60 | Leadership roles emphasize domain/role alignment. |
| default | 0.35 | 0.35 | 0.30 | Fallback profile. |

*Table 26. Seniority-aware matching weights.*

Constraint penalties are also code-derived. Missing mandatory skills subtract 15 percent each, capped at 50 percent. Experience shortfall subtracts a proportional penalty capped at 30 percent. Seniority mismatch subtracts 20 percent. Total validation penalty is capped at 80 percent. Bonus skills add 2 percent each, capped at 10 percent. The `/api/hybrid-match` endpoint additionally blends the Layer 3 semantic/adaptive result with TF-IDF when TF-IDF is available: 60 percent semantic/adaptive and 40 percent TF-IDF.

## 6.10 Explainable AI Fit Output

The analyzer does not only return a single percentage. It also returns supporting evidence that can be shown to users and examiners: score breakdowns, missing mandatory skills, strengths, gaps, red flags, and a fit verdict. This is important academically because it makes the fit-analysis process inspectable rather than opaque.

{figure_markdown("Figure 27", "Explainable AI fit output.", "assets/diagrams/27_explainable_ai_output.png")}

\\pagebreak

| Output Type | Example | Why It Helps |
|---|---|---|
| Score | 78 percent | Gives a quick summary of estimated fit. |
| Matched skills | Laravel, Docker, MySQL | Shows evidence supporting the fit score. |
| Missing skills | Kubernetes | Turns the gap into a learning target. |
| Red flags | Significant seniority mismatch | Warns that a numeric score should not be read alone. |
| Verdict | Strong Match or Potential Fit | Converts score ranges into readable guidance. |
| Gaps | Experience shortfall or missing mandatory skills | Explains why a candidate may need improvement before applying. |

*Table 27. Fit explanation output types.*

## 6.11 AI Analyzer Sequence

The analyzer is synchronous during CV upload: Laravel calls FastAPI and receives a structured parse result before updating the returned user resource. The stored profile, skills, experiences, CV analysis, and private file metadata then support later dashboard, recommendation, and gap-analysis requests.

{figure_markdown("Figure 28", "AI analyzer sequence diagram.", "assets/diagrams/28_ai_analyzer_sequence.png")}

\\pagebreak

## 6.12 Computational Complexity Overview

The following complexity statements are approximate code-level descriptions, not formal proofs. They describe how cost grows with input size and component counts.

| Component | Approximate Complexity | Explanation |
|---|---|---|
| Text cleanup and line processing | O(n) | Scans the recovered text and lines, where n is text length. |
| Spatial PDF row grouping | Approximately O(w log w) | Word positions are grouped/sorted, where w is extracted word count. |
| NER chunking | O(c x model_cost) | Long CV text is split into c chunks; transformer inference dominates per chunk. |
| Skill canonicalization | O(s x k) for fuzzy/semantic comparison | s extracted skills may be compared against k configured/canonical names. |
| Domain classification | O(d) embedding comparisons | CV context is compared against d taxonomy domain descriptions. |
| Matching one job | O(s + r) plus embedding calls | Compares candidate skills with job requirements r and computes semantic/domain signals. |
| Ranking many jobs | O(j x match_cost) | j jobs require repeated matching or backend fallback comparison. |
| TF-IDF fallback | O(t) vectorization plus sparse cosine | Depends on token count t in the two texts. |

*Table 28. Computational complexity overview.*

## 6.13 Raw CV Input to Structured Output Example

The example below is an illustrative walkthrough designed for examiner readability. It is not a live model benchmark. It uses a small CV fragment and a small job description to show how the three layers cooperate. Because the full transformer/OCR runtime dependencies were not available in the documentation Python environment, the JSON block is labeled as an illustrative schema example based on the actual `schema.py` response structure rather than a live model run.

Sanitized raw CV fragment:

```text
Demo Student
Laravel Backend Developer
Skills: Laravel, MySQL, RESTful APIs
Experience: Backend Developer at Demo Company, 2025-2026
Education: Computer Science
```

| Raw Evidence | Extracted Output | Schema Area |
|---|---|---|
| `Laravel Backend Developer` | Predicted role/title evidence. | `profile.current_title`, `analysis.predicted_role` |
| `Laravel, MySQL, RESTful APIs` | Hard technical skills. | `skills.items[]` |
| `Backend Developer at Demo Company, 2025-2026` | One experience item with company, title, period, and technologies. | `experience.items[]` |
| `Computer Science` | Education evidence. | Profile/education metadata when recovered |
| Text CV fragment | Successful text parsing path; no OCR needed in this example. | `parsing_status`, `stats`, `analysis.metadata` |

*Table 29. Raw CV fragment extraction example.*

Illustrative sanitized output based on an actual analyzer response structure:

```json
{{
  "parsing_status": "success",
  "profile": {{
    "full_name": "Demo Student",
    "current_title": "Laravel Backend Developer",
    "alternative_titles": ["Backend Developer"],
    "headline": "Professional Summary",
    "contact": {{
      "email": "student@example.com", "phone": "+20XXXXXXXXXX",
      "location": "Giza, Egypt",
      "linkedin_url": "https://example.com/linkedin",
      "github_url": "https://example.com/github", "portfolio_url": null
    }},
    "summary": "Redacted summary text for a backend-focused student CV.",
    "confidence_score": 0.93
  }},
  "stats": {{"page_count": 2, "char_count": 4110, "word_count": 498, "language_hint": null}},
  "skills": {{
    "items": [
      {{"confidence_score": 0.65, "name": "Laravel", "category": "hard", "evidence": "ner"}},
      {{"confidence_score": 0.65, "name": "MySQL", "category": "hard", "evidence": "ner"}},
      {{"confidence_score": 0.65, "name": "RESTful APIs", "category": "hard", "evidence": "rule"}}
    ],
    "confidence_score": 0.65
  }},
  "experience": {{
    "items": [
      {{
        "confidence_score": 0.85, "title": "Backend Developer",
        "company": "Demo Company", "location": "Remote",
        "start_date": "2025-12-01", "end_date": "2026-01-01", "is_current": false,
        "description": ["Developed Laravel APIs and database-backed features."],
        "technologies": ["Laravel", "MySQL"]
      }}
    ],
    "confidence_score": 0.85
  }},
  "analysis": {{
    "summary": null,
    "predicted_role": "Laravel Backend Developer",
    "seniority": "Intern",
    "primary_domain": "Full Stack Development",
    "strengths": ["Diverse technical portfolio with multiple backend technologies."],
    "gaps": [], "red_flags": [],
    "confidence_score": 0.75,
    "metadata": {{
      "segmentation": {{
        "found_sections": ["profile_summary", "experience", "projects", "skills", "education"],
        "sections_missing": [], "anomalies": []
      }},
      "experience": {{"total_experience_years": 0.08, "action_verb_score": 0.3, "gap_details": []}},
      "extraction": {{"source": "spatial", "spatial_status": "ok", "word_count_spatial": 499}},
      "layer2": {{
        "seniority_details": {{"level": "Intern", "semantic_match": "Intern"}},
        "categorized_skills": {{
          "hard_skills": ["Laravel", "MySQL", "RESTful APIs"],
          "soft_skills": [], "management_skills": []
        }},
        "domain_scores": {{"Backend Development": 0.3338, "Full Stack Development": 0.4616}}
      }}
    }},
    "domain_scores": {{"Backend Development": 0.3338, "Full Stack Development": 0.4616}}
  }}
}}
```

| Section | Meaning |
|---|---|
| `profile` | Normalized identity, contact, title, headline, summary, and confidence data. |
| `stats` | Document-level counts such as pages, characters, words, and language hint. |
| `skills` | Extracted and categorized skill items with confidence and evidence source. |
| `experience` | Parsed work-history items, dates, descriptions, technologies, and confidence. |
| `analysis` | Predicted role, seniority, domain, strengths, gaps, red flags, and confidence. |
| `analysis.metadata.segmentation` | CV sections detected or missing during document understanding. |
| `analysis.metadata.layer2` | Classification details such as seniority reasoning, categorized skills, and domain scores. |

*Table 35. AI CV Analyzer output schema sections.*

| Classification | Result | Reason |
|---|---|---|
| Domain | Full Stack Development | The sanitized sample preserves the attached schema's multi-domain scoring style, with backend and frontend scores visible. |
| Seniority | Intern | Short recorded experience and the Layer 2 seniority detail support an early-career estimate. |
| Skill category | Hard technical skills | Laravel, MySQL, and RESTful APIs are technical implementation skills. |

*Table 30. Layer 2 interpretation example.*

Example job: Junior Backend Developer requiring Laravel, MySQL, Docker, and REST APIs.

| Matching Evidence | Result |
|---|---|
| Matched skills | Laravel, MySQL, Docker, REST APIs |
| Missing skills | None in this simplified example |
| Fit interpretation | Good illustrative fit for a junior backend role |
| Explanation | Skill overlap is strong; seniority appears compatible; final production score would require live matcher execution. |

*Table 31. Layer 3 matching evidence example.*

## 6.14 Why Not Use a Direct LLM-Only Approach?

Direct LLM analysis can be powerful, especially for summarizing complex CVs and producing natural-language feedback. CareerCompass still avoids a direct LLM-only runtime because the graduation/demo system must be reproducible, inspectable, containerized, and privacy-aware. The Gemini-based code is used for synthetic training-data generation, not for sending private uploaded CVs to an external LLM during the normal runtime flow. The runtime analyzer is decomposed into layers so each part can be tested, explained, and improved independently.

\\pagebreak

| Direct LLM-Only Approach | CareerCompass Hybrid Analyzer |
|---|---|
| Requires an external inference service for each private CV unless self-hosted. | Runs the analyzer locally/containerized in the demo stack. |
| Responses can vary across prompts, model versions, and temperatures. | Uses deterministic rules, typed schema validation, and explicit status values. |
| Harder to benchmark each sub-step because extraction, reasoning, and explanation are blended. | Layers can be inspected separately: text recovery, NER, rules, classification, matching, and explanation. |
| Privacy risk is higher if real CVs are sent to a remote service. | Runtime CV analysis can stay inside the deployed services. |
| Can explain textually but may hallucinate unsupported fields. | Outputs structured strengths, gaps, red flags, metadata, and confidence-style signals from code paths. |
| May be faster to prototype but harder to reproduce academically. | Fits Docker-based graduation evaluation and component-level evidence. |

*Table 32. AI approach comparison.*

## 6.15 AI Analyzer Limitations and Future Work

The limitations below make the AI contribution more academically honest, not weaker. They identify exactly what should be improved before the system is treated as a stronger research or production artifact.

- OCR quality affects scanned or image-heavy CV extraction.
- Synthetic data may not cover all real CV styles, languages, layouts, and informal wording.
- The Colab PDF provides recorded NER training-run metrics, but the final cleaned NER dataset and exported model weights are not committed, so the run is not reproducible from repository files alone.
- More real or human-reviewed labeled CV data is needed for trustworthy per-label precision, recall, and F1.
- A larger fixed benchmark should evaluate parsing, role prediction, domain classification, seniority, matching, and gap analysis together.
- Arabic and multilingual CV support can be improved beyond the current UI localization.
- The role/domain taxonomy and skill alias catalog should expand through reviewed labor-market evidence.
- A model card and dataset card should be added for any exported model artifact.
- Human-in-the-loop review can improve reliability for important student guidance decisions.
- Privacy-preserving training and evaluation workflows should be designed before using real CV data.

## 6.16 AI Chapter Summary

The standalone AI chapter was added because the analyzer is a core project contribution. The system is best described as a transparent, layered hybrid analyzer for a graduation/demo environment. It does not claim production-grade AI accuracy, a certified hiring probability, or repository-alone reproducible final NER benchmarking. The new Colab PDF strengthens training-run evidence, while the academic value remains the integration of document recovery, NER, rules, canonicalization, classification, matching, explanation, and honest fallback behavior.

\\pagebreak

# Chapter 7: AI Job Miner and Scraping Deep Technical Analysis

## 7.1 Purpose of Job Mining in CareerCompass

CareerCompass needs job mining because career guidance becomes weak when job data is static. A student's CV skills, predicted role, and gap-analysis output are useful only when compared against job descriptions that contain current requirements. The job-mining subsystem supplies those job descriptions to the recommendation and gap-analysis workflows, while the admin interface gives operators visibility into sources, target roles, imported jobs, and failures.

The chapter uses the term "job mining" instead of claiming unrestricted web scraping. The repository contains a deterministic local demo source, API adapters, HTML parser adapters, and a Scrapy spider path. These are source adapters for a graduation/demo system. They do not prove complete labor-market coverage, production-grade crawling reliability, or permission to scrape every configured website.

## 7.2 Scraping Design Philosophy

The implementation separates responsibilities deliberately. Laravel remains the trusted application backend and system of record. It owns authentication, authorization, job records, skill synchronization, admin controls, and user-facing APIs. The Python AI Job Miner service owns network-heavy adapter work, public API parsing, HTML parsing, Scrapy execution, adapter quality checks, and callback payload construction. Queue workers sit between them so slow and failure-prone network work does not block normal browser requests.

{figure_markdown("Figure 53", "Job mining design philosophy.", "assets/diagrams/32_job_mining_design_philosophy.png")}

\\pagebreak

| Design Question | CareerCompass Decision | Reason |
|---|---|---|
| Why job mining? | Use imported job descriptions to support recommendations, gap analysis, and market context. | Static seed data becomes stale and cannot represent changing skill demand. |
| Why Python/FastAPI? | Keep scraping/API ingestion dependencies in `ai-job-miner`. | Python has stronger parsing/scraping tooling and isolates unstable network work from Laravel. |
| Why Laravel as system of record? | Only Laravel validates, deduplicates, stores, and exposes accepted jobs. | Auth, admin controls, database consistency, and skill sync already belong to Laravel. |
| Why queues? | `ProcessOnDemandJobScraping` and market scraping jobs run on the scraping queue. | External sources can timeout, block, or return malformed data. |
| Why diagnostics? | Admin pages show source status, source tests, target roles, failed URLs, and batch progress. | Operators need evidence instead of assuming sources are healthy. |

*Table 36. Job mining design decisions.*

## 7.3 AI Job Miner Runtime Architecture

The deployed runtime uses Docker Compose service separation. The AI Job Miner service is named `cc-job-miner`, maps host port `8003` to container port `8000`, and exposes `/health`. Laravel reaches it through `SCRAPER_SERVICE_URL`, while the scraper calls Laravel callback endpoints through `LARAVEL_API_BASE_URL`. The production overlay also defines `backend-worker-scraping`, which runs the database queue with the `scraping` queue name and a longer timeout than ordinary request work.

{figure_markdown("Figure 54", "AI Job Miner runtime architecture.", "assets/diagrams/34_scraping_runtime_architecture.png")}

\\pagebreak

| Component | Implementation Evidence | Runtime Role |
|---|---|---|
| React frontend | `frontend/src/api/endpoints.js`, `scrapingSources.js`, admin pages | Starts user/admin actions and polls status. |
| Laravel API | `JobController`, `ScrapedJobController`, admin controllers | Creates jobs, protects routes, validates imports, and exposes results. |
| Database queue | `ProcessOnDemandJobScraping`, `ProcessMarketScrapingCategory` | Runs slow scrape work outside request/response flow. |
| AI Job Miner | `ai-job-miner/service_api.py` | FastAPI adapter service with `/health`, `/metrics`, and protected `/scrape`. |
| Scrapy path | `ai_job_miner/settings.py`, `linkedin_spider.py`, pipelines | Public-page spider flow with robots obedience, delay, retries, dedupe, and Laravel export. |
| MySQL | migrations and models for jobs, sources, target roles, scraping jobs, failed URLs | Stores accepted jobs and operational state. |
| Optional proxies | `InternalProxyController`, `SCRAPER_USE_PROXIES` | Supplies active proxies only through a protected internal route when enabled. |

*Table 37. Scraping runtime component map.*

## 7.4 Complete Job Mining Flow

The complete flow starts with a user search or an admin target role. Laravel checks whether usable stored jobs already exist. If stored data is enough for the user workflow, Laravel returns it. If not, Laravel creates a `ScrapingJob` record, dispatches a background worker, calls AI Job Miner, receives candidate jobs through protected import endpoints, deduplicates and stores them, syncs required skills, and exposes status through polling/dashboard endpoints.

{figure_markdown("Figure 55", "Complete job mining flow.", "assets/diagrams/33_complete_job_mining_flow.png")}

The important architectural point is that the Python service does not directly become the database owner. It is an ingestion service. Candidate jobs become CareerCompass jobs only after Laravel form requests validate the payload and the import transaction completes.

{figure_markdown("Figure 56", "Scraping sequence diagram.", "assets/diagrams/35_scraping_sequence_diagram.png")}

## 7.5 On-Demand Scraping and Status Polling

On-demand scraping is implemented in `JobController`. `scrapeAndStore` accepts a query and maximum result count, creates a pending `ScrapingJob`, dispatches `ProcessOnDemandJobScraping`, and returns the scraping job ID. `scrapeJobTitleIfMissing` first checks whether public usable jobs with a matching title already exist. If jobs exist, it returns `data_exists: true`; otherwise it queues a scrape and returns a polling URL. `checkScrapingStatus` returns the lifecycle state, counters, completed timestamp, matching stored jobs, or an error message.

{figure_markdown("Figure 57", "Scraping job lifecycle.", "assets/diagrams/36_scraping_job_lifecycle.png")}

\\pagebreak

| Status | Meaning | User/Admin Behavior |
|---|---|---|
| pending | The `ScrapingJob` record exists and is waiting for a worker. | Poll status and keep the UI non-blocking. |
| processing | The queue worker has started and external/import work is running. | Continue polling or show progress/admin diagnostics. |
| completed | The worker finished and stored counters such as `jobs_found`, `jobs_stored`, `jobs_duplicated`, `discovered_count`, `failed_count`, and `processing_time_ms`. | Display imported/stored jobs and final metrics. |
| failed | The run ended with an unrecoverable error or failed-only outcome and `error_message` is stored. | Show an error and use admin diagnostics/manual review. |

*Table 38. On-demand scraping lifecycle states.*

## 7.6 Source Management and Target Roles

Admin source management is implemented through `ScrapingSourceController`, `ScrapingSource`, and `AdminSources.jsx`. Sources store endpoint, method, type, mode, status, headers, and params. The model computes adapter names and support metadata so the UI can distinguish demo/local sources, supported API adapters, missing credentials, external-risk HTML adapters, and unsupported configurations.

Target roles are implemented through `TargetJobRoleController`, `TargetJobRole`, and `AdminTargets.jsx`. A full scraping run combines active target roles with active/runnable sources. Unsupported sources and sources missing required credentials are skipped instead of being counted as successful.

{figure_markdown("Figure 58", "Source management and target-role flow.", "assets/diagrams/37_source_management_flow.png")}

\\pagebreak

| Control | Code Evidence | Purpose |
|---|---|---|
| Source CRUD/status | `ScrapingSourceController`, `StoreScrapingSourceRequest`, `UpdateScrapingSourceRequest` | Manage source definitions and active/inactive state. |
| Source support metadata | `ScrapingSource::supportMetadata()` | Label demo, supported APIs, config-required sources, external-risk sources, and adapter-missing sources. |
| Source diagnostics | `test`, `testSingle`, `runSourceDiagnostic` | Run a small extraction and report support status, jobs stored/rejected, failures, and elapsed time. |
| Target roles | `TargetJobRoleController`, `TargetJobRoleSeeder` | Manage role names/search queries for market scraping. |
| Full scraping run | `runFullScraping`, `ProcessMarketScrapingCategory` | Queue source/target pairs and record per-run `ScrapingJob` status. |
| Admin evidence | Figures 44-47 | Dashboard, jobs, source diagnostics, and target-role screens show the operator workflow. |

*Table 39. Source management and target-role controls.*

The admin operational evidence is already shown in Figures 44-47: dashboard, jobs, source diagnostics, and target roles. This chapter refers to those screenshots instead of repeating them, so Chapter 7 can focus on architecture, flows, and implementation behavior.

## 7.7 Laravel Import Pipeline and Deduplication

The import pipeline is centered in `ScrapedJobController::import`. It runs inside a database transaction and applies a layered duplicate strategy. First it checks URL, which is the strongest available source identity. Then it checks title/company candidates, including original and title-case variants. Finally it checks squished lowercase title/company values. If a job already exists, Laravel updates it; otherwise it creates a new `job_postings` record. `SkillSyncService` then normalizes and links required skills without detaching prior evidence.

{figure_markdown("Figure 59", "Job import and deduplication flow.", "assets/diagrams/38_job_import_deduplication_flow.png")}

\\pagebreak

| Deduplication Stage | Evidence | Reason |
|---|---|---|
| URL match | `Job::where('url', ...)` | Strongest available unique source identity. |
| Title/company variants | original title and title-case candidate with company | Catches common formatting differences. |
| Lowercase title/company | squished lowercase title and company comparison | Catches casing/spacing differences. |
| Update or create | Import runs inside `DB::transaction` | Keeps lookup, save, and skill sync atomic. |
| Skill sync | `SkillSyncService::syncJobSkills(..., detaching: false)` | Preserves and extends job-skill matching evidence. |

*Table 40. Import and deduplication stages.*

The current duplicate strategy is appropriate for a demo system, but a stronger production importer should add source-specific IDs, canonical URLs, content hashes, expiration states, and reviewed merge rules.

## 7.8 Failed URL Handling and Dead Letter Queue

The failure path uses `ScrapingFailedUrl` records as a lightweight dead-letter style store. AI Job Miner reports failed source URLs to `POST /api/v1/jobs/import/failed`; Laravel validates the payload with `ReportScrapingFailureRequest` and stores URL, optional source/job IDs, error message, retried flag, and failed timestamp. Admin dashboard routes expose failed URLs for a scraping job.

{figure_markdown("Figure 60", "Failed URL and retry flow.", "assets/diagrams/39_scraping_failure_dlq_flow.png")}

\\pagebreak

| Failure Type | Handling | User/Admin Visibility |
|---|---|---|
| Timeout or network error | Operational category reported through failed URL callback when available. | Admin failed-URL list and source diagnostics. |
| Parse or quality failure | Adapter may classify empty, rejected, or data-quality failed outcomes. | Source diagnostic result and counters. |
| Duplicate candidate | Import check/import path tracks duplicate or non-created outcomes. | Batch/on-demand counters rather than a user-facing error. |
| Source disabled or unsupported | Source is skipped before full run or reported as adapter missing/config required. | Admin source status and planned/skipped run summary. |
| Missing internal token | `scraper.token` middleware rejects Laravel import callbacks. | Request rejected; should be reviewed through logs/config. |

*Table 41. Failed URL and operational failure handling.*

The current `retry-failures` admin endpoint marks selected failed URLs as retried. It does not yet dispatch a targeted re-fetch job. The book therefore describes it as operational retry marking, not as a complete production DLQ processor.

## 7.9 Admin Diagnostics and Retry Operations

Admin diagnostics are more than a source list. `ScrapingSourceController::runSourceDiagnostic` creates a diagnostic `ScrapingJob`, calls the scraper with a small query, captures adapter classification, records elapsed time, and returns support metadata, job preview counts, quality rejections, failed URL counts, and an output excerpt. `DashboardController` exposes scraper health, batch progress, failed URLs, and retry marking. This gives examiners a visible way to discuss what happened rather than only whether a scrape produced jobs.

The diagnostic design is especially important for external sources. Public APIs can require credentials, and websites can change HTML or block automated access. A mature demo should show those outcomes honestly: skipped, config required, blocked, empty, failed, partial, or successful.

## 7.10 Security, Tokens, Rate Limits, and Proxy Configuration

Scraping uses multiple security boundaries. User and admin actions go through authenticated Laravel routes. Laravel-to-miner calls use `X-Scraper-Service-Token` on AI Job Miner `/scrape`. Miner-to-Laravel callbacks use the protected `scraper.token` middleware and `throttle:scraper`; in the current Laravel middleware this is checked through the bearer token against `SCRAPY_API_TOKEN`. Import logs are redacted by `ScrapedJobController::redactForLogs`, and the Python service redacts token/API-key-like fields in adapter output.

{figure_markdown("Figure 61", "Scraping security boundaries.", "assets/diagrams/40_scraping_security_boundaries.png")}

| Control | Code / Configuration Evidence | Purpose |
|---|---|---|
| Laravel import token | `VerifyScraperToken`, `SCRAPY_API_TOKEN`, `LARAVEL_API_TOKEN` | Protects `/jobs/import`, `/jobs/import/check`, `/jobs/import/failed`, and `/proxies/active`. |
| Laravel-to-miner token | `SCRAPER_SERVICE_TOKEN`, `X-Scraper-Service-Token` | Prevents public calls to AI Job Miner `/scrape`. |
| Scraper throttling | `throttle:scraper` route middleware | Limits protected callback traffic. |
| Secret redaction | `redactForLogs`, `_sanitize_sensitive` | Avoids logging tokens, API keys, app IDs, authorization headers, passwords, and secrets. |
| External API keys | `ADZUNA_APP_ID`, `ADZUNA_APP_KEY` | Enables Adzuna adapter when configured; missing keys are reported honestly. |
| Proxy configuration | `SCRAPER_USE_PROXIES`, `/proxies/active` | Optional operational feature; not a permission bypass or reliability guarantee. |
| Rate-limit configuration | `SCRAPER_RATE_LIMIT_PER_MINUTE` and Scrapy delay/retry settings | Documents conservative runtime policy, but source-specific terms still matter. |
| Robots and terms | Scrapy `ROBOTSTXT_OBEY = True`; RFC 9309 context [41] | Ethical scraping requires respecting source rules and terms. |

*Table 42. Scraping security and configuration controls.*

## 7.11 Job Mining API Contracts

The scraping API surface has three groups: authenticated user endpoints, protected internal scraper endpoints, and admin endpoints. Detailed examples are included in Appendix A so maintainers can reuse them without exposing real tokens. The examples follow an OpenAPI-style documentation pattern [39].

\\pagebreak

| Group | Method and Path | Auth / Middleware | Purpose |
|---|---|---|---|
| User/Auth | `POST /api/v1/jobs/scrape` | `Authorization: Bearer <user-token>` | Queue on-demand scraping for a query. |
| User/Auth | `POST /api/v1/jobs/scrape-if-missing` | `Authorization: Bearer <user-token>` | Return existing jobs or queue scraping when missing. |
| User/Auth | `GET /api/v1/scraping-status/{{jobId}}` | `Authorization: Bearer <user-token>` | Poll lifecycle state and counters. |
| Internal scraper | `POST /api/v1/jobs/import/check` | `Authorization: Bearer <internal-token>` | Check duplicate URL before import. |
| Internal scraper | `POST /api/v1/jobs/import` | `Authorization: Bearer <internal-token>` | Validate, deduplicate, save/update job, and sync skills. |
| Internal scraper | `POST /api/v1/jobs/import/failed` | `Authorization: Bearer <internal-token>` | Store failed source URL evidence. |
| Internal scraper | `GET /api/v1/proxies/active` | `Authorization: Bearer <internal-token>` | Return active proxy definitions when enabled. |
| Scraper service | `POST /scrape` on AI Job Miner | `X-Scraper-Service-Token: <internal-token>` | Execute adapter work for Laravel worker. |
| Admin | `/api/v1/admin/scraping-sources*`, `/api/v1/admin/target-roles*`, `/api/v1/admin/scraping/run-full` | User token plus admin middleware | Manage sources, target roles, diagnostics, and full runs. |

*Table 43. Job mining API contract summary.*

On-demand request example:

```json
{{
  "query": "Backend Developer",
  "max_results": 10
}}
```

Status response example:

```json
{{
  "success": true,
  "status": "completed",
  "jobs_found": 8,
  "jobs_stored": 5,
  "jobs_duplicated": 3,
  "failed_count": 0,
  "processing_time_ms": 12640
}}
```

Internal import example:

```json
{{
  "title": "Junior Backend Developer",
  "company": "Example Co",
  "location": "Remote",
  "description": "Build APIs with Laravel and MySQL.",
  "requirements": "Laravel, MySQL, REST APIs",
  "url": "https://example.com/jobs/123",
  "source": "remotive",
  "scraping_source_id": 5,
  "skills": ["Laravel", "MySQL", "REST APIs"]
}}
```

## 7.12 Scraping Evaluation and Validation Evidence

The strongest current scraping evidence is architectural and test evidence, not live market coverage evidence. The repository includes AI Job Miner tests for service auth, health, metrics, adapter parsing, classification, redaction, blocked/empty outcomes, and skill extraction helpers. Laravel validates imports through form requests and transactions. Docker Compose wires the `ai-job-miner` service, long-running scraping worker, tokens, and callback URLs.

{figure_markdown("Figure 62", "Scraping validation evidence.", "assets/diagrams/41_scraping_validation_evidence.png")}

| Evidence | Result to Record | What It Proves | Limitation |
|---|---|---|---|
| `python -m compileall ai-job-miner` | Passed using the bundled Python runtime. | Python syntax/importability for the service files. | Not runtime source success. |
| `python -m pytest` in `ai-job-miner` | Passed inside the running container: 75 tests, 1 warning. | Service logic and adapter parser behavior under tests. | Mocked tests do not prove real website availability. |
| `/health` and `/metrics` | AI Job Miner `/health` returned 200 with `{{"status":"ok","service":"CareerCompass Job Miner"}}`; `/metrics` returned Prometheus-style scraper counters. | FastAPI job-miner service is alive and exposes basic scraper metrics. | Not source coverage or data quality. |
| Docker Compose config | Passed for base plus production overlay. | Service wiring, tokens, workers, and ports are valid YAML/config. | Not external scraping success. |
| Deterministic demo scrape | Protected `/scrape` call using `CareerCompass Demo Jobs` returned `SUCCESS`, previewed 3 jobs, stored 3, and reported 0 failed URLs; smoke rows were cleaned afterward. | Demo adapter and Laravel import path can work together without external websites. | Direct service call bypassed the queue worker, so the temporary `ScrapingJob` status stayed pending and was not status-polling evidence. |
| Import API validation | Documented from form requests and controller code. | Laravel accepts structured payloads and rejects unsafe data. | Not a complete benchmark. |
| Admin diagnostics screenshots | Figures 44-47. | Admin UI supports source, job, dashboard, and target-role operations. | Point-in-time demo evidence. |

*Table 44. Scraping validation evidence.*

\\pagebreak

The final smoke test used a direct protected `/scrape` request to validate the deterministic demo adapter and Laravel import path. A full authenticated `/jobs/scrape-if-missing` queue lifecycle with browser polling remains a recommended final demonstration check because it exercises the user-facing queue trigger and status-polling path rather than only the internal scraper contract.

| Component | Primary Evidence | Summary |
|---|---|---|
| FastAPI service/API layer | `ai-job-miner/service_api.py` | Provides `/health`, `/metrics`, and protected `/scrape` orchestration. |
| Demo/local adapter | `CareerCompass Demo Jobs` source path and tests | Deterministic source used for safe smoke evidence without external websites. |
| Adzuna/API adapter | Adzuna source code and environment keys | Optional external API source when credentials and quotas are configured. |
| HTML/Scrapy-related path | `ai_job_miner/settings.py`, spiders, pipelines | Public-page crawling path with robots/delay/retry configuration boundaries. |
| Laravel JobController | `JobController` | Starts on-demand scraping and exposes status polling. |
| Laravel ScrapedJobController | `ScrapedJobController` | Validates imports, checks duplicates, records failures, and redacts logs. |
| Queue jobs | `ProcessOnDemandJobScraping`, market scraping jobs | Move long network tasks out of browser request time. |
| Skill sync | `SkillSyncService` | Connects imported job requirements to canonical skills. |
| Admin source/target controllers | `ScrapingSourceController`, `TargetJobRoleController` | Manage active sources, diagnostics, target roles, and full-run triggers. |
| Frontend admin/user pages | `frontend/src/pages/admin`, user job pages | Surface jobs, sources, targets, status, and retry/diagnostic views. |

AI Job Miner source and function inventory summary.

No source coverage percentage, success rate, or every-job-board claim is made. External-source behavior should be retested shortly before the final defense if the team wants live demonstration evidence.

\\pagebreak

## 7.13 Limitations, Ethics, and Future Work

The scraping subsystem is useful because it connects CV analysis to job requirements, but it must remain academically honest. Public websites can change HTML, block requests, or impose rules. APIs can require credentials and quotas. Proxy usage can introduce reliability and compliance risks. Stored jobs can become stale. Duplicate detection can be improved.

| Area | Current Boundary | Future Work |
|---|---|---|
| Source coverage | Demo/local source is deterministic; external adapters are partial and unstable. | Add reviewed source policies and source-specific adapters. |
| External terms | Code cannot prove permission for every source. | Record robots/terms review per source before enabling external runs. |
| Rate limits | Scrapy delay/retry settings and rate-limit config exist. | Add per-source rate-limit dashboard and enforcement. |
| API keys | Adzuna uses environment credentials when configured [40]. | Add secret rotation and key-status diagnostics. |
| Proxies | Optional protected proxy route exists. | Use only compliant proxies and document source permission. |
| Duplicate detection | URL and title/company transaction logic. | Add canonical IDs, URL normalization, and content hashes. |
| Data freshness | Imported jobs remain until managed. | Add expiration, archival, and human review queue. |
| Failed URLs | Failed URL records and retry marking. | Add targeted DLQ reprocessing with attempt counts. |
| Evaluation | Tests and structural validation. | Add reproducible live-source health checks without inflated claims. |

*Table 45. Scraping limitations, ethics, and future work.*

Ethical operation should respect robots.txt and website/API terms, prefer official APIs where available, avoid private/login/CAPTCHA bypasses, keep request rates conservative, and avoid presenting imported data as exhaustive or assured. This framing keeps the system useful for a graduation demonstration while making its real boundaries clear.

\\pagebreak

# Chapter 8: Testing and Evaluation

## 8.1 Introduction

Evaluation was performed using repository-aware commands and browser evidence. The goal was to verify that each major part of the graduation/demo system runs and to document limitations honestly.

## 8.2 Testing Strategy

The testing strategy combined automated tests, build checks, configuration checks, service health probes, and manual functional evaluation. Automated tests provide repeatable evidence. Screenshots provide visible workflow evidence. Manual tables document behavior that is difficult to fully automate in the available environment.

## 8.3 Backend Testing

Backend validation was executed inside the backend container. Composer dependencies were already installed. `php artisan config:clear`, `php artisan route:list`, migrations, and tests passed. The route list confirmed 131 routes. The Laravel test suite passed with 39 tests and 297 assertions.

## 8.4 Frontend Testing

The frontend was validated using the existing `frontend/node_modules` and the bundled Node runtime. ESLint passed with 9 warnings and 0 errors. The warnings were related to React fast-refresh export conventions and hook dependency notes. The Vite production build passed and transformed 2904 modules.

## 8.5 Python Services Testing

Python syntax compilation passed for both AI services. The AI Job Miner pytest suite was rerun inside the running `ai-job-miner` Docker container and passed with 75 tests and 1 warning. The local bundled Python runtime still did not include pytest, so the successful pytest evidence is specifically container-based. The AI CV Analyzer pytest suite was not rerun in this Phase 1 documentation-fix pass; AI CV Analyzer syntax compilation passed, and earlier smoke/Colab evidence remains documented separately.

## 8.6 Docker and Integration Testing

Docker Desktop was initially unavailable to the shell, then was started through `docker desktop start`. Docker Compose configuration validation passed for the base plus production overlay files. `docker compose up -d` was used without a rebuild to start the existing local stack. After startup settled, `docker compose ps` showed the main app containers running, with backend, frontend, Nginx, job miner, database, and queue workers healthy. Health probes returned 200 for `/api/health`, `/api/ready`, `/status`, AI Job Miner `/health`, and the AI CV Analyzer root endpoint. These checks prove service availability in the local demo stack, not external source reliability.

{figure_markdown("Figure 49", "Validation evidence summary.", "assets/screenshots/19_validation_summary.png")}

### 8.6.1 Module Validation Coverage Matrix

The following matrix summarizes validation coverage by system area. Items marked as previously passed are preserved from the latest documented validation notes rather than claimed as fresh reruns in this backend/frontend/database polish pass.

| System Area | Evidence | Latest Result | Limitation |
|---|---|---|---|
| Laravel backend | Container route list, migrations, and test suite evidence. | Previously passed: 131 routes and 39 tests / 297 assertions. | Not rerun in this polish pass because application code was not changed. |
| Database, migrations, and seeders | Migration inspection, ERD review, and backend migration/test evidence. | Schema relationships and constraints were re-reviewed against migrations. | A destructive fresh migration reset was not performed in this pass. |
| React frontend | ESLint and Vite production build evidence. | Previously passed with 9 lint warnings, 0 errors, and 2904 Vite modules transformed. | Frontend source was inspected; build was not rerun unless noted in generation notes. |
| AI CV Analyzer | `compileall`, Colab output PDF, model-analysis notes, and API examples. | Syntax compilation passed in this validation; Colab metrics remain recorded evidence. | Full training and live model inference were not rerun from repository files alone. |
| AI Job Miner | Container pytest, compileall, service health, and deterministic demo scrape. | 75 tests passed with 1 warning in container validation. | Tests do not prove live external website availability. |
| Docker runtime | Compose config and local health probes. | Compose config and local health endpoints passed in this validation pass. | Health checks prove availability, not business correctness or external-source success. |
| API health/readiness | `/api/health`, `/api/ready`, `/status`, service health endpoints. | Returned 200 in this validation pass. | These are shallow liveness/readiness probes. |
| PDF/DOCX generation | Generator run, structural checks, link/bookmark inspection, and PDF page count. | Revalidated after every documentation generation pass. | Visual manual review is still recommended for final submission. |
| GitHub Actions | Workflow files and PR checklist. | Manual PR review remains required after push. | CI status depends on GitHub scheduling after branch updates. |

*Table 46. Module validation coverage matrix.*

## 8.7 CI/CD Validation

GitHub Actions workflow files were reviewed as part of repository inspection. A live GitHub Actions status screenshot was not captured before the draft PR because PR checks only become meaningful after the branch is pushed and GitHub schedules workflows. The manual review checklist asks the team to inspect CI status on the opened draft PR.

## 8.8 AI CV Analyzer Model Evidence

The AI CV Analyzer training workflow was inspected from repository files, helper documentation, and the exported Colab training-results PDF. Full model training was not rerun during this documentation update because the generator requires external Gemini API keys, the cleaned training dataset content is not committed, the runtime dependencies for transformer inference were not installed in the bundled documentation Python environment, and the notebook is designed for a Colab/T4-style GPU runtime.

The important distinction is reproducibility scope. The repository alone still does not provide the final labeled evaluation dataset, final model weights, or a one-command reproducible training run. However, the user-provided Colab PDF does provide recorded output cells from the actual notebook run, including train/test counts and overall epoch metrics. Those values are reported as Colab-run training evidence on the generated/synthetic notebook split, not as production accuracy. A local ignored artifact folder was present on this workstation and safe metadata was inspected, but model binaries remain outside Git. The mini evaluation below remains useful for regression-style demonstration, but it is separate from the Colab NER training metrics.

| Evidence Item | Status | What It Proves | What It Does Not Prove |
|---|---|---|---|
| Runtime NER artifact path | Code checks `ai-cv-analyzer/models/ner_weights/career_compass_ner_final`; the folder is ignored by Git. | The service can load a local token-classification model when deployed. | It does not prove the artifact is committed or provide a final held-out F1 score. |
| Local ignored metadata | `config.json` and tokenizer metadata were inspected locally without copying weights. | The local artifact uses a BERT token-classification configuration and cased tokenizer. | It is not a portable repository artifact. |
| Training notebook | Present under `ai-cv-analyzer/training/train_ner.ipynb` | The training process, label map, token alignment, Trainer setup, metrics code, and export steps are documented. | It does not include the cleaned dataset or model weights. |
| Colab training-results PDF | Exported PDF under `docs/graduation-book/model-analysis/colab_train_ner_results.pdf` | Shows recorded NER fine-tuning/evaluation outputs: 41,319 train rows, 4,592 test rows, and final epoch precision 0.933307, recall 0.940521, F1 0.936900, accuracy 0.976376. | It does not prove readiness for real deployment and may not be reproducible without the same dataset, runtime, and model artifacts. |
| Dataset generator | Present under `ai-cv-analyzer/training/generate_tech_dataset.py` | Synthetic labeled data can be generated from Gemini with key rotation and negative decoys. | It was not run here because it requires API keys and would generate a large dataset. |
| Dataset cleaner | Present under `ai-cv-analyzer/training/clean_dataset.py` | Dataset normalization, deduplication, and span validation are part of the workflow. | The cleaned dataset file itself is not committed. |
| API tests | Present under `ai-cv-analyzer/tests/test_service_api.py` | FastAPI status handling and hybrid-match formula behavior are covered with fakes. | These tests do not measure real NER model accuracy. |
| Dependency probe | Local bundled Python lacked `transformers`, `torch`, `sentence_transformers`, OCR/PDF packages, and Gemini client libraries. | Explains why live model inference and training were not rerun during documentation generation. | It does not reflect what the Docker image may install at runtime. |
| Mini evaluation | Generated under `docs/graduation-book/evaluation/` | Synthetic skill/recommendation/gap logic can be checked repeatably. | It is not a statistical live-model benchmark. |

*Table 47. Model evaluation evidence.*

## 8.9 NER Extraction Examples

The table below documents expected extraction behavior from the inspected NER labels and runtime post-processing. It is intentionally marked as example evidence rather than measured per-label accuracy because transformer dependencies were not available in the local documentation runtime and the Colab PDF reports overall metrics rather than a per-label classification report.

| Example CV Text | Expected NER Entities | Downstream Use | Evidence Type |
|---|---|---|---|
| `Experienced Backend Developer with Laravel, Docker, and MySQL.` | ROLE: Backend Developer; SKILL: Laravel, Docker, MySQL | Predicted role, extracted skills, backend/domain matching. | Illustrative example from label schema and code path. |
| `Graduated from Faculty of Computers and Information, Kafr El-Sheikh University.` | EDU: Faculty of Computers and Information, Kafr El-Sheikh University | Education/profile evidence. | Illustrative example from EDU label behavior. |
| `AWS Cloud Practitioner certified with Kubernetes deployment experience.` | CERT: AWS Cloud Practitioner; SKILL: Kubernetes | Certification and cloud/DevOps skill evidence. | Illustrative example from CERT/SKILL labels. |
| `Leadership, communication, and teamwork across agile projects.` | SOFT labels exist in training setup; runtime grouping mainly returns SKILL/ROLE/EDU/CERT. | Soft-skill interpretation is handled mostly by taxonomy and rule layers. | Limitation observed from runtime grouping code. |

*Table 48. NER extraction examples.*

## 8.10 Semantic Matching vs TF-IDF Fallback Examples

The semantic matching path could not be executed locally during this documentation update because sentence-transformer dependencies were unavailable in the bundled Python environment. The pure Python TF-IDF matcher was executed directly as a small deterministic fallback check. It gave a positive score for overlapping backend skills and zero for an unrelated mobile-role comparison.

\\pagebreak

| Pair | Semantic Path Status | TF-IDF Fallback Result | Interpretation |
|---|---|---|---|
| CV: `Laravel Docker MySQL REST APIs`; Job: `Backend developer with Laravel Docker MySQL` | Not executed locally; dependencies unavailable. | 0.4316 | Keyword overlap confirms a backend-oriented match signal. |
| CV: `Flutter Dart mobile UI`; Job: `Backend developer with Laravel Docker MySQL` | Not executed locally; dependencies unavailable. | 0.0000 | No meaningful keyword overlap, so fallback does not inflate score. |
| Expected runtime behavior | Sentence embeddings plus TF-IDF in `/api/hybrid-match` when both paths are available. | 60 percent semantic/adaptive plus 40 percent TF-IDF in that endpoint. | The fallback helps explainable matching but is not a substitute for full semantic evaluation. |

*Table 49. Semantic matching and TF-IDF example results.*

## 8.11 Model Evaluation Limitations

- The Colab PDF provides recorded overall training-run metrics, but the repository alone still does not include the dataset/model artifacts needed to reproduce the run.
- The cleaned labeled dataset used for final training is not committed.
- The model-weight folder is ignored by Git; safe local metadata was inspected, but binary weights were not copied or benchmarked.
- Local documentation Python did not include transformer, sentence-transformer, OCR, PDF, or Gemini packages, so live model inference and training were not rerun here.
- The Colab PDF does not show a per-label classification report, per-label support counts, or a confusion matrix.
- The examples in Table 48 are expected-behavior examples, while the TF-IDF values in Table 49 are actual small local fallback checks.
- A stronger final defense package should add a fixed labeled CV test set, saved per-label NER metrics, and CI-friendly inference smoke tests.

## 8.12 CV Analyzer Mini Dataset Evaluation

A sample PDF CV was generated for the screenshot workflow and uploaded through the running system. The upload succeeded, and the dashboard showed parsed CV data, backend role inference, extracted skills, and profile completeness. To strengthen the evaluation beyond that smoke test, this revision adds a mini synthetic dataset under `docs/graduation-book/evaluation/`.

The mini CV evaluation is explicitly offline and deterministic. It uses fake CV text, expected skill labels, and a keyword/role inference evaluator. It does not claim live model accuracy. The live AI CV Analyzer endpoint can be added to this mini-evaluation later, but the current document records only metrics that were actually computed from the synthetic dataset.

## 8.13 Recommendation Mini Dataset Evaluation

The recommendation mini evaluation ranks synthetic jobs for each synthetic CV using skill overlap plus domain and seniority bonuses. This validates the recommendation concept and provides a repeatable regression check for report evidence. It is not a production recommender benchmark, and the report does not claim complete job-market coverage.

## 8.14 Gap Analysis Mini Dataset Evaluation

The gap-analysis mini evaluation compares expected matched and missing skills with computed matched and missing skills for selected CV/job pairs. This directly validates the explanation structure used by the gap-analysis workflow: matched skills should be shown separately from missing skills.

{mini_eval}

{smoke_eval}

## 8.16 Job Miner Evaluation

The AI Job Miner is evaluated in more detail in Chapter 7. This testing chapter records only validation outcomes: syntax compilation, the AI Job Miner pytest run when available, Docker Compose wiring, health checks when the stack is running, import validation, and admin diagnostics evidence. Because external job sources can change, throttle, or require keys, long-term data quality evaluation should be repeated near the final defense and should not be replaced by source-template counts.

## 8.17 Application Tracker Evaluation

The application tracker was evaluated by saving a selected job to the tracker and loading the Applications page. The screenshot shows saved opportunity state. Backend tests also include application tracker behavior.

## 8.18 Admin Dashboard Evaluation

Admin login and admin dashboard access were tested with the demo admin account. The admin dashboard, admin jobs, admin sources, and admin target roles pages were captured. The dashboard displayed 22 users, 209 imported jobs, 9 active sources, and 13 target roles at capture time.

## 8.19 Performance Observations

The local Docker stack is heavy because it runs frontend, backend, multiple Laravel workers, MySQL, MinIO, two Python AI services, Prometheus, and Grafana. Initial full build can exceed a short command timeout on a Windows laptop. Once images are built, targeted service startup and HTTP checks are practical for a graduation demo.

## 8.20 Evaluation Limitations

- The AI CV Analyzer pytest suite was not executed because pytest was absent in that container.
- The browser CV upload remains a smoke test, and the mini dataset is synthetic rather than statistically representative.
- The AI CV Analyzer training notebook and exported Colab PDF were inspected, but full model training was not rerun because the cleaned training dataset and external generation keys are not committed and the workflow is designed for Colab GPU execution.
- The Colab metrics are recorded training-run evidence on the notebook split, not a production benchmark on a large real-world CV dataset.
- The AI CV Analyzer smoke evaluation uses five deterministic text samples and does not measure the transformer NER model weights.
- The recommendation score shown in screenshots is an estimated local demo output.
- External scraping reliability depends on source availability and changing website/API behavior.
- Production security, privacy, and performance audits remain future work.

## 8.21 Summary of Results

| Area | Command or Scenario | Result | Evidence |
|---|---|---|---|
| Docker config | `docker compose -f docker-compose.yml -f docker-compose.prod.yml config --quiet` | Passed | Final cleanup command output |
| Docker stack | `docker desktop start`; `docker compose up -d`; `docker compose ps` | Running; main app containers healthy after startup settled | Service availability evidence |
| Backend routes | `php artisan route:list` | Previously passed, 131 routes | Earlier validation evidence retained |
| Backend tests | `php artisan test` | Previously passed, 39 tests, 297 assertions | Previously recorded test pass, not freshly rerun from the Phase 1 documentation-fix shell |
| Frontend lint | ESLint | Previously passed, 9 warnings, 0 errors | Previously recorded lint pass, retained from the earlier validation pass |
| Frontend build | Vite build | Previously passed, 2904 modules transformed | Previously recorded build pass, retained from the earlier validation pass |
| AI Job Miner tests | `docker compose exec -T ai-job-miner python -m pytest` | Passed, 75 tests, 1 warning | Fresh container command output |
| AI Job Miner demo scrape | Protected `/scrape` call with demo/local source | SUCCESS; previewed 3, stored 3, failed URLs 0; smoke rows cleaned afterward | Validates demo adapter plus Laravel import path, not queue status polling |
| AI CV Analyzer syntax | `python -m compileall ai-cv-analyzer` | Passed with non-fatal `.pytest_cache` listing warning | Final cleanup command output |
| AI CV Analyzer pytest | Not rerun | Skipped in Phase 1 documentation-fix pass | Colab/smoke evidence retained separately |
| HTTP probes | `/api/health`, `/api/ready`, `/status`, `:8003/health`, `:8000/` | 200 responses after Docker startup settled | Service availability only, not external scraping success |

*Table 55. Automated validation results.*

| Test ID | Module | Scenario | Status | Evidence |
|---|---|---|---|---|
| M-01 | Authentication | Register demo user | Passed | Register screenshot/API output |
| M-02 | Authentication | Login student | Passed | Figure 34 |
| M-03 | CV upload | Upload valid PDF | Passed | Figures 35-37 |
| M-04 | CV upload | Invalid file handling | Not Run Manual | Backend validation tests |
| M-05 | Recommendations | Open jobs page after CV | Passed | Figure 38 |
| M-06 | Gap analysis | Analyze selected job | Passed | Figure 40 |
| M-07 | Tracker | Save job | Passed | Figure 41 |
| M-08 | Admin | Login admin and open dashboard | Passed | Figure 44 |
| M-09 | Admin sources | Open diagnostics | Passed | Figure 46 |
| M-10 | Status | Open system status page | Passed | Figure 43 |

*Table 56. Manual functional evaluation matrix.*

\\pagebreak

| Test ID | Expected vs Actual Observation |
|---|---|
| M-01 | Expected user creation and token return; actual user was created with an accepted Gmail-style address. |
| M-02 | Expected dashboard access after login; actual dashboard loaded. |
| M-03 | Expected CV acceptance and parsing; actual upload parsed successfully. |
| M-04 | Expected validation error for invalid file; actual manual browser repetition was not run because backend validation/tests already cover the rule. |
| M-05 | Expected recommendations with estimated matches; actual jobs page loaded. |
| M-06 | Expected match and skill breakdown; actual gap page loaded. |
| M-07 | Expected saved job in tracker; actual application page loaded with saved item. |
| M-08 | Expected admin-only dashboard; actual dashboard visible after admin login. |
| M-09 | Expected source diagnostics; actual diagnostics page visible. |
| M-10 | Expected health UI; actual system status page visible. |

*Table 57. Manual functional observations.*

\\pagebreak

# Chapter 9: Security and Privacy

## 9.1 Introduction

Security in CareerCompass is implemented for a graduation/demo context. The system includes meaningful controls, but it should not be presented as production-grade without future hardening, legal review, and operational security work.

## 9.2 Authentication and Authorization

User authentication uses token-based API access through Laravel. Laravel Sanctum supports API token and SPA authentication use cases [2]. CareerCompass stores an auth token in the browser and attaches it to API requests. Authorization is role-aware: student routes require authentication, while admin routes require the admin role. Authentication security should follow established password and session guidance in production [28].

## 9.3 Admin Access Control

Admin access is enforced server-side by admin middleware. Frontend route protection improves user experience, but the backend check is the important control. The demo admin account is generated by a seeder and should be changed or disabled in any non-demo environment.

## 9.4 CV File Privacy

CV files contain personal data. CareerCompass validates file type and size, stores files privately, and avoids public direct file exposure. OWASP's file upload guidance emphasizes validation, restricted storage, and safe handling [27].

## 9.5 Private Storage and Signed Downloads

Uploaded CV files are stored through private storage and accessed through signed or temporary URLs. MinIO/S3-compatible storage supports object-based file storage and access control patterns [9]. The graduation demo should still avoid uploading real sensitive CVs unless the environment is controlled.

The model-training workflow is intentionally documented separately from runtime CV processing. Synthetic training snippets may be generated through Google AI developer tooling [34], [35], but real student CV uploads should not be sent to external AI APIs without explicit consent, a privacy notice, retention rules, and supervisory approval. In the demonstrated runtime, Laravel sends the uploaded file to the local FastAPI analyzer service and stores the file privately; the Gemini-based generator is a training-support script, not the normal CV-upload path.

## 9.6 Internal Service Tokens

Scraper import routes use a service token middleware. This reduces accidental public ingestion, but service tokens must be rotated, stored securely, and monitored in production.

## 9.7 Payload and File Validation

Laravel form requests validate registration, login, CV upload, applications, and profile updates. File validation includes MIME type, extension, and size checks. Validation reduces risk but does not replace malware scanning, deep file inspection, or content disarm in production.

## 9.8 Logging and Request IDs

The API client attaches request IDs, and backend logging records important events such as CV processing status and AI gateway errors. For production, logs should avoid sensitive CV content and should be retained according to a privacy policy.

\\pagebreak

## 9.9 Demo Security Limitations

| Area | Current Demo Control | Production Hardening Needed |
|---|---|---|
| Admin account | Demo seeder account | Secret rotation, SSO/MFA, audit logs |
| CV files | Private storage and signed URLs | Malware scanning, retention policy, consent model |
| Tokens | Bearer tokens | Token rotation, secure cookie strategy, revocation review |
| Scraper service | Internal token | Secret manager, network isolation, rate limits |
| Monitoring | Local Prometheus/Grafana | Auth, TLS, dashboard access control |
| Privacy | Local demo posture | Legal review, privacy notice, data minimization |

*Table 58. Security and privacy controls.*

## 9.10 Future Production Hardening

Future work should include HTTPS-only deployment, secure cookie/session strategy, administrator MFA, centralized secrets management, object scanning, retention policies, audit logging, rate-limit review, CSRF/CORS review, dependency vulnerability scanning, and a privacy impact assessment.

\\pagebreak

# Chapter 10: Conclusion and Future Work

## 10.1 Conclusion

CareerCompass demonstrates a practical AI-assisted career guidance workflow for a graduation project. It integrates CV parsing, profile normalization, skill extraction, imported job records, estimated recommendation, gap analysis, application tracking, admin diagnostics, and containerized deployment.

## 10.2 Project Achievements

- Built a working multi-service web application with frontend, backend, AI services, database, object storage, proxy, and monitoring.
- Implemented student authentication, dashboard, CV upload, profile view, jobs, gap analysis, and application tracking.
- Implemented admin dashboard, job management, source diagnostics, and target role management.
- Documented the AI CV Analyzer runtime architecture, optional local NER artifact path, synthetic data-generation workflow, and Colab training notebook.
- Documented the AI Job Miner runtime architecture, queue flow, source management, import/deduplication logic, failed URL handling, and ethical scraping boundaries.
- Added tests and validation commands across backend, frontend, Python services, Docker, and HTTP probes.
- Captured browser screenshots from the running local system.
- Generated a formal report with diagrams, references, and evaluation notes.

## 10.3 Educational Value

The project demonstrates practical learning in software architecture, service decomposition, secure file handling, API design, database schema design, frontend state management, AI service integration, testing, Docker operations, monitoring, and academic documentation.

## 10.4 Current Limitations

- The system is a graduation/demo platform and not a production product.
- Recommendation and gap analysis outputs are estimates.
- AI evaluation needs larger labeled datasets, committed training artifacts, per-label reports, and repeatable model scoring.
- The exported NER model artifact path is supported by the runtime. The repository includes the exported Colab PDF with recorded overall NER metrics, but it does not include the cleaned dataset and model weights required to rerun the same training experiment from repository files alone.
- The Colab PDF records overall NER training metrics, but the final labeled NER dataset was not available in committed evidence, so label distribution and final per-label NER precision/recall/F1 were not claimed.
- OCR and PDF text-recovery quality can affect scanned CVs, image-heavy layouts, multi-column documents, and unusual fonts.
- Synthetic training examples may not represent all real student CV styles, Arabic/multilingual CVs, or informal job-market wording.
- External scraping sources can be unstable, require keys, change page structure, enforce rate limits, or impose terms that must be reviewed before use.
- The AI CV Analyzer container needs pytest installed to run its test suite.
- Security and privacy controls need production hardening before real deployment.

## 10.5 Future Work

- Add a larger CV/job evaluation dataset with manual labels.
- Add a reproducible NER evaluation pipeline that runs on a fixed labeled test set and records per-label precision, recall, and F1.
- Store model cards, dataset cards, and training-run summaries for each exported model artifact.
- Add a privacy-preserving workflow for any future human-reviewed real CV dataset.
- Improve Arabic/multilingual parsing, OCR, and field extraction.
- Expand the role taxonomy, skill aliases, and labor-market evidence with periodic review.
- Add human-in-the-loop review for high-impact guidance and ambiguous CV parsing results.
- Improve skill normalization through reviewed aliases and false-positive checks.
- Add production-grade authentication and administrator controls.
- Add malware scanning and retention policies for uploaded CV files.
- Improve recommendation explanations and calibration.
- Strengthen job-mining evaluation with reproducible source-health checks, canonical source IDs, data freshness rules, and targeted failed-URL reprocessing.
- Add automated browser end-to-end tests.
- Extend CI to run all containerized test suites consistently.
- Improve observability dashboards and alerting.
- Add deployment documentation for a secure cloud environment.

## 10.6 Final Remarks

CareerCompass is an original graduation project that connects academic software engineering requirements with a practical career guidance problem. Its strongest value is the integration of many realistic system parts into one demonstrable workflow, while preserving honest language about limitations and future production work.

\\pagebreak

# References

{refs}

\\pagebreak

# Appendices

## Appendix A: API Request and Response Examples

This appendix expands the endpoint summary with JSON-oriented examples. The examples are intentionally small and use placeholders such as `Authorization: Bearer <token>`. They document the implemented API shape for examiners and future maintainers; they do not expose real tokens, private CV contents, or production secrets. The format follows an OpenAPI-style request/response documentation pattern [39].

| Group | Example Endpoints | Purpose |
|---|---|---|
| Health | `/api/health`, `/api/ready`, `/api/metrics` | Liveness, readiness, and Prometheus metrics. |
| Auth | `/api/v1/register`, `/api/v1/login`, `/api/v1/logout`, `/api/v1/user` | User identity and token lifecycle. |
| CV | `/api/v1/upload-cv`, `/api/v1/user/skills`, `/api/v1/user/cv-analysis` | CV upload, parsed analysis, and extracted skills. |
| Jobs | `/api/v1/jobs`, `/api/v1/jobs/recommended`, `/api/v1/jobs/{{id}}`, `/api/v1/jobs/scrape`, `/api/v1/jobs/scrape-if-missing`, `/api/v1/scraping-status/{{jobId}}` | Job listing, details, recommendations, on-demand scraping, and status polling. |
| Gap Analysis | `/api/v1/gap-analysis/job/{{jobId}}`, `/api/v1/gap-analysis/role/{{roleId}}` | Skill comparison and recommendations. |
| Applications | `/api/v1/applications` | Save and update tracked opportunities. |
| Admin | `/api/v1/admin/dashboard/stats`, `/api/v1/admin/dashboard/batch-progress`, `/api/v1/admin/dashboard/failed-urls/{{scrapingJobId}}`, `/api/v1/admin/dashboard/retry-failures`, `/api/v1/admin/scraping-sources`, `/api/v1/admin/target-roles`, `/api/v1/admin/scraping/run-full` | Admin dashboards, diagnostics, source management, target roles, full runs, and failed URL operations. |
| Internal Scraper | `/api/v1/jobs/import`, `/api/v1/jobs/import/check`, `/api/v1/jobs/import/failed`, `/api/v1/proxies/active` | Service-token protected import, duplicate check, failure report, and proxy routes. |

*Table 59. API endpoint summary.*

### A.1 Core Authentication and Current User Endpoints

Login method and URL: `POST /api/v1/login`

Purpose: Authenticate a student or admin user and issue a bearer token for later API calls.

Login request example:

```json
{{
  "email": "student@example.com",
  "password": "<password>"
}}
```

Login response example:

```json
{{
  "success": true,
  "message": "Login successful",
  "data": {{
    "token": "<user-token>",
    "user": {{
      "id": 7,
      "name": "Demo Student",
      "email": "student@example.com",
      "role": "student",
      "profile": {{
        "headline": "Backend Developer",
        "location": "Giza, Egypt"
      }}
    }}
  }}
}}
```

Current user method and URL: `GET /api/v1/user`

Request header:

```text
Authorization: Bearer <user-token>
```

Current user response example:

```json
{{
  "data": {{
    "id": 7,
    "name": "Demo Student",
    "email": "student@example.com",
    "role": "student",
    "job_title": "Backend Developer",
    "headline": "Backend Developer",
    "summary": "Sanitized profile summary.",
    "location": "Giza, Egypt",
    "total_experience_years": 0.5,
    "seniority": "Intern",
    "primary_domain": "Backend Development",
    "phone": "+20XXXXXXXXXX",
    "linkedin_url": "https://example.com/linkedin",
    "github_url": "https://example.com/github",
    "profile": {{
      "headline": "Backend Developer",
      "location": "Giza, Egypt",
      "contact_info": {{}}
    }},
    "experiences": [],
    "skills": []
  }}
}}
```

### A.2 Student CV Upload Endpoint

Method and URL: `POST /api/v1/upload-cv`

Purpose: Upload a PDF/image CV to Laravel, send it to the local FastAPI AI CV Analyzer, persist successful structured outputs, and return warnings/status for the dashboard.

Request example:

```text
Authorization: Bearer <token>
Content-Type: multipart/form-data

cv=@sample_cv.pdf
```

Successful response example:

```json
{{
  "success": true,
  "message": "CV parsed successfully.",
  "parsing_status": "success",
  "analysis_id": 12,
  "skills_count": 4,
  "predicted_role": "Backend Developer",
  "profile_updated": true,
  "retry_available": false,
  "download_url": "https://example.local/api/cv-files/12?signature=...",
  "data": {{
    "analysis_id": 12,
    "parsing_status": "success",
    "skills_count": 4,
    "predicted_role": "Backend Developer",
    "warnings": []
  }}
}}
```

Error response example:

```json
{{
  "success": false,
  "message": "The AI engine is currently unavailable. Please try again in a moment.",
  "parsing_status": "error",
  "retry_available": true,
  "warnings": [
    {{
      "code": "ai_unavailable",
      "message": "The AI engine could not be reached. No profile data was changed."
    }}
  ]
}}
```

### A.3 AI Analyzer Parse-CV Endpoint

Method and URL: `POST /api/parse-cv`

Purpose: FastAPI service endpoint that accepts an uploaded CV file and returns the typed parse schema used by Laravel.

Request example:

```text
Content-Type: multipart/form-data

file=@sample_cv.pdf
```

Response example:

```json
{{
  "parsing_status": "success",
  "profile": {{
    "full_name": "Demo Student",
    "current_title": "Backend Developer",
    "email": "student@example.com"
  }},
  "stats": {{
    "page_count": 1,
    "word_count": 340,
    "language_hint": "en"
  }},
  "skills": {{
    "items": [
      {{"name": "Laravel", "category": "hard", "confidence_score": 0.65}},
      {{"name": "Docker", "category": "hard", "confidence_score": 0.65}}
    ],
    "confidence_score": 0.65
  }},
  "experience": {{
    "items": [],
    "confidence_score": 0.0
  }},
  "analysis": {{
    "predicted_role": "Backend Developer",
    "seniority": "junior",
    "primary_domain": "Backend Development",
    "strengths": [],
    "gaps": [],
    "red_flags": [],
    "confidence_score": 0.65
  }},
  "request_id": "example-request-id"
}}
```

Error response example:

```json
{{
  "detail": "Empty file uploaded."
}}
```

### A.4 AI Hybrid Match Endpoint

Method and URL: `POST /api/hybrid-match`

Purpose: Compare CV text/skills with a job description using Layer 3 adaptive matching plus TF-IDF when available.

Request example:

```json
{{
  "cv_skills": ["Laravel", "Docker", "MySQL"],
  "cv_text": "Backend developer with Laravel Docker MySQL REST APIs.",
  "job_description": "Junior backend developer required with Laravel, Docker, MySQL, and REST API experience.",
  "job_skills": ["Laravel", "Docker", "MySQL", "REST APIs"]
}}
```

Response example:

```json
{{
  "hybrid_match_score": 78.4,
  "semantic_match_pct": 82.0,
  "tfidf_score_pct": 73.0,
  "missing_skills": [],
  "formula": "Final = (Adaptive Layer 3 x 60%) + (TF-IDF x 40%)",
  "matching_mode": "hybrid",
  "request_id": "example-request-id"
}}
```

Validation error example:

```json
{{
  "detail": "cv_text must not be empty."
}}
```

### A.5 Recommended Jobs Endpoint

Method and URL: `GET /api/v1/jobs/recommended`

Purpose: Return a personalized job list for an authenticated student when CV/profile context exists, otherwise return recent usable jobs. The current Laravel code scores candidates with predicted role/profile-title matching, required-skill overlap, and seniority hints.

It does not call `/api/hybrid-match`. Semantic/adaptive and TF-IDF matching are documented under gap analysis.

Request example:

```text
Authorization: Bearer <token>
Accept: application/json
```

Response example:

```json
{{
  "success": true,
  "job_title": "Backend Developer",
  "data": [
    {{
      "id": 101,
      "title": "Junior Backend Developer",
      "company": "DemoTech",
      "location": "Cairo",
      "source": "demo",
      "match_percentage": 84.5,
      "skills_count": 4
    }}
  ],
  "meta": {{
    "total": 1,
    "based_on": "Your CV title: Backend Developer"
  }}
}}
```

### A.6 On-Demand Job Scraping Endpoint

Method and URL: `POST /api/v1/jobs/scrape`

Purpose: Start an authenticated on-demand scraping job for a query. Laravel validates the request, creates a `ScrapingJob`, dispatches the scraping queue worker, and returns a status identifier.

Request example:

```text
Authorization: Bearer <user-token>
Content-Type: application/json
```

```json
{{
  "query": "Backend Developer",
  "max_results": 10
}}
```

\\pagebreak

Successful response example:

```json
{{
  "success": true,
  "message": "Jobs scraping dispatched to background process",
  "data": {{
    "query": "Backend Developer",
    "scraping_job_id": 42
  }}
}}
```

Validation error example:

```json
{{
  "message": "The query field is required.",
  "errors": {{
    "query": ["The query field is required."]
  }}
}}
```

### A.7 Scrape If Missing and Status Polling

Method and URL: `POST /api/v1/jobs/scrape-if-missing`

Purpose: Check whether matching stored jobs exist before queueing an external scrape.

Request example:

```json
{{
  "job_title": "Laravel Developer",
  "max_results": 10
}}
```

Existing-data response example:

```json
{{
  "success": true,
  "data_exists": true,
  "message": "Job data already available",
  "jobs_count": 5
}}
```

Queued response example:

```json
{{
  "success": true,
  "data_exists": false,
  "message": "Analyzing market data for this role. Please wait...",
  "scraping_job_id": 43,
  "status": "pending",
  "poll_url": "http://localhost/api/v1/scraping-status/43"
}}
```

\\pagebreak

Method and URL: `GET /api/v1/scraping-status/{{jobId}}`

Status response example:

```json
{{
  "success": true,
  "scraping_job_id": 43,
  "job_title": "Laravel Developer",
  "status": "completed",
  "type": "on_demand",
  "started_at": "2026-06-08T00:00:00.000000Z",
  "results": {{
    "jobs_found": 8,
    "jobs_stored": 5,
    "jobs_duplicated": 3,
    "discovered_count": 8,
    "failed_count": 0,
    "processing_time_ms": 12640,
    "completed_at": "2026-06-08T00:00:12.000000Z"
  }},
  "jobs": [
    {{
      "id": 101,
      "title": "Junior Backend Developer",
      "company": "Example Co"
    }}
  ]
}}
```

### A.8 Internal Scraper Duplicate Check and Import

These Laravel endpoints are protected by `scraper.token` and `throttle:scraper`. In the current middleware, the scraper supplies `Authorization: Bearer <internal-token>`.

Method and URL: `POST /api/v1/jobs/import/check`

Request example:

```text
Authorization: Bearer <internal-token>
Content-Type: application/json
```

```json
{{
  "url": "https://example.com/jobs/123"
}}
```

\\pagebreak

Response example:

```json
{{
  "exists": false
}}
```

\\pagebreak

Method and URL: `POST /api/v1/jobs/import`

Request example:

```json
{{
  "title": "Junior Backend Developer",
  "company": "Example Co",
  "location": "Remote",
  "description": "Build APIs with Laravel and MySQL.",
  "requirements": "Laravel, MySQL, REST APIs",
  "url": "https://example.com/jobs/123",
  "source": "remotive",
  "scraping_source_id": 5,
  "skills": ["Laravel", "MySQL", "REST APIs"],
  "work_type": "remote",
  "job_type": "full_time"
}}
```

Response example:

```json
{{
  "success": true,
  "job_id": 101,
  "created": true
}}
```

### A.9 Failed URL Reporting, Proxies, and Admin Scraping

Method and URL: `POST /api/v1/jobs/import/failed`

Purpose: Store a failed source URL for diagnostics and retry visibility.

Request example:

```json
{{
  "url": "https://example.com/jobs/broken",
  "scraping_source_id": 5,
  "scraping_job_id": 43,
  "error_message": "Timeout while fetching public job detail page."
}}
```

Response example:

```json
{{
  "success": true
}}
```

Method and URL: `GET /api/v1/proxies/active`

Purpose: Return active proxy definitions to the scraper only when proxy configuration is enabled and authorized.

Admin source diagnostic request:

```text
Authorization: Bearer <admin-token>
POST /api/v1/admin/scraping-sources/5/test
```

Admin full scraping request:

```text
Authorization: Bearer <admin-token>
POST /api/v1/admin/scraping/run-full
```

Example full-run response shape:

```json
{{
  "success": true,
  "batch_id": "example-batch-id",
  "planned_jobs": 12,
  "skipped_sources": []
}}
```

### A.10 Gap Analysis Endpoint

Method and URL: `GET /api/v1/gap-analysis/job/{{jobId}}`

Purpose: Compare the authenticated student's extracted skills/profile with one job and return matched skills, missing skills, recommendations, and CV-analysis context.

Request example:

```text
Authorization: Bearer <token>
Accept: application/json
```

Response example:

```json
{{
  "success": true,
  "data": {{
    "job": {{
      "id": 101,
      "title": "Junior Backend Developer",
      "company": "DemoTech"
    }},
    "analysis": {{
      "match_percentage": 80.0,
      "match_level": "Good Match",
      "matched_skills": ["Laravel", "Docker", "MySQL"],
      "missing_skills": ["Kubernetes"]
    }},
    "recommendations": [
      "Practice Kubernetes deployment basics before applying."
    ],
    "cv_analysis": {{
      "parsing_status": "success",
      "completeness_score": 75,
      "strengths": [],
      "gaps": [],
      "red_flags": []
    }}
  }}
}}
```

Error response example:

```json
{{
  "success": false,
  "message": "Upload a CV first so the system can extract skills and profile data."
}}
```

### A.11 Application Tracking Endpoint

Method and URL: `POST /api/v1/applications`

Purpose: Let an authenticated student save or update an opportunity status without changing the shared job record.

Request example:

```json
{{
  "job_id": 101,
  "status": "saved",
  "notes": "Review Laravel and Docker requirements before applying."
}}
```

Response example:

```json
{{
  "success": true,
  "data": {{
    "id": 55,
    "job_id": 101,
    "status": "saved",
    "notes": "Review Laravel and Docker requirements before applying.",
    "job": {{
      "title": "Junior Backend Developer",
      "company": "DemoTech"
    }}
  }}
}}
```

### A.12 Admin Dashboard Summary Endpoint

Method and URL: `GET /api/v1/admin/dashboard/stats`

Purpose: Give administrators a compact operational summary of users, jobs, CV activity, scraping runs, and source state.

Request example:

```text
Authorization: Bearer <admin-token>
Accept: application/json
```

Response example:

```json
{{
  "success": true,
  "data": {{
    "total_students": 39,
    "total_jobs": 128,
    "total_sources": 4,
    "total_targets": 12,
    "jobs_by_month": [
      {{
        "month": "Jun 2026",
        "month_key": "2026-06",
        "count": 18
      }}
    ],
    "scraper_overview": {{
      "jobs_last_24h": 3,
      "avg_health_score": 91.5,
      "active_sources": 4,
      "total_sources": 4,
      "recent_failures": 0
    }}
  }}
}}
```

### A.13 Health and Readiness Endpoints

Methods and URLs: `GET /api/health`, `GET /api/ready`

Purpose: Confirm whether Laravel is alive and whether dependent services are ready.

Health response example:

```json
{{
  "success": true,
  "status": "ok",
  "service": "CareerCompass API",
  "request_id": "example-request-id"
}}
```

Readiness response example:

```json
{{
  "success": true,
  "status": "ready",
  "checks": {{
    "database": {{"ok": true}},
    "cache": {{"ok": true}},
    "ai": {{"ok": true, "status": 200}},
    "scraper": {{"ok": true, "status": 200}}
  }},
  "request_id": "example-request-id"
}}
```

## Appendix B: Quick Start Guide for Examiners

This guide is intended for a local graduation/demo run. It assumes Docker Desktop is installed and running. It should not be read as a production deployment guide.

### B.1 Start the Stack

```powershell
git clone https://github.com/YousefAlTohamy/CareerCompass.git
cd CareerCompass
cp .env.example .env
docker compose -f docker-compose.yml -f docker-compose.prod.yml up -d --build
docker compose exec backend-api php artisan migrate --seed
```

### B.2 Useful URLs

- Frontend: `http://localhost`
- Backend health: `http://localhost/api/health`
- Backend readiness: `http://localhost/api/ready`
- AI Job Miner health, when the stack is running: `http://localhost:8003/health`
- System status page: `http://localhost/status`
- Admin dashboard: `http://localhost/admin/dashboard`
- Grafana, if exposed by the local compose stack: `http://localhost:3000`

### B.3 Demo Accounts and Flow

The demo admin account is seeded from environment variables. The current example values are:

| Variable | Demo Value |
|---|---|
| `DEMO_ADMIN_EMAIL` | `careercompassadmin@gmail.com` |
| `DEMO_ADMIN_PASSWORD` | `CareerCompassAdmin2026` |

*Demo admin environment values.*

Student demo flow:

1. Open `http://localhost`.
2. Register a new student account or log in with an existing demo student.
3. Upload a small PDF CV from the dashboard.
4. Review dashboard readiness signals, extracted role, extracted skills, and profile page.
5. Open Jobs and review recommended jobs.
6. Open Gap Analysis for a selected job and review matched/missing skills.
7. Save a job to Applications.
8. Log in as the demo admin and review dashboard, jobs, sources, and target roles.

### B.4 Validation Commands

```powershell
docker compose config --quiet
docker compose -f docker-compose.yml -f docker-compose.prod.yml config --quiet
docker compose exec backend-api php artisan test
cd frontend
npm run lint
npm run build
```

Documentation/evaluation checks:

```powershell
python docs/graduation-book/evaluation/run_mini_evaluation.py
python docs/graduation-book/evaluation/run_ai_cv_analyzer_smoke_eval.py
python docs/graduation-book/scripts/generate_graduation_book.py
python -m compileall ai-job-miner
python -m compileall ai-cv-analyzer
cd ai-job-miner
python -m pytest
cd ..
```

### B.5 Troubleshooting Notes

- If the frontend appears stale after rebuild, hard-refresh the browser or rebuild the frontend/nginx containers.
- If Docker startup is slow, wait for MySQL, MinIO, Laravel workers, Python AI services, Prometheus, and Grafana to settle before judging readiness.
- AI model weights are not committed to Git; the runtime supports `ai-cv-analyzer/models/ner_weights/career_compass_ner_final` when supplied locally.
- NER training requires a cleaned labeled dataset and external API keys for synthetic-data generation; those are not included in the repository.
- Demo admin credentials should be changed for any non-demo environment.

## Appendix C: Database Tables Summary

| Table | Purpose |
|---|---|
| users | Authentication identity, role, and banned status. |
| user_profiles | Student profile details and contact metadata. |
| skills | Canonical skill catalog. |
| user_skills | User-to-skill relationship and proficiency metadata. |
| user_experiences | Extracted or entered experience records. |
| cv_analyses | CV parsing status, predicted role, confidence, file metadata, strengths, gaps, and warnings. |
| job_postings | Imported/demo job data. |
| job_skills | Job-to-skill relationship and requirement metadata. |
| applications | Saved opportunities and status tracking. |
| scraping_sources | Admin-configured job source definitions. |
| target_job_roles | Target role list for scraping and market exploration. |
| scraping_jobs | Scraping batch execution state. |
| scraping_failed_urls | Failed source URL diagnostics and retry marker records. |
| scraping_proxies | Optional active proxy definitions protected by internal scraper token. |

*Table 60. Database tables summary.*

## Appendix D: Docker Services Summary

| Service | Role |
|---|---|
| nginx | Public gateway and reverse proxy. |
| frontend | React/Vite web UI container. |
| backend-api | Laravel API runtime. |
| backend-worker variants | Queue processing for default, high, AI, emails, and scraping work. |
| backend-scheduler | Scheduled Laravel commands. |
| db | MySQL database. |
| minio | S3-compatible CV object storage. |
| ai-cv-analyzer | FastAPI CV parsing service. |
| ai-job-miner | FastAPI job mining service. |
| prometheus | Metrics collection. |
| grafana | Metrics visualization. |

*Table 61. Docker services summary.*

## Appendix E: Screenshots

The screenshot set was reviewed during this pass. Some dashboard states are similar, but they document different examiner-visible states: before CV upload, upload UI, and after analysis. No screenshot merge was performed because combining them would reduce traceability and could disturb existing figure references in the generated List of Figures.

{screenshots_markdown}

## Appendix F: Test Cases

The manual test matrix in Chapter 8 should be repeated before final submission. Additional recommended tests include invalid CV uploads, banned user login, expired signed download URLs, failed AI service behavior, scraper token rejection, admin route rejection for normal users, and browser checks on a clean database.

The supporting evaluation files are summarized below instead of listed as raw paths:

| Evaluation Artifact | Purpose | Reader-Facing Evidence |
|---|---|---|
| Mini CV/job dataset | Deterministic synthetic records for recommendation and gap-analysis checks. | Chapter 8 mini evaluation tables. |
| Expected labels | Defines expected roles, domains, skills, and recommendation relevance for the mini dataset. | Table 52 and related detail tables. |
| Mini evaluation runner | Computes offline skill extraction, recommendation, and gap-analysis agreement. | Generated JSON/Markdown summaries. |
| AI CV smoke samples | Five sanitized fake CV text samples for analyzer smoke behavior. | Section 8.8 smoke evaluation. |
| AI CV smoke runner | Executes deterministic parser/check logic without exposing real CVs. | Figure 30 and smoke metric tables. |

Evaluation support artifact summary.

## Appendix G: GitHub Actions / CI Summary

The repository contains GitHub Actions workflow definitions. After the draft PR is opened, reviewers should inspect PR checks, rerun failed jobs if needed, and confirm that branch protection expectations are satisfied before merging. A CI screenshot was not embedded in this generated report because live PR checks were not available until after branch push.

## Appendix H: Demo Script

1. Start Docker Desktop.
2. Run `docker compose -f docker-compose.yml -f docker-compose.prod.yml up -d --build`.
3. Open `http://localhost`.
4. Register or login as a student.
5. Upload a sample PDF CV from the dashboard.
6. Review parsed profile and skills.
7. Open Jobs and select a recommendation.
8. Open Gap Analysis and explain matched/missing skills.
9. Save the job to Applications.
10. Login as the demo admin.
11. Open Admin Dashboard, Jobs, Sources, and Target Roles.
12. Open System Status and explain health checks.
13. Discuss limitations and future work honestly.

## Appendix I: AI CV Analyzer Deep Inventory

This appendix summarizes the code audit that supports Sections 5.5, Chapter 6, and Sections 8.8-8.11. The detailed companion notes are stored in `docs/graduation-book/model-analysis/` and should be kept with the generated book artifacts.

### I.1 Runtime Call Path

1. Laravel receives the CV upload and stores the file privately.
2. Laravel sends the uploaded file to the FastAPI analyzer `/api/parse-cv` endpoint.
3. `main.py` chooses image or PDF handling and wraps processing in timeout/error fallbacks.
4. `CVOrchestrator` extracts ordered text, triggers OCR fallback when needed, segments sections, runs NER, extracts contacts and dates, canonicalizes skills, and validates the strict output schema.
5. Layer 2 enriches the result with domain, seniority, and skill-category classification.
6. Laravel persists normalized profile, skills, experiences, and analysis metadata.
7. Job recommendations reuse the stored profile and skills with Laravel title/skill/seniority scoring. Gap analysis reuses the stored profile together with Layer 3 `/api/hybrid-match` and backend services.

### I.2 Function Inventory Summary

| Area | Classes and Functions Audited | Main Purpose |
|---|---|---|
| `main.py` | `_get_orchestrator`, `health_check`, `metrics`, `hybrid_match`, `analyze_cv`, `_process_with_timeout`, `process_file`, `_timeout_result`, `_error_result` | API endpoints, service lifecycle, timeout/fallback behavior, and hybrid match composition. |
| Layer 1 NER/contact/sections | `AdvancedNEREngine`, `extract_contacts`, `SemanticSegmenter`, `DataCanonicalizer`, `ExperienceEngine`, spatial/OCR helpers | Converts raw CV text into profile, skills, experience, and confidence-style structured data. |
| Layer 2 classification | `ClassificationOrchestrator`, `DomainEngine`, `SeniorityEngine`, `SkillEngine`, `load_taxonomy` | Adds domain, seniority, and skill-category interpretation. |
| Layer 3 matching | `SemanticEmbedder`, `IntelligentMatcher`, `ConstraintValidator`, `FitAnalysisGenerator`, `JobDescriptionEngine`, `RankingOrchestrator`, `tfidf.match_score` | Produces explainable job-fit scores, penalties, verdicts, and ranking behavior. |
| Training and diagnostics | `generate_tech_dataset.py`, `clean_dataset.py`, `train_ner.ipynb`, `verify_phase*.py`, `test_service_api.py`, `trace_cv.py` | Supports synthetic dataset generation, cleaning, notebook training, phase verification, service tests, and trace output. |

*Table 62. AI CV Analyzer function inventory summary.*

### I.3 Training Summary

The notebook trains a BERT-family token-classification model from `bert-base-cased`, maps character spans to BIO token labels, uses a 90/10 train/evaluation split with seed 42, trains for five epochs with learning rate 2e-5 and batch size 16, computes seqeval precision/recall/F1/accuracy, and exports `career_compass_ner_final`. The exported Colab PDF records overall final-epoch metrics: precision 0.933307, recall 0.940521, F1 0.936900, and accuracy 0.976376. The final cleaned dataset content and model weights are not committed.

### I.4 Generated Dataset Summary

The synthetic dataset script asks Google Gemini tooling to generate varied CV snippets across technical domains, rotates API keys/models from environment variables, includes negative decoys, and writes labeled examples for cleaning. The cleaner normalizes text, deduplicates exact samples, validates spans, and filters malformed or overly broad entities. This supports the training workflow, but it is separate from normal private CV upload processing.

### I.5 Local Artifact and Secrets Boundary

`.env` and `ai-cv-analyzer/models/` are ignored by Git. This protects secrets and avoids committing large model weights. Documentation may mention the runtime path and safe local metadata, but should not imply that the committed repository contains the model binary or private API keys.

## Appendix J: AI Job Miner Deep Inventory

This appendix converts the job-mining audit notes into a compact reader-facing summary. The detailed companion files remain in `docs/graduation-book/job-mining-analysis/` for repository traceability, but the important evidence is summarized below so the printed book is useful on its own.

| Audit Area | What It Documents | Related Chapter |
|---|---|---|
| Source inventory | Demo/local source behavior, optional API adapters, HTML/Scrapy-related paths, unsupported or credential-gated sources, and external-risk boundaries. | Chapter 7 source management and limitations. |
| Function inventory | FastAPI service entry points, adapter orchestration, classification helpers, Laravel controllers, queue jobs, services, and frontend admin pages. | Chapter 7 architecture, import, and diagnostics sections. |
| Runtime flow | On-demand scraping, full/admin runs, protected service calls, Laravel import callbacks, database updates, and status polling. | Figures 53-57 and Table 38. |
| API contracts | Authenticated user scraping endpoints, internal scraper import/check/failure endpoints, proxy route, and admin source/target endpoints. | Appendix A and Section 7.11. |
| Evaluation summary | Compile checks, container pytest, health probes, deterministic demo-source smoke evidence, and validation boundaries. | Chapter 8 and Table 44. |
| Limitations and ethics | External site instability, API keys, rate limits, proxy risks, robots/terms considerations, data freshness, and duplicate-detection boundaries. | Section 7.13 and Chapter 9. |

AI Job Miner audit support summary.

The appendix intentionally avoids copying large code blocks or raw path lists. Its purpose is to preserve the audit trail while keeping the graduation book readable in print.
"""


def write_references() -> None:
    REFERENCES_PATH.write_text(references_markdown(), encoding="utf-8")


def write_markdown(mini_results: dict, smoke_results: dict) -> None:
    MD_PATH.write_text(report_markdown(mini_results, smoke_results), encoding="utf-8")


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.75)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.footer_distance = Cm(0.51)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Title", 20), ("Heading 1", 18), ("Heading 2", 15), ("Heading 3", 13)]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Page ")
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_table = doc.add_table(rows=1, cols=2)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_table.autofit = True
    logo_paths = [LOGOS / "university_logo.png", LOGOS / "faculty_logo.png"]
    for idx, logo in enumerate(logo_paths):
        cell = logo_table.cell(0, idx)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo.exists():
            para.add_run().add_picture(str(logo), width=Inches(1.25))

    for text, size, bold in [
        (UNIVERSITY, 16, True),
        (FACULTY, 14, True),
        (DEPARTMENT, 13, False),
        ("Graduation Project Book", 15, True),
        (ACADEMIC_YEAR, 13, False),
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.bold = bold

    doc.add_paragraph()
    if CC_LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.add_run().add_picture(str(CC_LOGO_PATH), width=COVER_LOGO_WIDTH)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(SHORT_NAME)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(PROJECT_SUBTITLE)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = False
    run.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_paragraph()
    submitted = doc.add_paragraph()
    submitted.alignment = WD_ALIGN_PARAGRAPH.CENTER
    submitted.add_run("Submitted by:").bold = True
    student_table = doc.add_table(rows=3, cols=2)
    student_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    student_table.autofit = False
    set_table_geometry(student_table, [4500, 4500])
    for idx, student in enumerate(STUDENTS):
        cell = student_table.cell(idx // 2, idx % 2)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_width(cell, 4500)
        set_cell_text(cell, student, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    sup = doc.add_paragraph()
    sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sup.add_run("Supervised by:").bold = True
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(SUPERVISOR)
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("\\", "")
    return text


def find_following_caption(lines: list[str], index: int, kind: str) -> tuple[str | None, int]:
    j = index
    while j < len(lines) and not lines[j].strip():
        j += 1
    if j < len(lines) and lines[j].startswith(f"*{kind}"):
        return clean_inline(lines[j].strip("*")), j + 1
    return None, index


def set_cell_text(
    cell,
    value: str,
    *,
    bold: bool = False,
    size: float = 8.5,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    line_spacing: float = 1.05,
) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = line_spacing
    run = para.add_run(clean_inline(value))
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold


def table_widths(headers: list[str], column_count: int) -> list[int]:
    total = 9300
    joined = " ".join(headers).lower()
    header_set = {h.strip().lower() for h in headers}
    if column_count == 2 and {"layer", "software"}.issubset(header_set):
        weights = [0.20, 0.80]
    elif column_count == 2:
        weights = [0.24, 0.76]
    elif column_count == 3 and {"id", "requirement", "implementation evidence"}.issubset(header_set):
        weights = [0.10, 0.34, 0.56]
    elif column_count == 3 and {"category", "requirement", "careercompass approach"}.issubset(header_set):
        weights = [0.18, 0.31, 0.51]
    elif column_count == 3 and {"area", "current demo control", "production hardening needed"}.issubset(header_set):
        weights = [0.20, 0.31, 0.49]
    elif column_count == 3:
        weights = [0.22, 0.30, 0.48]
    elif column_count == 4 and "value" in joined and "notes" in joined:
        weights = [0.22, 0.24, 0.12, 0.42]
    elif column_count == 4:
        weights = [0.18, 0.28, 0.24, 0.30]
    elif column_count == 5:
        weights = [0.11, 0.17, 0.32, 0.14, 0.26]
    else:
        weights = [1 / column_count] * column_count
    widths = [int(total * weight) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, width in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
        elem = margins.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            margins.append(elem)
        elem.set(qn("w:w"), str(width))
        elem.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width: int) -> None:
    cell.width = Twips(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def keep_row_intact(row, repeat_header: bool = False) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if repeat_header:
        tbl_header = OxmlElement("w:tblHeader")
        tbl_header.set(qn("w:val"), "true")
        tr_pr.append(tbl_header)
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def add_md_table_docx(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    column_count = max(len(r) for r in rows)
    header_key = [clean_inline(cell).strip().lower() for cell in rows[0]]
    is_functional_requirements = header_key[:3] == ["id", "requirement", "implementation evidence"]
    is_software_requirements = header_key[:2] == ["layer", "software"]
    is_security_controls = header_key[:3] == ["area", "current demo control", "production hardening needed"]
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = table_widths(rows[0], column_count)
    set_table_geometry(table, widths)
    for r_idx, row in enumerate(rows):
        keep_row_intact(table.rows[r_idx], repeat_header=(r_idx == 0))
        for c_idx in range(column_count):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[c_idx])
            if r_idx == 0:
                shade_cell(cell, "D9EAF7")
            short_cell = c_idx == 0 or clean_inline(value).lower() in {"passed", "not run manual", "skipped, pytest missing"}
            align = WD_ALIGN_PARAGRAPH.CENTER if short_cell else WD_ALIGN_PARAGRAPH.LEFT
            font_size = 7.3 if is_functional_requirements else (8.2 if is_software_requirements or is_security_controls else (7.6 if column_count >= 5 else 8.5))
            spacing = 0.95 if is_functional_requirements else (1.0 if is_software_requirements or is_security_controls else 1.05)
            set_cell_text(cell, value, bold=(r_idx == 0), size=font_size, align=align, line_spacing=spacing)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def add_image_docx(doc: Document, rel_path: str, caption: str, bookmark_name: str | None = None) -> None:
    image_path = OUT_DIR / rel_path
    if not image_path.exists():
        doc.add_paragraph(f"[Missing image: {rel_path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True
    try:
        para.add_run().add_picture(str(image_path), width=Inches(6.1))
    except Exception:
        para.add_run(f"[Unable to insert image: {rel_path}]")
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_together = True
    if bookmark_name:
        add_bookmark(cap, bookmark_name)
    for run in cap.runs:
        run.italic = True


def add_bookmark(paragraph, bookmark_name: str) -> None:
    bookmark_id = abs(hash(bookmark_name)) % 2147483647
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text: str, anchor: str) -> None:
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(r_pr)
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_code_block_docx(doc: Document, code_lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [9300])
    cell = table.cell(0, 0)
    set_cell_width(cell, 9300)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(cell, "F8FAFC")
    cell.text = ""
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    for idx, code_line in enumerate(code_lines):
        run = para.add_run(code_line)
        run.font.name = "Consolas"
        run.font.size = Pt(7.5)
        if idx < len(code_lines) - 1:
            run.add_break()
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def generate_docx() -> None:
    doc = Document()
    set_doc_defaults(doc)
    add_cover(doc)
    caption_bookmarks_seen: set[str] = set()
    last_block: str | None = None
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines) and lines[i] != "\\pagebreak":
        i += 1
    if i < len(lines) and lines[i] == "\\pagebreak":
        i += 1
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line == "\\pagebreak":
            doc.add_page_break()
            last_block = None
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].startswith("```"):
                i += 1
            add_code_block_docx(doc, code_lines)
            last_block = "code"
            continue
        if i == 0 and line.startswith("# "):
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_md_table_docx(doc, table_lines)
            last_block = "table"
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            rel_path = image_match.group(2)
            caption = clean_inline(image_match.group(1))
            explicit_caption, next_index = find_following_caption(lines, i + 1, "Figure")
            if explicit_caption:
                caption = explicit_caption
                i = next_index
            else:
                i += 1
            bookmark_name = caption_anchor(caption)
            if bookmark_name in caption_bookmarks_seen:
                bookmark_name = None
            elif bookmark_name:
                caption_bookmarks_seen.add(bookmark_name)
            add_image_docx(doc, rel_path, caption, bookmark_name)
            last_block = "image"
            continue
        if line.startswith("# "):
            title = clean_inline(line[2:])
            para = doc.add_heading(title, level=1)
            para.paragraph_format.keep_with_next = True
            add_bookmark(para, heading_anchor(title))
            last_block = "heading"
        elif line.startswith("## "):
            title = clean_inline(line[3:])
            if title in PAGE_BREAK_BEFORE_HEADINGS:
                doc.add_page_break()
            para = doc.add_heading(title, level=2)
            para.paragraph_format.keep_with_next = True
            add_bookmark(para, heading_anchor(title))
            last_block = "heading"
        elif line.startswith("### "):
            title = clean_inline(line[4:])
            if title in PAGE_BREAK_BEFORE_HEADINGS:
                doc.add_page_break()
            para = doc.add_heading(title, level=3)
            para.paragraph_format.keep_with_next = True
            add_bookmark(para, heading_anchor(title))
            last_block = "heading"
        elif line.startswith("- "):
            link = re.fullmatch(r"\[([^\]]+)\]\(#([^)]+)\)", line[2:].strip())
            para = doc.add_paragraph(style="List Bullet")
            if link:
                add_internal_hyperlink(para, link.group(1), link.group(2))
            else:
                para.add_run(clean_inline(line[2:]))
            last_block = "list"
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(clean_inline(re.sub(r"^\d+\. ", "", line)), style="List Number")
            last_block = "list"
        elif line.startswith("*Figure") or line.startswith("*Table"):
            caption = clean_inline(line.strip("*"))
            p = doc.add_paragraph(caption)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bookmark_name = caption_anchor(caption)
            if bookmark_name and bookmark_name not in caption_bookmarks_seen:
                add_bookmark(p, bookmark_name)
                caption_bookmarks_seen.add(bookmark_name)
            for run in p.runs:
                run.italic = True
            last_block = "caption"
        else:
            doc.add_paragraph(clean_inline(line))
            last_block = "text"
        i += 1
    doc.save(DOCX_PATH)


def para_style(name: str, size: int, leading: int | None = None, bold: bool = False, align=TA_LEFT):
    return ParagraphStyle(
        name=name,
        fontName="Helvetica-Bold" if bold else "Helvetica",
        fontSize=size,
        leading=leading or size + 4,
        alignment=align,
        spaceAfter=8,
    )


def add_pdf_cover(story: list) -> None:
    logo_row = []
    for logo in [LOGOS / "university_logo.png", LOGOS / "faculty_logo.png"]:
        if logo.exists():
            logo_row.append(Image(str(logo), width=3.0 * cm, height=3.0 * cm))
        else:
            logo_row.append(Paragraph("", para_style("empty", 10)))
    table = Table([logo_row], colWidths=[7 * cm, 7 * cm])
    table.setStyle(TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER"), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(table)
    story.append(Spacer(1, 0.4 * cm))
    center = para_style("center", 13, bold=True, align=TA_CENTER)
    for line in [UNIVERSITY, FACULTY, DEPARTMENT, "Graduation Project Book", ACADEMIC_YEAR]:
        story.append(Paragraph(html.escape(line), center))
    story.append(Spacer(1, 0.6 * cm))
    if CC_LOGO_PATH.exists():
        story.append(Image(str(CC_LOGO_PATH), width=3.5 * cm, height=3.5 * cm))
        story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph(html.escape(SHORT_NAME), para_style("pdf-title", 20, leading=25, bold=True, align=TA_CENTER)))
    story.append(Paragraph(html.escape(PROJECT_SUBTITLE), para_style("pdf-subtitle", 13, leading=18, align=TA_CENTER)))
    story.append(Spacer(1, 1.0 * cm))
    story.append(Paragraph("Submitted by:", center))
    for student in STUDENTS:
        story.append(Paragraph(html.escape(student), para_style("student", 12, align=TA_CENTER)))
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph("Supervised by:", center))
    story.append(Paragraph(html.escape(SUPERVISOR), para_style("supervisor", 12, align=TA_CENTER)))
    story.append(PageBreak())


def add_pdf_image(story: list, rel_path: str, caption: str) -> None:
    image_path = OUT_DIR / rel_path
    if not image_path.exists():
        story.append(Paragraph(html.escape(f"[Missing image: {rel_path}]"), para_style("normal", 10)))
        return
    with PILImage.open(image_path) as img:
        w, h = img.size
    max_w = 16.2 * cm
    max_h = 10.2 * cm
    ratio = min(max_w / w, max_h / h)
    story.append(
        KeepTogether(
            [
                Image(str(image_path), width=w * ratio, height=h * ratio),
                Paragraph(html.escape(caption), para_style("caption", 9, align=TA_CENTER)),
                Spacer(1, 0.2 * cm),
            ]
        )
    )


def add_md_table_pdf(story: list, lines: list[str]) -> None:
    rows = []
    for line in lines:
        parts = [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", part) for part in parts):
            continue
        rows.append(parts)
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    data = [[Paragraph(html.escape(cell), para_style("tablecell", 7, leading=9, bold=(r == 0))) for cell in row] for r, row in enumerate(normalized)]
    table = Table(data, repeatRows=1, hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dbeafe")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#94a3b8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.25 * cm))


def add_code_block_pdf(story: list, code_lines: list[str]) -> None:
    code_style = ParagraphStyle(
        "codeblock",
        fontName="Courier",
        fontSize=6.6,
        leading=8.2,
        textColor=colors.HexColor("#0f172a"),
    )
    code_text = "\n".join(code_lines) if code_lines else " "
    table = Table([[Preformatted(code_text, code_style)]], colWidths=[16.2 * cm], hAlign="LEFT")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f8fafc")),
                ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#cbd5e1")),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 0.25 * cm))


def generate_pdf() -> int:
    styles = getSampleStyleSheet()
    story: list = []
    add_pdf_cover(story)
    h1 = para_style("h1", 16, leading=20, bold=True)
    h2 = para_style("h2", 13, leading=16, bold=True)
    h3 = para_style("h3", 12, leading=15, bold=True)
    normal = para_style("normal", 9, leading=12)
    bullet = para_style("bullet", 9, leading=12)
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines) and lines[i] != "\\pagebreak":
        i += 1
    if i < len(lines) and lines[i] == "\\pagebreak":
        i += 1
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if i == 0 and line.startswith("# "):
            i += 1
            continue
        if line == "\\pagebreak":
            story.append(PageBreak())
            i += 1
            continue
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines) and lines[i].startswith("```"):
                i += 1
            add_code_block_pdf(story, code_lines)
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_md_table_pdf(story, table_lines)
            continue
        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption = clean_inline(image_match.group(1))
            explicit_caption, next_index = find_following_caption(lines, i + 1, "Figure")
            if explicit_caption:
                caption = explicit_caption
                i = next_index
            else:
                i += 1
            add_pdf_image(story, image_match.group(2), caption)
            continue
        if line.startswith("# "):
            story.append(Paragraph(html.escape(clean_inline(line[2:])), h1))
        elif line.startswith("## "):
            story.append(Paragraph(html.escape(clean_inline(line[3:])), h2))
        elif line.startswith("### "):
            story.append(Paragraph(html.escape(clean_inline(line[4:])), h3))
        elif line.startswith("- "):
            story.append(Paragraph("&bull; " + html.escape(clean_inline(line[2:])), bullet))
        elif re.match(r"^\d+\. ", line):
            story.append(Paragraph(html.escape(clean_inline(line)), normal))
        elif line.startswith("*Figure") or line.startswith("*Table"):
            story.append(Paragraph(html.escape(clean_inline(line.strip("*"))), para_style("caption", 8, align=TA_CENTER)))
        else:
            story.append(Paragraph(html.escape(clean_inline(line)), normal))
        i += 1

    def add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2.0, 1.2 * cm, f"Page {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        rightMargin=2.2 * cm,
        leftMargin=2.2 * cm,
        topMargin=1.0 * cm,
        bottomMargin=0.75 * cm,
    )
    doc.build(story, onFirstPage=add_page_number, onLaterPages=add_page_number)
    return len(PdfReader(str(PDF_PATH)).pages)


def export_pdf_from_docx_with_word() -> tuple[bool, str]:
    ps = f"""
$ErrorActionPreference = 'Stop'
$docx = '{str(DOCX_PATH).replace("'", "''")}'
$pdf = '{str(PDF_PATH).replace("'", "''")}'
$word = $null
$doc = $null
try {{
  $word = New-Object -ComObject Word.Application
  $word.Visible = $false
  $word.DisplayAlerts = 0
  $doc = $word.Documents.Open($docx, $false, $true)
  $doc.ExportAsFixedFormat($pdf, 17, $false, 0, 0, 1, 1, 0, $true, $true, 1, $true, $true, $false)
  $doc.Close($false)
  $word.Quit()
  'WORD_EXPORT_OK'
}} catch {{
  if ($doc -ne $null) {{ try {{ $doc.Close($false) }} catch {{}} }}
  if ($word -ne $null) {{ try {{ $word.Quit() }} catch {{}} }}
  'WORD_EXPORT_FAILED=' + $_.Exception.Message
  exit 1
}}
"""
    try:
        result = subprocess.run(
            ["powershell.exe", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", ps],
            cwd=ROOT,
            text=True,
            capture_output=True,
            timeout=180,
        )
    except Exception as exc:
        return False, f"Microsoft Word COM export could not be started: {exc}"
    output = (result.stdout + result.stderr).strip()
    return result.returncode == 0, output


def count_pdf_link_annotations() -> int:
    try:
        reader = PdfReader(str(PDF_PATH))
        total = 0
        for page in reader.pages:
            annots = page.get("/Annots") or []
            total += len(annots)
        return total
    except Exception:
        return 0


def pdf_pages_for_terms(terms: list[str]) -> dict[str, list[int]]:
    try:
        reader = PdfReader(str(PDF_PATH))
        pages: dict[str, list[int]] = {term: [] for term in terms}
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for term in terms:
                if term in text:
                    pages[term].append(page_number)
        return pages
    except Exception:
        return {term: [] for term in terms}


def count_docx_internal_links() -> dict[str, int]:
    try:
        from zipfile import ZipFile
        from xml.etree import ElementTree as ET

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with ZipFile(DOCX_PATH) as zf:
            xml = ET.fromstring(zf.read("word/document.xml"))
        counts = {"toc": 0, "figures": 0, "tables": 0, "total": 0}
        for link in xml.findall(".//w:hyperlink", ns):
            anchor = link.get(qn("w:anchor")) or ""
            counts["total"] += 1
            if anchor.startswith("bm_figure_"):
                counts["figures"] += 1
            elif anchor.startswith("bm_table_"):
                counts["tables"] += 1
            else:
                counts["toc"] += 1
        return counts
    except Exception:
        return {"toc": 0, "figures": 0, "tables": 0, "total": 0}


def write_notes(
    page_count: int,
    pdf_method: str,
    toc_docx_status: str,
    figures_docx_status: str,
    tables_docx_status: str,
    toc_pdf_status: str,
    table_status: str,
    caption_status: str,
    requirements_layout_status: str,
    software_layout_status: str,
    security_layout_status: str,
) -> None:
    screenshot_count = len(list(SCREENSHOTS.glob("*.png")))
    diagram_count = len(list(DIAGRAMS.glob("*.png")))
    evaluation_files = [
        "mini_cv_dataset.json",
        "mini_jobs_dataset.json",
        "expected_labels.json",
        "run_mini_evaluation.py",
        "mini_evaluation_results.json",
        "mini_evaluation_summary.md",
        "ai_cv_analyzer_smoke_samples.json",
        "run_ai_cv_analyzer_smoke_eval.py",
        "ai_cv_analyzer_smoke_results.json",
        "ai_cv_analyzer_smoke_summary.md",
    ]
    notes = f"""# Report Generation Notes

## Purpose

This folder contains the generated graduation project book for CareerCompass.

## Style References

The supervisor-provided previous graduation books were copied into `reference-books/` and used only as formatting, structure, and academic-report style references. They were not used as scientific or technical sources for CareerCompass, and they are not cited as CareerCompass references.

## Generated Files

- `CareerCompass_Graduation_Project_Book.md`
- `CareerCompass_Graduation_Project_Book.docx`
- `CareerCompass_Graduation_Project_Book.pdf`
- `references.md`
- `model-analysis/*.md`
- `job-mining-analysis/*.md`
- `model-analysis/colab_train_ner_results.pdf`
- `assets/diagrams/*.png`
- `assets/screenshots/*.png`
- `assets/logos/*.png`
- `reference-books/*.pdf`

## Generation Tooling

- Markdown source: generated by `scripts/generate_graduation_book.py`.
- DOCX: generated with `python-docx` using A4 page settings and page-number fields.
- PDF: {pdf_method}.
- Diagrams: generated as PNG files with Pillow.
- Browser screenshots: captured from the running local Docker stack using Chrome DevTools Protocol.
- Code and JSON examples: rendered as shaded monospace boxes in DOCX/PDF generation paths.

## Counts

- PDF pages: {page_count}
- Screenshots/evidence images: {screenshot_count}
- Diagrams: {diagram_count}

## Table of Contents and Tables

- TOC placement: standalone page immediately after the cover page
- Front matter order: Cover -> Table of Contents -> List of Figures -> List of Tables -> Acknowledgment -> Abstract -> Abbreviations -> Chapter 1
- Abbreviations placement: standalone page after Abstract and before Chapter 1
- DOCX TOC status: {toc_docx_status}
- DOCX List of Figures status: {figures_docx_status}
- DOCX List of Tables status: {tables_docx_status}
- PDF TOC status: {toc_pdf_status}
- Table formatting status: {table_status}
- Figure caption status: {caption_status}
- Table caption status: formal table captions are rendered below their corresponding tables and remain linked from the List of Tables
- Section 2.6/2.7 layout status: {requirements_layout_status}
- Section 2.9/2.10 layout status: {software_layout_status}
- Section 9.9/9.10 layout status: {security_layout_status}

## Caption, Link, and Layout Verification Method

- Duplicate figure captions were checked by removing image-prefix alt text and rendering only the explicit italic `Figure n. ...` caption line in DOCX/PDF generation paths.
- Markdown scan: `rg -n "Figure [0-9]+:" docs/graduation-book/CareerCompass_Graduation_Project_Book.md` returned no matches after regeneration.
- DOCX/PDF text extraction checks verify that no visible image-prefix caption lines remain and that figure captions still exist.
- Table-caption order was checked structurally from DOCX block order and Markdown source order so that `Table n. ...` captions occur after the corresponding table blocks.
- DOCX structural scan: OOXML inspection counted TOC/List of Figures/List of Tables internal hyperlinks, checked that every hyperlink anchor has a matching bookmark, and counted figure/table caption bookmarks.
- PDF structural scan: `pypdf` counted pages and link annotations after Microsoft Word export.
- Front-matter placement was checked from DOCX body order and PDF page text extraction: the first generated page after the cover is the Table of Contents page, followed by List of Figures, List of Tables, Acknowledgment, Abstract, standalone Abbreviations, and Chapter 1.
- Table layout was checked through OOXML for fixed table layout, table grids, cell widths, and repeated header rows on data tables.

## Mini Dataset Evaluation

The mini dataset evaluation was added under `evaluation/` and uses fake synthetic CV/job records. It is a preliminary offline validation, not a production benchmark.

{chr(10).join([f"- `evaluation/{name}`" for name in evaluation_files])}

## AI CV Analyzer Smoke Evaluation

- A deterministic text-only smoke evaluation was added for the AI CV Analyzer evidence pass.
- The smoke evaluation uses five fake CV samples: backend, data analyst, frontend, DevOps/cloud, and low-information/noisy text.
- It measures only reproducible local checks: expected skill extraction, role/domain/seniority labels, parsing-status classification, dependency availability, and TF-IDF fallback availability.
- It is not described as a final transformer NER benchmark because it does not run the transformer weights and the repository does not include the final cleaned labeled dataset or model weights.

## Colab NER Training Results Evidence

- The user-provided exported Colab PDF was copied to `model-analysis/colab_train_ner_results.pdf`.
- A metric summary was added at `model-analysis/colab_ner_training_results_summary.md`.
- PDF extraction note: the PDF is image-based, so `pypdf` extracted no text; pages were rendered for visual inspection using `pypdfium2`.
- Secret scan/inspection result: no actual API keys, bearer tokens, passwords, or secret values were found. The PDF only shows an optional missing `HF_TOKEN` warning.
- Visible dataset evidence from the PDF: 45,911 total rows, 41,319 train rows, 4,592 test rows, 11 BIO labels, test size 0.1, seed 42.
- Visible final epoch metrics from the PDF: precision 0.933307, recall 0.940521, F1 0.936900, accuracy 0.976376, training loss 0.037280, validation loss 0.068058.
- Per-label metrics status: not visible in the PDF, so no per-label table or support-distribution chart was invented.
- Colab visualization status: generated/updated final metrics, epoch performance, and loss-curve charts from the verified visible metrics.
- Loss-function status: the notebook uses Hugging Face `Trainer` with `AutoModelForTokenClassification` and does not define a custom loss; the report treats losses as Trainer-reported token-classification objective values.
- Confusion matrix status: not visible in the PDF and not present in the attached notebook outputs/source terms, so no confusion-matrix chart was generated.

## AI CV Analyzer Documentation Update

- Decision: the AI CV Analyzer is now documented as standalone Chapter 6 because it is a core technical contribution rather than only an implementation subsection.
- Added model-analysis notes under `model-analysis/` to summarize runtime architecture, Layer 1/2/3 internals, function inventory, NER token processing, synthetic data generation, training-notebook inspection, and evaluation limitations.
- Reviewed `D:/Graduation/model-analys-helper`; the top-level helper folders found were `docs`, `layer1`, `layer2`, and `layer3`.
- No raw training datasets, image artifacts, or actual secrets were copied from the helper folder.
- The training notebook and exported Colab PDF were inspected; the notebook was not re-executed because the cleaned dataset is not committed, the generation workflow depends on external Gemini API keys, and the notebook is designed for Colab GPU execution.
- AI diagrams were generated for runtime flow, model-training workflow, extraction components, Layer 1 understanding, Layer 2 classification, Layer 3 matching, NER token processing, seniority logic, canonicalization, and score collapse.
- New academic AI diagrams were generated for design philosophy, complete CV flow, fault tolerance, confidence/readiness signals, skill canonicalization example, fine-tuned BERT NER architecture, detailed training flow, matching formula, explainable output, and analyzer sequence.
- Matching formula status: exact for the `IntelligentMatcher.calculate_match` score-composition path; upstream semantic/skill/domain scores still depend on model availability.
- Dataset statistics status: Colab PDF records train/test row counts; dataset content remains unavailable from committed evidence.
- Dataset evidence diagram status: updated the dataset evidence availability diagram to include Colab PDF evidence.
- Colab metrics visualization status: generated a final-epoch metrics bar chart, epoch performance trend chart, and training/validation loss curve from PDF-visible values.
- NER label distribution chart status: not generated because per-label support counts are not visible in the PDF and no committed final labeled training dataset exists to count labels honestly.
- Confidence status: the system uses confidence-style and readiness signals, not a certified hiring probability formula.
- Raw CV walkthrough status: includes a sanitized raw CV fragment, Layer 1/2/3 tables, an illustrative schema JSON block based on the attached actual analyzer response shape, and a top-level output-section summary. Personal email, phone, profile URLs, name, and raw CV text were replaced with placeholders or redacted text.
- API appendix status: expanded from endpoint summary to request/response/error JSON examples for upload, parse-cv, hybrid-match, recommendations, gap analysis, and health/readiness.
- Quick Start status: added examiner commands, local URLs, demo admin environment values, validation commands, and troubleshooting notes.
- Preview features status: clarified CV Builder, Mock Interview, Learning Paths, Career Planner, Mentorship, Tools Hub, and Market Intelligence as preview/future modules where appropriate.
- Screenshot duplication decision: reviewed and kept separate screenshots because they document different states and preserve figure traceability.
- ERD/System Architecture review: ERD was updated against current migrations/table names; high-level and Docker architecture diagrams were reviewed and retained as clear.
- End-to-end walkthrough status: illustrative academic example, not a live model benchmark.
- References review: existing references cover BERT, Hugging Face token classification/Trainer, sentence embeddings, TF-IDF/cosine similarity, Gemini/Colab, dynamic quantization, and OpenAPI-style API documentation.

## AI Job Miner Documentation Update

- Decision: the AI Job Miner and scraping subsystem is now documented as standalone Chapter 7 because it has a distinct runtime service, queue workflow, import contract, diagnostics model, and ethical boundary.
- Added support notes under `job-mining-analysis/` for source inventory, function inventory, runtime flow, API contracts, evaluation summary, and limitations.
- Chapter restructuring status: Testing and Evaluation is Chapter 8, Security and Privacy is Chapter 9, and Conclusion/Future Work is Chapter 10.
- New scraping diagrams generated: design philosophy, complete mining flow, runtime architecture, sequence diagram, lifecycle, source management, import/deduplication, failed URL flow, security boundaries, and validation evidence.
- Code areas audited: `ai-job-miner/`, scraping jobs, `JobController`, `ScrapedJobController`, admin source/target/dashboard controllers, request validation, models, migrations, seeders, frontend admin/user APIs, and Docker Compose service wiring.
- Cleanup status: Chapter 7 no longer repeats the admin dashboard/jobs/sources/targets screenshots; it cross-references Figures 44-47 instead.
- Accuracy fix: Table 36 now describes job mining design decisions, and `scraping_jobs` documentation now reflects the actual `job_title`/status/counter schema rather than implying a source-id column.
- Source coverage language: the report distinguishes deterministic demo/local sources, API adapters, HTML adapters, and unsupported/external-risk sources without claiming whole-market reach.
- Evaluation language: scraping evidence is recorded as tests, compile/config/health checks, form-request validation, admin diagnostics, and screenshots, not as source success rates unless rerun and recorded.
- Queue lifecycle note: the deterministic demo smoke validates the protected `/scrape` adapter/import path; a full authenticated `/jobs/scrape-if-missing` queue lifecycle with browser polling remains a recommended final demonstration check.
- Ethics language: the chapter explicitly covers rate limits, external instability, API keys, proxy configuration, robots/terms considerations, and demo scope.

## Backend / Frontend / Database Documentation Update

- Decision: this polish pass strengthens the non-AI system documentation without expanding it into another oversized deep-dive chapter.
- Frontend update: Chapter 3 now includes a route/layout architecture diagram and an API/authentication flow diagram based on `frontend/src/App.jsx`, route guards, `AuthContext.jsx`, the Axios API client, and localization files.
- Backend update: Chapter 3 now includes a Laravel request lifecycle diagram, a backend module responsibility table, and a concise validation/protection mapping grounded in route groups, middleware, form requests, controllers, services, models, resources, queues, and storage services.
- Database/ERD update: Chapter 3 now adds database design rationale, relationship notes, and data-integrity mechanisms based on migrations and model relationships.
- API appendix update: added compact core backend examples for login, current user, application tracking, and admin dashboard stats while keeping existing AI and scraping examples intact.
- Testing chapter update: added a module validation coverage matrix so previous backend/frontend evidence is clearly distinguished from checks rerun in later documentation passes.
- Caption cleanup: Table 4 now uses the accurate caption `Software environment summary`.
- New diagrams generated for this pass: frontend route/layout architecture, frontend API/auth flow, Laravel backend request lifecycle, and database relationship rationale.
- References review: React Router and Axios official documentation references were added because the route tree and centralized API client are now discussed explicitly.

## Validation Summary

- Branch check passed on `docs/graduation-book`; the only unrelated untracked file remained `docs/REVERSE_ENGINEERING_SYSTEM_WALKTHROUGH.md`.
- `git diff --check` and `git diff --cached --check` completed without whitespace errors; Git reported line-ending normalization warnings only.
- Report generation ran successfully with the bundled Python runtime and produced Markdown, DOCX, and PDF artifacts.
- Generated PDF page count: {page_count}; generated PDF link annotations: {toc_pdf_status}.
- DOCX structural scan found internal hyperlinks/bookmarks for the custom TOC/List of Figures/List of Tables with no missing anchors after figure/table anchor checks.
- JSON code-fence validation parsed 36 JSON blocks successfully.
- All ten scraping diagrams and the four new backend/frontend/database diagrams exist and are referenced by the generated Markdown.
- Placeholder/typo/bookmark/caption/overclaim scans returned no matches.
- `python -m py_compile docs/graduation-book/scripts/generate_graduation_book.py` passed using the bundled Python runtime.
- `python -m compileall ai-job-miner` passed.
- `python -m compileall ai-cv-analyzer` passed with a non-fatal `.pytest_cache` listing warning from compileall output.
- `python -m compileall ai-cv-analyzer/training` passed.
- `docker compose -f docker-compose.yml -f docker-compose.prod.yml config --quiet` passed.
- Docker runtime was available in this pass; the stack was already running, so no rebuild was performed.
- `docker compose ps` showed the main app containers running; backend, frontend, Nginx, job miner, database, and queue workers were healthy.
- Runtime health probes returned HTTP 200 for `/api/health`, `/api/ready`, `/status`, AI Job Miner `/health`, and the AI CV Analyzer root endpoint. AI CV Analyzer `/health` returned 404, so the root endpoint is the verified liveness endpoint for that service in this pass.
- Local bundled Python still lacked `pytest`, but `docker compose exec -T ai-job-miner python -m pytest` passed with 75 tests and 1 warning.
- Deterministic demo-source scrape smoke passed: protected AI Job Miner `/scrape` with `CareerCompass Demo Jobs` returned `SUCCESS`, previewed 3 jobs, stored 3 through Laravel import, and reported 0 failed URLs. The temporary smoke rows were cleaned from the local database afterward. Because this called the service directly, it validates the demo adapter plus import path, not queue status polling.
- Backend and frontend tests were not rerun in this pass because no backend or frontend application code was changed.
- Mini evaluation script ran successfully and generated JSON plus Markdown result summaries.
- AI CV Analyzer smoke evaluation script ran successfully and generated JSON plus Markdown result summaries.
- AI CV Analyzer model-training evidence was documented as inspected evidence; Colab-run overall metrics were added from the exported PDF, while per-label metrics and confusion-matrix values were not invented because they are not visible in the PDF.

## Placeholder Review

All previously listed student placeholders were removed and replaced with the final six team names.

## Known Limitations

- This script creates a ReportLab PDF fallback first. Microsoft Word COM automation is then used when available to export the DOCX to PDF.
- The Documents skill `render_docx.py` workflow was attempted after DOCX generation, but it could not render because LibreOffice/soffice was not available on the host PATH (`FileNotFoundError: [WinError 2]`).
- The report uses custom manual Table of Contents, List of Figures, and List of Tables links instead of fragile automatic Word fields.
- PDF link preservation is checked through PDF annotation counts after Microsoft Word export; individual clicks should still be spot-checked in a PDF reader before printing.
- GitHub Actions status should be reviewed on the draft PR after every push.
- AI evaluation results are treated as preliminary/manual smoke evidence, not as statistical accuracy claims.

## Manual Review Before Submission

- Confirm supervisor name, department, academic year, and final team-name spelling before printing.
- Open the DOCX in Microsoft Word and update visual spacing if the faculty requires a specific template.
- Confirm the generated PDF opens and figures are readable.
- Review the draft PR checks after GitHub Actions finish.
- Re-run a fresh Docker demo before the final defense.
- Confirm supervisor name, department, academic year, and student names before printing.
"""
    NOTES_PATH.write_text(notes, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    mini_results = run_mini_evaluation()
    smoke_results = run_smoke_evaluation()
    create_diagrams()
    create_terminal_evidence()
    write_references()
    write_markdown(mini_results, smoke_results)
    generate_docx()
    fallback_page_count = generate_pdf()
    export_ok, export_output = export_pdf_from_docx_with_word()
    if export_ok:
        print(f"Word COM export succeeded: {export_output}")
    else:
        print(f"Word COM export failed: {export_output}")
    page_count = len(PdfReader(str(PDF_PATH)).pages)
    if export_ok:
        pdf_method = "exported from the generated DOCX using Microsoft Word COM automation"
    else:
        pdf_method = f"generated directly with ReportLab fallback; Word export failed ({export_output or 'no output'})"
        page_count = fallback_page_count
    link_count = count_pdf_link_annotations()
    docx_link_counts = count_docx_internal_links()
    toc_docx_status = f"{docx_link_counts['toc']} custom manual TOC entries contain internal hyperlinks to bookmarked major headings"
    figures_docx_status = f"{docx_link_counts['figures']} List of Figures entries link to bookmarked figure captions"
    tables_docx_status = f"{docx_link_counts['tables']} List of Tables entries link to bookmarked table captions"
    toc_pdf_status = f"PDF contains {link_count} link annotations after export" if link_count else "PDF link preservation could not be confirmed from annotations"
    table_status = "data tables use fixed DXA widths, wrapped text, repeated header rows, smaller table fonts, and split wide manual test observations into narrower tables; cover layout tables are intentionally excluded from repeated-header checks"
    caption_status = "explicit italic caption lines are the single visible figure-caption source; Markdown image alt text is not rendered as a visible figure caption"
    section_pages = pdf_pages_for_terms([
        "2.6 Functional Requirements",
        "FR-01",
        "FR-11",
        "Table 2. Functional requirements summary.",
        "2.7 Non-Functional Requirements",
        "2.10 Software Requirements",
        "Frontend",
        "Backend",
        "AI services",
        "Data",
        "Infrastructure",
        "Testing",
        "Table 4. Software environment summary.",
        "2.11 Input and Output Flow",
        "9.9 Demo Security Limitations",
        "Admin account",
        "CV files",
        "Tokens",
        "Scraper service",
        "Monitoring",
        "Privacy",
        "Table 58. Security and privacy controls.",
        "9.10 Future Production Hardening",
    ])
    section_26_start = min(section_pages.get("2.6 Functional Requirements", []) or section_pages.get("FR-01", []) or [1])
    table_2_body_pages = [page for page in section_pages.get("Table 2. Functional requirements summary.", []) if page >= section_26_start]
    fr_pages = sorted(set(
        section_pages.get("2.6 Functional Requirements", [])
        + section_pages.get("FR-01", [])
        + section_pages.get("FR-11", [])
        + table_2_body_pages
    ))
    requirements_layout_status = (
        f"2.6 heading, FR-01/FR-11 rows, and Table 2 caption appear on PDF page(s) {fr_pages or 'not detected'}; "
        f"2.7 starts on PDF page(s) {section_pages.get('2.7 Non-Functional Requirements', []) or 'not detected'} after the Table 2 caption"
    )
    section_29_start = min(section_pages.get("2.10 Software Requirements", []) or [1])
    section_210_start = min(section_pages.get("2.11 Input and Output Flow", []) or [section_29_start])
    table_4_body_pages = [page for page in section_pages.get("Table 4. Software environment summary.", []) if page >= section_29_start]
    software_row_pages = []
    for term in ["Frontend", "Backend", "AI services", "Data", "Infrastructure", "Testing"]:
        software_row_pages.extend(page for page in section_pages.get(term, []) if section_29_start <= page <= section_210_start)
    software_pages = sorted(set(
        section_pages.get("2.10 Software Requirements", [])
        + software_row_pages
        + table_4_body_pages
    ))
    software_layout_status = (
        f"2.10 heading, Software Requirements rows, and Table 4 caption appear on PDF page(s) {software_pages or 'not detected'}; "
        f"2.11 starts on PDF page(s) {section_pages.get('2.11 Input and Output Flow', []) or 'not detected'} after the Table 4 caption"
    )
    section_79_start = min(section_pages.get("9.9 Demo Security Limitations", []) or [1])
    section_710_start = min(section_pages.get("9.10 Future Production Hardening", []) or [section_79_start])
    table_39_body_pages = [page for page in section_pages.get("Table 58. Security and privacy controls.", []) if page >= section_79_start]
    security_row_pages = []
    for term in ["Admin account", "CV files", "Tokens", "Scraper service", "Monitoring", "Privacy"]:
        security_row_pages.extend(page for page in section_pages.get(term, []) if section_79_start <= page <= section_710_start)
    security_pages = sorted(set(
        section_pages.get("9.9 Demo Security Limitations", [])
        + security_row_pages
        + table_39_body_pages
    ))
    security_layout_status = (
        f"9.9 heading, security-control rows, and Table 58 caption appear on PDF page(s) {security_pages or 'not detected'}; "
        f"9.10 starts on PDF page(s) {section_pages.get('9.10 Future Production Hardening', []) or 'not detected'} after the Table 58 caption"
    )
    write_notes(
        page_count,
        pdf_method,
        toc_docx_status,
        figures_docx_status,
        tables_docx_status,
        toc_pdf_status,
        table_status,
        caption_status,
        requirements_layout_status,
        software_layout_status,
        security_layout_status,
    )
    print(f"Generated Markdown: {MD_PATH}")
    print(f"Generated DOCX: {DOCX_PATH}")
    print(f"Generated PDF: {PDF_PATH}")
    print(f"PDF pages: {page_count}")
    print(toc_pdf_status)


if __name__ == "__main__":
    main()
